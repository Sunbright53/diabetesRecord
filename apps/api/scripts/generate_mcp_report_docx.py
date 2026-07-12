"""
สร้างรายงาน MCP (Model Context Protocol) — Cheewarun Health Platform
เชื่อมโยงกับ SmartBreath LSTM Training Data
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── สี ──────────────────────────────────────────────────────────────────────
C_NAVY    = RGBColor(0x1B, 0x4F, 0x72)
C_BLUE    = RGBColor(0x2E, 0x86, 0xC1)
C_LBLUE   = RGBColor(0xD6, 0xEA, 0xF8)
C_GREEN   = RGBColor(0x1E, 0x84, 0x49)
C_ORANGE  = RGBColor(0xCA, 0x6F, 0x1E)
C_RED     = RGBColor(0x92, 0x2B, 0x21)
C_GREY    = RGBColor(0x71, 0x7D, 0x7E)
C_GREY_R  = RGBColor(0xF2, 0xF3, 0xF4)
C_WARN    = RGBColor(0xFD, 0xEB, 0xD0)
C_PURPLE  = RGBColor(0x6C, 0x3D, 0xA8)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK   = RGBColor(0x1A, 0x1A, 0x1A)
C_TEAL    = RGBColor(0x00, 0x71, 0x77)

FONT = "TH SarabunPSK"


def set_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(rgb))
    tcPr.append(shd)


def rn(para, text, bold=False, size=10.5, color=None, italic=False, font=None):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = font or FONT
    r.font.color.rgb = color or C_BLACK
    rpr = r._r.get_or_add_rPr()
    f = OxmlElement("w:rFonts")
    fn = font or FONT
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        f.set(qn(attr), fn)
    rpr.append(f)
    return r


def H1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    rn(p, text, bold=True, size=15, color=C_NAVY)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "8")
    b.set(qn("w:space"), "1");    b.set(qn("w:color"), "2E86C1")
    pBdr.append(b); pPr.append(pBdr)
    return p


def H2(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    rn(p, text, bold=True, size=12.5, color=color or C_BLUE)
    return p


def H3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    rn(p, text, bold=True, size=11, color=C_NAVY)
    return p


def BD(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sa=5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(1)
    rn(p, text, size=10.5)
    return p


def NOTE(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    rn(p, text, size=9.5, italic=True, color=C_GREY)


def CODE(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    rn(p, text, size=9, font="Courier New", color=RGBColor(0x2C, 0x3E, 0x50))


def warn_box(doc, title, body_text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_bg(cell, C_WARN)
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    rn(p, f"⚠  {title}\n", bold=True, size=10.5, color=C_ORANGE)
    rn(p, body_text, size=10, color=RGBColor(0x5D, 0x40, 0x07))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def info_box(doc, title, body_text, bg=None, title_color=None):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_bg(cell, bg or C_LBLUE)
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    rn(p, f"{title}\n", bold=True, size=10.5, color=title_color or C_NAVY)
    rn(p, body_text, size=10, color=C_BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def tbl(doc, headers, rows, widths=None, hbg=None):
    n = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    hr = t.rows[0]
    for j, h in enumerate(headers):
        c = hr.cells[j]
        set_bg(c, hbg or C_NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        rn(p, h, bold=True, size=9.5, color=C_WHITE)

    for i, row in enumerate(rows):
        bg = C_GREY_R if i % 2 == 0 else C_WHITE
        tr = t.rows[i + 1]
        for j, val in enumerate(row):
            c = tr.cells[j]
            set_bg(c, bg)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)

            if isinstance(val, tuple):
                txt, col = val
            else:
                txt, col = str(val), None

            align = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            if widths and widths[j] <= 2.0:
                align = WD_ALIGN_PARAGRAPH.CENTER
            p.alignment = align
            rn(p, txt, size=9.5, color=col or C_BLACK)

    if widths:
        for row in t.rows:
            for j, c in enumerate(row.cells):
                c.width = Cm(widths[j])

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def SP(doc, n=0.2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(n * 28)


# ═══════════════════════════════════════════════════════════════════════════
def build(out_path: str):
    doc = Document()
    for s in doc.sections:
        s.page_width  = Cm(21)
        s.page_height = Cm(29.7)
        s.left_margin = s.right_margin = Cm(2.5)
        s.top_margin  = s.bottom_margin = Cm(2.5)

    # ─── ปก ────────────────────────────────────────────────────────────────
    SP(doc, 1.5)
    cov = doc.add_paragraph()
    cov.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(cov, "รายงาน Model Context Protocol (MCP)\nCheewarun Health Platform",
       bold=True, size=20, color=C_NAVY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(p2, "\nการเชื่อมโยง SmartBreath LSTM Training Data เข้ากับ AI Coach ผ่าน MCP\n"
           "NSC 2026 — เวอร์ชัน 1.0  |  วันที่: 12 กรกฎาคม 2569",
       size=12, color=C_BLUE)
    SP(doc, 0.5)

    abs_t = doc.add_table(rows=1, cols=1)
    abs_t.style = "Table Grid"
    ac = abs_t.cell(0, 0)
    set_bg(ac, C_LBLUE)
    ac.width = Cm(16)
    ap = ac.paragraphs[0]
    ap.paragraph_format.space_before = Pt(8)
    ap.paragraph_format.space_after  = Pt(8)
    rn(ap, "บทคัดย่อ\n", bold=True, size=11, color=C_NAVY)
    rn(ap,
       "รายงานนี้บันทึกสถาปัตยกรรม Model Context Protocol (MCP) ของ Cheewarun "
       "ซึ่งเป็นชั้น middleware ที่เชื่อม Claude AI Coach กับข้อมูลสุขภาพจริงของผู้ใช้ "
       "ครอบคลุม Schema of Context, Data Contract, Tool/Interface, Reasoning Flow "
       "และ Guardrail/Prompt Template ทั้งหมด "
       "พร้อมอธิบายบทบาทของ SmartBreath LSTM Training Data (58 samples, 16 sensors, 300 timesteps) "
       "ในฐานะ pretraining dataset สำหรับโมเดล LSTM temporal ของระบบ",
       size=10.5)

    doc.add_page_break()

    # ─── 1. ภาพรวม MCP ──────────────────────────────────────────────────────
    H1(doc, "1. ภาพรวม Model Context Protocol (MCP) คืออะไร")
    BD(doc,
       "Model Context Protocol (MCP) เป็น open protocol ที่ช่วยให้ Large Language Model (LLM) "
       "เช่น Claude สามารถเชื่อมต่อกับข้อมูลและเครื่องมือภายนอกได้อย่างมีโครงสร้าง "
       "แทนที่จะให้ผู้ใช้ copy-paste ข้อมูลเข้าไปเอง MCP จะทำให้ AI สามารถ:")
    for item in [
        "อ่านข้อมูล sensor readings จาก MetaBreath โดยตรง",
        "เรียก API endpoint ของ Cheewarun backend เพื่อดึงแนวโน้มและการพยากรณ์",
        "บันทึกข้อมูลมื้ออาหาร กิจกรรม และ calibration",
        "อ้างอิง reference range ของ TGS1820 และ Anderson 2015 เป็น Resource",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.8)
        rn(p, item, size=10.5)

    BD(doc,
       "ใน Cheewarun, MCP server ชื่อ cheewarun-mcp ทำงานเป็น stdio process "
       "รับคำสั่งจาก Claude Desktop และส่งผ่านมายัง FastAPI backend (port 8000) "
       "ผ่าน HTTP/JSON ที่มี Bearer Token authentication")

    H2(doc, "1.1 ตำแหน่งของ MCP ในสถาปัตยกรรมระบบ")
    info_box(doc,
        "Flow: ESP32 → MQTT → FastAPI → MCP Server → Claude AI Coach → ผู้ใช้",
        "MetaBreath (ESP32/TGS1820)  →  MQTT Broker  →  FastAPI Backend\n"
        "                                                        ↓\n"
        "                                              MCP Server (Python stdio)\n"
        "                                                        ↓\n"
        "                               Claude AI (ผ่าน Claude Desktop หรือ /ai/chat)\n"
        "                                                        ↓\n"
        "                              Cheewarun Web App (แสดงผลผู้ใช้)\n"
        "\nSmartBreath LSTM Training Data → LSTM Model (lstm_model.pt) → /ai/predict/lstm endpoint → MCP Tool")

    doc.add_page_break()

    # ─── 2. SmartBreath Dataset ──────────────────────────────────────────────
    H1(doc, "2. SmartBreath LSTM Training Data — บทบาทในระบบ")
    BD(doc,
       "SmartBreath LSTM Training Data เป็นชุดข้อมูลสัญญาณเซนเซอร์แก๊สจริง "
       "จากงานวิจัย Ziyatdinov et al. (2015) ที่บันทึกการตอบสนองของ array เซนเซอร์ 16 ตัว "
       "ต่อแก๊สอะซีโตน เอทานอล และส่วนผสม ภายใต้วงจรการไหล (flow modulation) "
       "จำลองการหายใจ 5 ครั้ง/นาที")

    H2(doc, "2.1 ลักษณะข้อมูล")
    tbl(doc,
        ["คุณสมบัติ", "ค่า", "หมายเหตุ"],
        [
            ["ที่มา", "UCI ML Repository / Ziyatdinov et al. 2015", "DOI: 10.24432/C5BG7G"],
            ["License", "CC BY 4.0", "ใช้ได้เชิงวิจัย ต้องอ้างอิง"],
            ["จำนวน Sample", "58 ตัวอย่าง", "แยกตาม gas class และ concentration"],
            ["จำนวน Sensor", "16 ตัว (MOS/SnO₂)", "5 รุ่น TGS: TGS2600/2602/2610/2620 + 1"],
            ["Timesteps/Sample", "300 จุด (1 จุด/วินาที)", "Downsample จาก 7,500 จุด (25 Hz × 5 min)"],
            ["จำนวน Row รวม", "17,400 แถว", "58 × 300 = 17,400"],
            ["Feature ต่อ Row", "S1_dR … S16_dR", "Normalized ΔR/R₀ ของแต่ละเซนเซอร์"],
        ],
        widths=[4.0, 5.5, 6.5])

    H2(doc, "2.2 การกระจาย Gas Type")
    tbl(doc,
        ["Gas Type", "จำนวน Sample", "Concentration (vol.%)", "บทบาทใน LSTM"],
        [
            ["air",     "8",  "0 (clean air baseline)", "สอน LSTM รู้จัก Baseline ก่อนมีแก๊ส"],
            ["acetone", "15", "0.1 / 0.3 / 1.0",        "สัญญาณหลัก — ตรงกับ metabolic acetone"],
            ["ethanol", "15", "0.1 / 0.3 / 1.0",        "Cross-sensitivity test (TGS1820 sensitive ~30%)"],
            ["mixture", "20", "acetone + ethanol mixed",  "Real-world breath complexity simulation"],
            [("รวม", C_NAVY), "58", "—", ""],
        ],
        widths=[2.5, 3.0, 4.5, 6.0])

    H2(doc, "2.3 โครงสร้าง 3 ไฟล์หลัก")
    tbl(doc,
        ["ไฟล์", "แถว", "รูปแบบ", "ใช้สำหรับ"],
        [
            ["LSTM_Training_Data.csv",  "17,400", "Wide: sample × time × S1_dR…S16_dR",      "เทรน LSTM — sequence (300×16) ต่อ sample"],
            ["Precomputed_Features.csv","58",      "Steady-state max response ต่อ sensor",     "เทรน XGBoost/RF fallback (non-sequence)"],
            ["Sample_Index.csv",        "58",      "Metadata: batch, gas, ace_conc, eth_conc", "แบ่ง Train/Test โดยไม่ให้ leak ข้าม sample"],
        ],
        widths=[4.5, 1.5, 5.5, 4.5])

    H2(doc, "2.4 การนำข้อมูลเข้าสู่ Pipeline LSTM ของ Cheewarun")
    BD(doc,
       "LSTM ของ Cheewarun ใช้ SmartBreath เป็น pretraining dataset "
       "เพื่อให้โมเดลเรียนรู้ pattern การตอบสนองของเซนเซอร์ MOS/SnO₂ ต่ออะซีโตน "
       "ก่อนที่จะ fine-tune ด้วยข้อมูลจาก MetaBreath TGS1820 จริง")
    tbl(doc,
        ["ขั้นตอน", "การดำเนินการ", "Output"],
        [
            ["1. Group by sample", "จัดกลุ่ม LSTM_Training_Data ตาม sample ID", "58 sequence แต่ละชุด (300 × 16)"],
            ["2. Sliding window", "ตัด subsequence ขนาด 60-step (overlap 30)", "~928 training windows สำหรับ augmentation"],
            ["3. StandardScaler", "Normalize ต่อ sensor channel", "ค่า mean=0, std=1 ต่อ channel"],
            ["4. Label mapping", "ace_conc → metabolic label ตาม threshold", "0 ppm=air, 0.1=fat_burning, 0.3+=ketosis"],
            ["5. Train LSTM", "2-layer LSTM (64→32) + Dropout + FC", "lstm_model.pt"],
            ["6. Deploy", "โหลดผ่าน ml_inference.py predict_risk_lstm()", "API /ai/predict/lstm"],
        ],
        widths=[3.0, 6.5, 6.5])

    warn_box(doc,
        "ข้อจำกัดสำคัญ: SmartBreath ≠ ลมหายใจมนุษย์จริง",
        "SmartBreath ใช้เซนเซอร์ตระกูล TGS (TGS2600/2602/2610/2620) ไม่ใช่ TGS1820 ของ MetaBreath "
        "และเป็นการสัมผัสแก๊สในห้องปฏิบัติการ ไม่ใช่ลมหายใจจากมนุษย์จริง "
        "ดังนั้น LSTM ที่เทรนจาก SmartBreath จึงเป็น surrogate pretraining เท่านั้น "
        "ต้องการ fine-tuning ด้วยข้อมูลจาก MetaBreath จริงก่อนนำไปใช้งานทางคลินิก")

    doc.add_page_break()

    # ─── 3. Schema of Context ──────────────────────────────────────────────
    H1(doc, "3. Schema of Context — บริบทที่ส่งให้ AI")
    BD(doc,
       "ก่อน Claude จะตอบคำถามสุขภาพของผู้ใช้ MCP จะรวบรวม Context Package "
       "ซึ่งประกอบด้วยข้อมูล 4 ชั้น:")

    H2(doc, "3.1 Context Package Structure")
    tbl(doc,
        ["ชั้น", "ชื่อ Context", "ข้อมูลที่รวม", "แหล่งที่มา"],
        [
            ["C1", "User Profile Context",
             "goal_type, age, activity_level, health_conditions",
             "FastAPI /profile"],
            ["C2", "Sensor Reading Context",
             "acetone_delta (ppm), quality_score, reliability_score,\nlabel, confidence_score, model_used",
             "MCP Tool: get_recent_readings → /sensor/readings"],
            ["C3", "Metabolic Trend Context",
             "trend_direction, slope_ppm_per_day, predicted_points[7],\nconfidence, n_readings_used",
             "MCP Tool: get_metabolic_trend → /ai/trend"],
            ["C4", "Reference Context",
             "acetone-ranges (Anderson 2015), TGS1820 datasheet specs",
             "MCP Resource: cheewarun://reference/*"],
        ],
        widths=[0.8, 3.5, 6.2, 5.5])

    H2(doc, "3.2 Context JSON Schema (SensorReading)")
    CODE(doc, '{')
    CODE(doc, '  "time":              "2026-07-12T08:30:00Z",   // ISO8601 timestamp')
    CODE(doc, '  "device_id":         "uuid-v4",                // MetaBreath device UUID')
    CODE(doc, '  "acetone_delta":     12.4,                     // breath VOC - baseline VOC (ppm)')
    CODE(doc, '  "quality_score":     87.0,                     // 0-100: sensor quality gate')
    CODE(doc, '  "reliability_score": 82.0,                     // 0-100: quality + drift + calibration age')
    CODE(doc, '  "label":             "nutritional_ketosis",     // Anderson 2015 five-class')
    CODE(doc, '  "confidence_score":  0.9430,                   // ML model confidence (0-1)')
    CODE(doc, '  "model_used":        "xgboost",                // xgboost | random_forest | lstm | rule_based')
    CODE(doc, '  "metabolic_risk_index": 2,                     // 0=basal, 1=light, 2=nutritional, 3=deep, 4=dka')
    CODE(doc, '  "pressure_mean":     1.8,                      // kPa breath pressure (XGZP6847A)')
    CODE(doc, '  "temperature":       26.5,                     // °C (SHT31)')
    CODE(doc, '  "humidity":          68.0                      // %RH (SHT31)')
    CODE(doc, '}')
    SP(doc, 0.3)

    H2(doc, "3.3 Context JSON Schema (TrendResponse)")
    CODE(doc, '{')
    CODE(doc, '  "trend_direction":    "increasing",')
    CODE(doc, '  "slope_ppm_per_day":  1.24,')
    CODE(doc, '  "confidence":         0.87,')
    CODE(doc, '  "n_readings_used":    14,')
    CODE(doc, '  "predicted_points": [')
    CODE(doc, '    { "time": "2026-07-13T00:00:00Z", "predicted_acetone": 13.6 },')
    CODE(doc, '    { "time": "2026-07-14T00:00:00Z", "predicted_acetone": 14.9 }')
    CODE(doc, '  ]')
    CODE(doc, '}')

    doc.add_page_break()

    # ─── 4. Data Contract ──────────────────────────────────────────────────
    H1(doc, "4. Data Contract (schema.json)")
    BD(doc,
       "Data Contract กำหนด \"สัญญา\" ระหว่าง MCP Server กับ FastAPI backend "
       "ว่าข้อมูลแต่ละ object มีฟิลด์อะไร ประเภทข้อมูลคืออะไร และค่าที่ valid คืออะไร "
       "ใช้ JSON Schema draft-07 เป็นมาตรฐาน")

    tbl(doc,
        ["Schema Object", "ฟิลด์หลัก", "ข้อกำหนดสำคัญ"],
        [
            ["SensorReading",
             "time, device_id, acetone_delta, quality_score,\nreliability_score, label, confidence_score,\nmetabolic_risk_index",
             "required: [time, device_id]\nlabel enum: healthy|fat_burning|ketosis|diabetes|unreliable\nmetabolic_risk_index enum: 0|1|2|3|null"],
            ["TrendResponse",
             "device_id, trend_direction, slope_ppm_per_day,\npredicted_points[], confidence, n_readings_used",
             "trend_direction enum: increasing|decreasing|stable|insufficient_data\nconfidence: 0.0–1.0"],
            ["CalibrationReport",
             "device_id, lod_ppm, repeatability_cv_pct,\ndrift_slope_ppm_per_day, needs_recalibration",
             "lod_ppm: Limit of Detection\nneeds_recalibration: boolean\ndrift_slope_ppm_per_day: ค่าลบ = drift ลง"],
            ["PilotSession",
             "cohort, day_number, timepoint, homa_ir,\nblood_glucose, blood_ketone_mmol",
             "required: [cohort, day_number, timepoint]\ncohort enum: 5day_20p|14day_10p\ntimepoint enum: fasting|post_meal_60|post_meal_120"],
            ["AcetoneLabelMap",
             "healthy, fat_burning, ketosis, diabetes_dka",
             "Reference mapping: ppm range → metabolic label\nยึดตาม Anderson 2015"],
        ],
        widths=[3.5, 6.0, 6.5])

    H2(doc, "4.1 Label Mapping — Anderson 2015 → Data Contract")
    tbl(doc,
        ["Anderson Label (Backend)", "ช่วง ppm", "Data Contract Label", "metabolic_risk_index"],
        [
            ["basal",                "0.5–2.0",   "healthy",     "0"],
            ["light_ketosis",        "2.0–4.0",   "fat_burning", "1"],
            ["nutritional_ketosis",  "4.0–30.0",  "ketosis",     "2"],
            ["deep_ketosis",         "30.0–75.0", "ketosis",     "2"],
            [("dka_risk", C_RED),    "≥ 75.0",    ("diabetes", C_RED), ("3", C_RED)],
            ["unreliable",           "—",          "unreliable",  "null"],
        ],
        widths=[4.0, 2.5, 3.5, 4.5])
    NOTE(doc, "backend ใช้ 5 class (Anderson), Data Contract ใช้ 4 class (simplified) เพื่อง่ายต่อการ Map ใน MCP")

    doc.add_page_break()

    # ─── 5. Tools / Interface ──────────────────────────────────────────────
    H1(doc, "5. Tool / Interface — เครื่องมือที่ Claude ใช้ได้")
    BD(doc,
       "MCP Server ลงทะเบียน Tool 6 ตัวให้ Claude เรียกใช้ได้โดยอัตโนมัติ "
       "เมื่อ Claude วิเคราะห์ว่าควรดึงข้อมูลหรือดำเนินการใด")

    tools = [
        ("get_recent_readings",
         "ดึงค่า breath acetone sensor 5 readings ล่าสุด",
         "device_id (required), days (default 7)",
         "list[SensorReading summary] — time, acetone_delta, label, quality_score, confidence_score",
         "/sensor/readings",
         "Claude วิเคราะห์ pattern จาก 5 readings ล่าสุด"),
        ("get_metabolic_trend",
         "พยากรณ์ทิศทาง 7 วัน จาก Linear Regression",
         "device_id (required), days (default 7)",
         "TrendResponse — direction, slope, predicted_points[7], confidence",
         "/ai/trend",
         "Claude บอกผู้ใช้ว่า ketosis กำลังเพิ่มหรือลด"),
        ("explain_reading",
         "แปลค่า ppm เป็นคำอธิบายภาษาธรรมชาติ",
         "acetone_ppm (required), context (optional)",
         "state, explanation_th, reference_ranges",
         "Internal (no API call)",
         "Claude อธิบายความหมายของค่าที่ผู้ใช้ถามโดยตรง"),
        ("log_meal",
         "บันทึกมื้ออาหารเพื่อ correlate กับค่า ketone",
         "name (required), kcal, carbs_g",
         "meal_id (UUID)",
         "/logs/meal",
         "Claude บันทึกข้อมูลหลังผู้ใช้บอกว่าเพิ่งกินอะไร"),
        ("log_activity",
         "บันทึกกิจกรรมออกกำลังกาย",
         "kind (required), duration_min (required), kcal",
         "activity_id (UUID)",
         "/logs/activity",
         "Claude บันทึกหลังผู้ใช้บอกว่าเพิ่งวิ่งหรือออกกำลังกาย"),
        ("calibrate_device",
         "Trigger zero-point calibration ของ MetaBreath",
         "device_id (required), baseline_voc (required), temp_c, humidity_pct",
         "CalibrationReport",
         "/sensor/device/{id}/calibrate",
         "Claude แนะนำ calibrate เมื่อ quality_score ต่ำ"),
    ]

    for i, (name, desc, params, ret, api, when) in enumerate(tools):
        H3(doc, f"5.{i+1}  {name}")
        tbl(doc,
            ["รายการ", "ค่า"],
            [
                ["คำอธิบาย", desc],
                ["พารามิเตอร์", params],
                ["ค่าที่ return", ret],
                ["API endpoint", api],
                ["Claude ใช้เมื่อ", when],
            ],
            widths=[3.5, 12.5],
            hbg=C_TEAL)

    doc.add_page_break()

    # ─── 6. Resources ─────────────────────────────────────────────────────
    H1(doc, "6. Resources — ฐานความรู้อ้างอิง")
    BD(doc,
       "Resource ใน MCP คือข้อมูล reference ที่ Claude สามารถ \"อ่าน\" ได้เหมือนเปิดเอกสาร "
       "ต่างจาก Tool ตรงที่ Resource เป็น read-only และไม่ trigger side effect")

    tbl(doc,
        ["URI", "ชื่อ", "เนื้อหา", "Claude ใช้เมื่อ"],
        [
            ["cheewarun://reference/acetone-ranges",
             "MetaBreath Acetone Reference Ranges",
             "4 class ranges: healthy (0.3–0.9), fat_burning (1–5),\nketosis (5–40), diabetes_risk (>75) ppm\nพร้อม measurement formula",
             "ก่อนอธิบายค่า ppm ให้ผู้ใช้ ต้องการ context ว่า range นี้หมายความว่าอะไร"],
            ["cheewarun://reference/tgs1820-datasheet",
             "TGS1820 Sensor Characteristics",
             "cross-sensitivity (Ethanol ~30%, H₂ ~15%),\noperating range (20–60°C, 10–95%RH),\ndrift characteristics (~0.2–0.5 ppm/month),\nLoD typical 0.01 ppm",
             "เมื่อผู้ใช้ถามว่าเซนเซอร์แม่นแค่ไหน หรือทำไมค่าถึงผันผวน"],
        ],
        widths=[5.0, 3.5, 5.0, 4.0])

    doc.add_page_break()

    # ─── 7. Reasoning Flow ────────────────────────────────────────────────
    H1(doc, "7. Reasoning Flow — วิธีที่ Claude ตัดสินใจตอบ")
    BD(doc,
       "เมื่อผู้ใช้ถามคำถามสุขภาพ Claude จะทำตาม Reasoning Flow 5 ขั้นตอน "
       "โดยใช้ MCP Tools เป็นเครื่องมือดึงข้อมูลก่อนสังเคราะห์คำตอบ")

    H2(doc, "7.1 Flow Diagram (ข้อความ)")
    info_box(doc,
        "Reasoning Flow — Cheewarun AI Coach",
        "ผู้ใช้ถาม: \"ค่าของฉันเป็นยังไงบ้างตอนนี้?\"\n"
        "│\n"
        "▼ Step 1: SAFETY CHECK (Guardrail pre-screen)\n"
        "  llm_guardrail.is_refusal_needed(user_message)\n"
        "  → ถ้า match pattern ยา/วินิจฉัย → คืน build_refusal_response() ทันที\n"
        "│\n"
        "▼ Step 2: CONTEXT GATHERING (MCP Tools)\n"
        "  [Tool] get_recent_readings(device_id) → ดึง 5 readings ล่าสุด\n"
        "  [Tool] get_metabolic_trend(device_id) → ดึง trend 7 วัน\n"
        "  [Resource] cheewarun://reference/acetone-ranges → โหลด reference\n"
        "│\n"
        "▼ Step 3: SIGNAL PROCESSING CHECK\n"
        "  ตรวจ quality_score ≥ 60 และ reliability_score ≥ 40\n"
        "  → ถ้าต่ำ: แจ้งผู้ใช้ว่าข้อมูลคุณภาพต่ำ แนะนำวัดใหม่\n"
        "│\n"
        "▼ Step 4: AI CLASSIFICATION (LSTM + XGBoost)\n"
        "  LSTM (≥5 readings) หรือ XGBoost (snapshot)\n"
        "  → ได้ label, confidence_score, metabolic_risk_index\n"
        "│\n"
        "▼ Step 5: SYNTHESIS + GUARDRAIL OUTPUT\n"
        "  สังเคราะห์คำตอบตาม SYSTEM_PROMPT_TEMPLATE\n"
        "  llm_guardrail.sanitise_response() → กรอง + เพิ่ม disclaimer\n"
        "  → ส่งคืนผู้ใช้ผ่าน /ai/chat หรือ Claude Desktop",
        bg=RGBColor(0xEA, 0xF2, 0xFB))

    H2(doc, "7.2 Reasoning Flow รายขั้นตอนแบบละเอียด")
    tbl(doc,
        ["ขั้นตอน", "Action", "Tool/Module ที่ใช้", "Output"],
        [
            ["1. Safety Check",
             "ตรวจ input ผู้ใช้ด้วย regex pattern 30+ รายการ",
             "llm_guardrail.is_refusal_needed()",
             "should_refuse: bool, reason: str"],
            ["2a. Fetch Readings",
             "ดึง 5 readings ล่าสุดจาก backend",
             "MCP Tool: get_recent_readings",
             "list[SensorReading] — acetone, label, quality"],
            ["2b. Fetch Trend",
             "ดึง trend 7 วันและ 7-day forecast",
             "MCP Tool: get_metabolic_trend",
             "TrendResponse — direction, slope, predictions"],
            ["2c. Load Reference",
             "โหลด reference range จาก MCP Resource",
             "MCP Resource: cheewarun://reference/*",
             "JSON: acetone-ranges, TGS1820 specs"],
            ["3. Quality Gate",
             "ตรวจ quality_score + reliability_score",
             "signal_processing.quality_score()",
             "ผ่าน / แจ้งวัดใหม่"],
            ["4. AI Classify",
             "เรียก LSTM (≥5) หรือ XGBoost (snapshot)",
             "ml_inference.predict_risk_lstm() / predict_risk()",
             "label, confidence, metabolic_risk_index"],
            ["5. Build Prompt",
             "รวม user_context + sensor_data เข้า system prompt",
             "llm_guardrail.build_system_prompt()",
             "Formatted system prompt"],
            ["6. LLM Generate",
             "Claude สร้างคำตอบตาม system prompt",
             "Claude API (via /ai/chat)",
             "Raw LLM response"],
            ["7. Sanitise",
             "กรอง banned phrases + เพิ่ม disclaimer",
             "llm_guardrail.sanitise_response()",
             "Safe response + DISCLAIMER_TH"],
        ],
        widths=[2.5, 4.5, 4.5, 4.5])

    doc.add_page_break()

    # ─── 8. Guardrail Template ────────────────────────────────────────────
    H1(doc, "8. Guardrail Template — ระบบความปลอดภัย")
    BD(doc,
       "Guardrail ทำงาน 2 ทิศทาง: กรอง Input จากผู้ใช้ และกรอง Output จาก LLM "
       "เพื่อป้องกันเนื้อหาที่อาจเป็นอันตรายทางการแพทย์")

    H2(doc, "8.1 ตาราง Banned Patterns")
    tbl(doc,
        ["หมวดหมู่", "ตัวอย่าง Pattern (EN)", "ตัวอย่าง Pattern (TH)", "การตอบสนอง"],
        [
            ["ปรับยา/ฉีดยา",
             "adjust insulin dose, how much metformin",
             "ปรับยา, ฉีดอินซูลินเท่าไหร่, เพิ่มยา",
             "build_refusal_response()"],
            ["วินิจฉัยโรค",
             "you have diabetes, diagnosed with DKA",
             "คุณเป็นเบาหวาน, วินิจฉัยโรค",
             "build_refusal_response()"],
            ["ปฏิเสธการพบแพทย์",
             "don't need a doctor",
             "ไม่ต้องไปหาหมอ, ไม่จำเป็นต้องพบแพทย์",
             "build_refusal_response()"],
            ["อดอาหารสุดขีด",
             "fast for 7 days, VLCD",
             "อดอาหาร 7 วัน, ไม่กิน 5 วัน",
             "build_refusal_response()"],
            [("ทำร้ายตัวเอง", C_RED),
             ("suicide, self-harm, kill myself", C_RED),
             ("อยากตาย, ฆ่าตัวตาย", C_RED),
             ("build_refusal_response() + emergency", C_RED)],
        ],
        widths=[3.0, 4.5, 4.5, 4.0])

    H2(doc, "8.2 ตัวอย่าง Refusal Response (ภาษาไทย)")
    info_box(doc, "build_refusal_response(lang='th')",
        "\"ขอโทษค่ะ คำถามนี้เกี่ยวกับการรักษาทางการแพทย์เฉพาะบุคคล\n"
        "ซึ่ง Cheewarun ไม่สามารถให้คำแนะนำได้\n"
        "กรุณาปรึกษาแพทย์หรือผู้เชี่ยวชาญด้านสุขภาพโดยตรง\n\n"
        "---\n"
        "⚠️ ข้อมูลนี้เพื่อการศึกษาเท่านั้น ไม่ใช่คำแนะนำทางการแพทย์\n"
        "ปรึกษาแพทย์หรือผู้เชี่ยวชาญด้านสุขภาพก่อนปรับเปลี่ยนพฤติกรรมหรือการรักษา\"",
        bg=RGBColor(0xF9, 0xEB, 0xEA), title_color=C_RED)

    H2(doc, "8.3 Mandatory Disclaimer (แนบท้ายทุก Response)")
    info_box(doc, "DISCLAIMER_TH (บังคับทุก response)",
        "\"\\n\\n---\\n"
        "⚠️ ข้อมูลนี้เพื่อการศึกษาเท่านั้น ไม่ใช่คำแนะนำทางการแพทย์\n"
        "ปรึกษาแพทย์หรือผู้เชี่ยวชาญด้านสุขภาพก่อนปรับเปลี่ยนพฤติกรรมหรือการรักษา\"")

    doc.add_page_break()

    # ─── 9. Prompt Template ───────────────────────────────────────────────
    H1(doc, "9. Prompt Template — System Prompt สำหรับ Claude")
    BD(doc,
       "SYSTEM_PROMPT_TEMPLATE คือ template ที่ build_system_prompt() ใช้สร้าง system prompt "
       "ก่อนส่งให้ Claude ทุกครั้ง ประกอบด้วย 5 ส่วน:")

    H2(doc, "9.1 โครงสร้าง System Prompt")
    tbl(doc,
        ["ส่วน", "ชื่อ", "เนื้อหา"],
        [
            ["ROLE",     "บทบาท AI",
             "Wellness assistant ด้าน ketogenic, IF, metabolic health\n"
             "อธิบาย sensor readings ในภาษาไทย (หรืออังกฤษถ้าถาม)\n"
             "ให้คำแนะนำ evidence-based nutrition และ exercise"],
            ["STRICT RULES", "กฎห้ามละเมิด 5 ข้อ",
             "1. ห้ามสั่งจ่ายยาหรือปรับขนาดยา\n"
             "2. ห้ามวินิจฉัยโรค\n"
             "3. ห้ามบอกว่าไม่ต้องพบแพทย์\n"
             "4. ต้องแนบ disclaimer ทุก response\n"
             "5. อาการฉุกเฉิน → แนะนำโทร 1669 ทันที"],
            ["USER CONTEXT", "ข้อมูลผู้ใช้",
             "goal_type, recent_readings[], trend_direction, streak_days\n"
             "(inject จาก FastAPI /profile + /sensor/readings)"],
            ["SENSOR DATA", "ข้อมูล sensor ล่าสุด",
             "acetone_delta, quality_score, reliability_score, label,\n"
             "confidence_score, model_used, metabolic_risk_index\n"
             "(inject จาก MCP Tool: get_recent_readings)"],
            ["REASONING FLOW", "ขั้นตอนการตอบ 5 ข้อ",
             "1. Acknowledge ค่าในภาษาธรรมชาติ\n"
             "2. เปรียบเทียบกับ reference range (low/moderate/high ppm)\n"
             "3. แนะนำ 1-2 lifestyle action\n"
             "4. ระบุ data quality concern ถ้าต้อง recalibrate\n"
             "5. แนบ disclaimer"],
        ],
        widths=[2.5, 3.0, 10.5])

    H2(doc, "9.2 Prompt Template (ย่อ)")
    CODE(doc, 'SYSTEM_PROMPT_TEMPLATE = """')
    CODE(doc, 'You are Cheewarun AI Coach — a wellness assistant specialising in')
    CODE(doc, 'ketogenic lifestyle, intermittent fasting, and metabolic health monitoring.')
    CODE(doc, '')
    CODE(doc, 'ROLE:')
    CODE(doc, '- Explain sensor readings in plain Thai (or English if asked)')
    CODE(doc, '- Provide evidence-based nutrition and exercise guidance')
    CODE(doc, '- Encourage healthy habits aligned with the user\'s goal_type')
    CODE(doc, '')
    CODE(doc, 'STRICT RULES (never break these):')
    CODE(doc, '1. Never prescribe or recommend specific medications or dosages')
    CODE(doc, '2. Never diagnose the user with any disease')
    CODE(doc, '3. Never tell a user they do not need to see a doctor')
    CODE(doc, '4. Always end every response with the disclaimer: "{disclaimer}"')
    CODE(doc, '5. Emergency symptoms → immediately say: "โปรดโทร 1669 หรือไปห้องฉุกเฉินทันที"')
    CODE(doc, '')
    CODE(doc, 'USER CONTEXT: {user_context}')
    CODE(doc, 'RECENT SENSOR DATA: {sensor_data}')
    CODE(doc, '')
    CODE(doc, 'REASONING FLOW:')
    CODE(doc, '1. Acknowledge reading in plain language')
    CODE(doc, '2. Compare to reference: low (<30 ppm) / moderate (30–74) / high (≥75)')
    CODE(doc, '3. Identify 1–2 actionable lifestyle suggestions')
    CODE(doc, '4. Note data quality concerns (low quality_score, needs recalibration)')
    CODE(doc, '5. Append disclaimer')
    CODE(doc, '"""')

    H2(doc, "9.3 MCP Prompts (Predefined Templates)")
    tbl(doc,
        ["Prompt Name", "คำอธิบาย", "พารามิเตอร์", "ผลลัพธ์ที่คาดหวัง"],
        [
            ["analyze_metabolic_state",
             "วิเคราะห์ metabolic state ครบถ้วนจาก sensor readings ล่าสุด",
             "device_id (required)",
             "Claude เรียก get_recent_readings อัตโนมัติ แล้วสังเคราะห์การวิเคราะห์ 3–5 ประโยค"],
            ["daily_coaching_message",
             "สร้างข้อความโค้ชประจำวันส่วนตัว",
             "goal_type: keto|if|exercise|diabetes_management\nstreak_days (optional)",
             "ข้อความสร้างแรงบันดาลใจ 2–3 ประโยค เหมาะกับเป้าหมายผู้ใช้"],
        ],
        widths=[4.0, 4.5, 5.0, 6.5])

    doc.add_page_break()

    # ─── 10. สรุปการเชื่อมโยง SmartBreath → MCP ─────────────────────────
    H1(doc, "10. สรุปการเชื่อมโยง SmartBreath → MCP → AI Coach")
    BD(doc,
       "แผนภาพต่อไปนี้สรุปว่าข้อมูล SmartBreath LSTM Training Data "
       "ไหลผ่านระบบอย่างไรจนถึงมือผู้ใช้:")

    info_box(doc,
        "Data Flow: SmartBreath → LSTM → MCP → Claude → ผู้ใช้",
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        " SmartBreath LSTM Training Data (58 samples, 16 sensors, 300 timesteps)\n"
        "        ↓ (notebook 04_lstm_temporal.ipynb)\n"
        " LSTM Model (lstm_model.pt) — 2-layer, ~37K params, PyTorch\n"
        "        ↓ (ml_inference.predict_risk_lstm())\n"
        " FastAPI endpoint: POST /ai/predict/lstm\n"
        "  → รับ: sequence 5 readings × 8 features\n"
        "  → คืน: label, confidence_score, metabolic_risk_index\n"
        "        ↓\n"
        " MCP Tool: get_recent_readings(device_id)\n"
        "  → Claude เรียกใช้ tool นี้อัตโนมัติ\n"
        "  → รับ JSON: [SensorReading × 5] รวม label จาก LSTM\n"
        "        ↓\n"
        " MCP Tool: get_metabolic_trend(device_id)\n"
        "  → รับ: trend_direction, slope_ppm_per_day, predicted_points\n"
        "        ↓\n"
        " build_system_prompt(user_context, sensor_data)\n"
        "  → inject ข้อมูลทั้งหมดเข้า SYSTEM_PROMPT_TEMPLATE\n"
        "        ↓\n"
        " Claude AI Coach สังเคราะห์คำตอบ (ผ่าน Guardrail)\n"
        "        ↓\n"
        " sanitise_response() → กรอง banned + เพิ่ม DISCLAIMER_TH\n"
        "        ↓\n"
        " ผู้ใช้รับคำแนะนำ wellness ที่ปลอดภัย ไม่ใช่ diagnosis",
        bg=RGBColor(0xE8, 0xF8, 0xF5), title_color=C_TEAL)

    H2(doc, "10.1 สรุปตาราง Component Mapping")
    tbl(doc,
        ["Component", "Input", "Process", "Output", "ใช้ใน MCP"],
        [
            ["SmartBreath CSV",
             "16 sensors × 300 timesteps × 58 samples",
             "Sliding window 60-step → LSTM training",
             "lstm_model.pt weights",
             "ทางอ้อม — เป็น pretraining foundation"],
            ["LSTM Model",
             "Sequence 5 readings × 8 features",
             "2-layer LSTM → softmax 3 class",
             "label, confidence",
             "เรียกผ่าน /ai/predict/lstm ที่ MCP Tool เรียก"],
            ["MCP get_recent_readings",
             "device_id, days",
             "HTTP GET /sensor/readings",
             "JSON SensorReading[5]",
             "Tool ใน server.py"],
            ["MCP get_metabolic_trend",
             "device_id, days",
             "HTTP GET /ai/trend → Linear Regression",
             "JSON TrendResponse",
             "Tool ใน server.py"],
            ["MCP explain_reading",
             "acetone_ppm, context",
             "Internal threshold logic",
             "state, explanation_th",
             "Tool ใน server.py (no API)"],
            ["Guardrail",
             "user message + LLM response",
             "Regex scan 30+ banned patterns",
             "safe text + disclaimer",
             "Pre/post Claude call"],
            ["System Prompt Template",
             "user_context, sensor_data",
             "String format injection",
             "Formatted system prompt",
             "llm_guardrail.build_system_prompt()"],
        ],
        widths=[3.5, 3.5, 3.5, 3.5, 2.0])

    doc.add_page_break()

    # ─── 11. References ───────────────────────────────────────────────────
    H1(doc, "11. บรรณานุกรม")
    refs = [
        "[1] Ziyatdinov A, Fonollosa J, Fernández L, Gutiérrez-Gálvez A, Marco S, Perera A. "
        "\"Data set from gas sensor array under flow modulation.\" "
        "Data in Brief (2015) 3:131-136. DOI: 10.1016/j.dib.2015.02.016. UCI ML Repository DOI: 10.24432/C5BG7G",
        "[2] Anthropic. \"Model Context Protocol (MCP) Specification.\" "
        "modelcontextprotocol.io, 2024. Open-source protocol for LLM-tool integration.",
        "[3] Anderson JC. \"Measuring breath acetone for monitoring fat loss.\" "
        "Obesity (2015) 23(12):2327-2334. doi:10.1002/oby.21242",
        "[4] Hochreiter S, Schmidhuber J. \"Long Short-Term Memory.\" "
        "Neural Computation (1997) 9(8):1735-1780.",
        "[5] Figaro Engineering Inc. \"TGS1820 Datasheet.\" Figaro USA, 2020.",
        "[6] Sensirion AG. \"SHT31 Digital Humidity and Temperature Sensor Datasheet.\" 2021.",
        "[7] Varma S, et al. \"Cross-sensitivity of metal oxide gas sensors: a review.\" "
        "Sensors and Actuators B: Chemical (2021).",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        rn(p, ref, size=10, color=RGBColor(0x2C, 0x3E, 0x50))

    SP(doc, 0.5)
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(foot,
       "Cheewarun MCP Report — NSC 2026 | สร้างเมื่อ 12 กรกฎาคม 2569 | "
       "เพื่อวัตถุประสงค์ทางวิชาการ ไม่ใช่สำหรับการวินิจฉัยโรค",
       size=8.5, italic=True, color=C_GREY)

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    out = os.path.abspath(os.path.join(
        os.path.dirname(__file__),
        "../../../Cheewarun_MCP_Technical_Report_NSC2026.docx"
    ))
    build(out)
