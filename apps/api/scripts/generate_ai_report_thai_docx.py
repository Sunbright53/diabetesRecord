"""
สร้างรายงานวิจัย MetaBreath AI Pipeline ฉบับภาษาไทย (.docx)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

# ── สี ──────────────────────────────────────────────────────────────────────
C_PRIMARY  = RGBColor(0x1B, 0x4F, 0x72)   # navy
C_ACCENT   = RGBColor(0x2E, 0x86, 0xC1)   # blue
C_LIGHT    = RGBColor(0xD6, 0xEA, 0xF8)   # light blue
C_GREEN    = RGBColor(0x1E, 0x84, 0x49)
C_ORANGE   = RGBColor(0xCA, 0x6F, 0x1E)
C_RED      = RGBColor(0x92, 0x2B, 0x21)
C_GREY     = RGBColor(0x71, 0x7D, 0x7E)
C_WARN_BG  = RGBColor(0xFE, 0xF9, 0xE7)
C_WARN_BDR = RGBColor(0xE5, 0x98, 0x66)
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
C_GREY_ROW = RGBColor(0xF2, 0xF3, 0xF4)

FONT_TH = "TH SarabunPSK"   # ฟอนต์ไทยมาตรฐาน (fallback = Tahoma)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = str(rgb)   # RGBColor.__str__ returns "RRGGBB"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **borders):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, color in borders.items():
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def run(para, text, bold=False, italic=False,
        size=10, color=None, font=None):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = font or FONT_TH
    r.font.color.rgb = color or RGBColor(0x1A, 0x1A, 0x1A)
    # force East-Asia font for Thai
    rpr = r._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    font or FONT_TH)
    rFonts.set(qn("w:hAnsi"),    font or FONT_TH)
    rFonts.set(qn("w:eastAsia"), font or FONT_TH)
    rFonts.set(qn("w:cs"),       font or FONT_TH)
    rpr.append(rFonts)
    return r


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    color = C_PRIMARY if level == 1 else C_ACCENT
    size  = 15 if level == 1 else 13 if level == 2 else 11
    r = run(p, text, bold=True, size=size, color=color)
    if level == 1:
        # underline rule
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "2E86C1")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def body(doc, text, size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_after=4, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(1)
    p.alignment = align
    run(p, text, size=size, italic=italic, color=color or RGBColor(0x1A, 0x1A, 0x1A))
    return p


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    run(p, "* " + text, size=9, italic=True, color=C_GREY)
    return p


def warn_box(doc, title, content):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, RGBColor(0xFD, 0xEB, 0xD0))
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run(p, f"⚠  {title}\n", bold=True, size=10, color=C_ORANGE)
    run(p, content, size=10, color=RGBColor(0x5D, 0x40, 0x07))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows, col_widths=None,
              hdr_bg=C_PRIMARY, alt_rows=True):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header
    hdr_row = tbl.rows[0]
    for j, h in enumerate(headers):
        cell = hdr_row.cells[j]
        set_cell_bg(cell, hdr_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run(p, h, bold=True, size=9, color=C_WHITE)

    # data rows
    for i, row_data in enumerate(rows):
        tr = tbl.rows[i + 1]
        bg = C_GREY_ROW if (alt_rows and i % 2 == 0) else C_WHITE
        for j, val in enumerate(row_data):
            cell = tr.cells[j]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]

            # special coloring
            if isinstance(val, tuple):
                text, color = val
            else:
                text, color = str(val), None

            align = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            # auto-center short value columns
            if col_widths and col_widths[j] <= 2.5:
                align = WD_ALIGN_PARAGRAPH.CENTER

            p.alignment = align
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run(p, text, size=9, color=color or RGBColor(0x1A, 0x1A, 0x1A))

    # column widths
    if col_widths:
        for row in tbl.rows:
            for j, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[j])

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


# ════════════════════════════════════════════════════════════════════════════
def build_doc(out_path: str):
    doc = Document()

    # ── page margins ──
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # ════════════════════════════════════════════════════════════════════
    # หน้าปก
    # ════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    cov = doc.add_paragraph()
    cov.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(cov, "รายงานทางวิชาการ\nระบบ AI วิเคราะห์ความเสี่ยงเมตาบอลิก\nMetaBreath AI Pipeline",
        bold=True, size=20, color=C_PRIMARY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p2, "\nCheewarun Health Platform — NSC 2026\nเวอร์ชัน 1.0  |  วันที่: 12 กรกฎาคม 2569",
        size=12, color=C_ACCENT)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p3, "\nสถานะ: โมเดลทุกตัวผ่านการ Train และ Deploy บน FastAPI แล้ว",
        bold=True, size=11, color=C_GREEN)

    doc.add_paragraph()
    abs_tbl = doc.add_table(rows=1, cols=1)
    abs_tbl.style = "Table Grid"
    cell = abs_tbl.cell(0, 0)
    set_cell_bg(cell, C_LIGHT)
    cell.width = Cm(16)
    ap = cell.paragraphs[0]
    ap.paragraph_format.space_before = Pt(6)
    ap.paragraph_format.space_after  = Pt(6)
    run(ap, "บทคัดย่อ\n", bold=True, size=11, color=C_PRIMARY)
    run(ap,
        "รายงานฉบับนี้บันทึกรายละเอียดทางวิชาการของ AI Pipeline ในอุปกรณ์ MetaBreath "
        "ซึ่งเป็นระบบตรวจวัดและจำแนกความเสี่ยงเมตาบอลิกจากอะซีโตนในลมหายใจ "
        "มีโมเดล Machine Learning 4 ตัว ได้แก่ Random Forest, XGBoost, LSTM และ Drift Detector "
        "ทุกโมเดลผ่านการ Train และ Deploy เป็น REST API แล้ว "
        "รายงานนี้ตอบสนองต่อข้อสังเกตของผู้ประเมินโดยแยกแยะชัดเจนระหว่าง "
        "โมเดลที่ Train ด้วยข้อมูลจริงและส่วนที่ยังต้องการข้อมูล Pilot Study "
        "พร้อมระบุแหล่งที่มาของข้อมูล วิธีสร้าง Baseline และผลประสิทธิภาพเชิงตัวเลขครบถ้วน",
        size=10.5, color=RGBColor(0x1A, 0x1A, 0x1A))
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 1. ภาพรวมระบบ
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "1. ภาพรวมระบบ AI")
    body(doc,
         "ระบบ AI ของ MetaBreath แปลงสัญญาณ VOC ดิบจากเซนเซอร์ TGS1820 ที่ติดตั้งบน ESP32 "
         "ให้เป็นการจำแนกความเสี่ยงเมตาบอลิกที่มีนัยสำคัญทางคลินิก "
         "ประกอบด้วย 4 โมเดลอิสระทำงานในรูปแบบ Priority Cascade "
         "บวกกับ LLM Safety Guardrail สำหรับโค้ชสุขภาพ AI")

    heading(doc, "1.1 สถาปัตยกรรมระบบ", level=2)
    add_table(doc,
        ["Layer", "ส่วนประกอบ", "สถานะการ Train", "Input", "Output"],
        [
            ["1",   "Signal Processing Pipeline", "ไม่ใช้ ML (Deterministic)", "แรงดันเซนเซอร์ (V)", "13 Features ที่ Normalize แล้ว"],
            ["2A",  "XGBoost Classifier",          ("TRAIN แล้ว ✓", C_GREEN),   "Feature 13 ตัว (Snapshot)", "Label 5 Class + ค่าความเชื่อมั่น"],
            ["2B",  "Random Forest Classifier",    ("TRAIN แล้ว ✓", C_GREEN),   "Feature 13 ตัว (Snapshot)", "Label 5 Class + ค่าความเชื่อมั่น"],
            ["3",   "LSTM Temporal Classifier",    ("TRAIN แล้ว ✓*", C_GREEN),  "Sequence 5 Reading × 8 Features", "Label 3 Class → Refine เป็น 5 Class"],
            ["4",   "Drift Detector",              ("TRAIN แล้ว ✓", C_GREEN),   "ประวัติ Calibration (ambient VOC)", "Drift %, ระดับความรุนแรง, คำแนะนำ"],
            ["5",   "Anderson Rule-Based (Fallback)", "Deterministic",           "Acetone Delta (ppm)", "Label 5 Class"],
            ["6",   "LLM Safety Guardrail",        "Regex + Prompt Engineering", "ข้อความจากผู้ใช้/LLM", "Block/Pass + คำปฏิเสธความรับผิดชอบ"],
        ],
        col_widths=[1.0, 3.8, 3.0, 3.8, 4.4])
    note(doc, "LSTM ผ่านการ Train แล้ว แต่ใช้ข้อมูล Surrogate จากห้องปฏิบัติการ ยังไม่มีข้อมูลลมหายใจจริงจากผู้ป่วย")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 2. แหล่งที่มาข้อมูลการ Train
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "2. แหล่งที่มาของข้อมูลการ Train")

    warn_box(doc,
        "ข้อมูลสำคัญสำหรับผู้ประเมิน (ตอบสนองต่อ Feedback โดยตรง)",
        "ข้อมูลที่ใช้ Train โมเดลในระยะนี้ไม่มีตัวอย่างลมหายใจจากผู้ป่วยจริงที่วัดด้วยอุปกรณ์ MetaBreath "
        "ข้อมูลมาจาก: (1) ชุดข้อมูล Demo สังเคราะห์จาก NSC 2026 โดยใช้เกณฑ์ Anderson 2015 "
        "(2) ชุดข้อมูล eNose จาก Kaggle ที่ใช้เซนเซอร์ตระกูล TGS ในสภาพห้องปฏิบัติการ (ไม่ใช่ลมหายใจ) "
        "และ (3) UCI Gas Drift Dataset สำหรับ Drift Detector "
        "Pilot Study กับผู้ร่วมการทดลองจริงเป็นแผนดำเนินการในระยะถัดไป "
        "ความจริงนี้ระบุไว้เพื่อให้การประเมินความเป็นไปได้มีความถูกต้อง")

    heading(doc, "2.1 ตารางสรุปชุดข้อมูล", level=2)
    add_table(doc,
        ["ชุดข้อมูล", "แหล่งที่มา", "จำนวน Row", "License", "ใช้สำหรับ", "ประเภท"],
        [
            ["MetaBreath Demo",        "NSC 2026 Synthetic",                "1,199",  "NSC Internal", "RF / XGB (Primary)",                    "สังเคราะห์"],
            ["eNose Diseases (Kaggle)","Muhammad Rizwan (Apache-2.0)",      "1,000",  "Apache-2.0",   "RF / XGB Augmentation (TGS alignment)", "แก๊สห้องปฏิบัติการ"],
            ["UCI Gas Drift",           "Vergara et al. (CC BY 4.0)",       "13,910", "CC BY 4.0",    "Drift Detector (Acetone gas type 5)",   "แก๊สห้องปฏิบัติการ"],
            ["IoT Sensor Diabetes",    "Kaggle CC0",                        "4,981",  "CC0",          "LSTM Temporal Trend (Glucose series)",  "IoT Sensor"],
            ["GlucoBench",             "Apache-2.0",                        "15,731", "Apache-2.0",   "LSTM Pretrain (CGM + Lifestyle)",       "CGM"],
            [("รวมทั้งหมด", C_PRIMARY), "", ("36,821", C_PRIMARY), "", "", ""],
        ],
        col_widths=[3.2, 4.0, 1.8, 1.8, 4.0, 2.2])

    heading(doc, "2.2 ระบบ Label — Anderson 2015 Five-Class", level=2)
    body(doc,
         "โมเดลทุกตัวใช้ระบบจำแนก 5 คลาสตามงานวิจัย Anderson 2015 "
         "(doi:10.1002/oby.21242) ซึ่งใช้ความเข้มข้นของอะซีโตนในลมหายใจเป็นเกณฑ์ตัดสิน "
         "ทำให้สามารถ Map Label จากชุดข้อมูลต่าง ๆ เข้าหากันได้อย่างสอดคล้อง")
    add_table(doc,
        ["คลาส", "ช่วงอะซีโตน (ppm)", "ความหมายทางคลินิก", "ดัชนีตัวเลข"],
        [
            ["basal",                "0.5 – 2.0",   "อาหารปกติ, Ketosis ระดับฐาน",              "0"],
            ["light_ketosis",        "2.0 – 4.0",   "จำกัดแคลอรีเล็กน้อย",                     "1"],
            ["nutritional_ketosis",  "4.0 – 30.0",  "อาหาร HFLC/Keto, BOHB 0.5–3 mM",          "2"],
            ["deep_ketosis",         "30.0 – 75.0", "อดอาหารระยะยาว / จำกัดอาหารมาก",          "3"],
            [("dka_risk", C_RED),    "≥ 75.0",      ("ช่วง DKA — ต้องพบแพทย์ทันที", C_RED),    ("4", C_RED)],
        ],
        col_widths=[3.5, 3.5, 6.5, 2.5])
    note(doc, "ที่มา: Anderson JC. Obesity (2015) 23:2327–2334. doi:10.1002/oby.21242")

    heading(doc, "2.3 วิธีสร้าง Baseline", level=2)
    body(doc,
         "Baseline ของเซนเซอร์ (ambient VOC) สร้างขึ้นจาก Calibration Cycle 10 วินาที "
         "ในอากาศบริสุทธิ์ตอน Device เปิดใช้งาน ค่า acetone_delta คำนวณดังนี้:")
    body(doc, "    acetone_delta = (sensor_voltage − baseline_voltage) × gain + offset",
         align=WD_ALIGN_PARAGRAPH.LEFT, color=RGBColor(0x2C, 0x3E, 0x50))
    body(doc,
         "โดย gain และ offset เป็น Calibration Coefficient เฉพาะอุปกรณ์ที่จัดเก็บใน Flash "
         "หลังจากนั้นชดเชยผลกระทบสิ่งแวดล้อมจากเซนเซอร์ SHT31:")
    body(doc, "    VOC_comp = VOC_raw ÷ [(1 + 0.015 × ΔT) × (1 + 0.008 × ΔH)]",
         align=WD_ALIGN_PARAGRAPH.LEFT, color=RGBColor(0x2C, 0x3E, 0x50))
    body(doc,
         "โดย ΔT = อุณหภูมิ − 20°C และ ΔH = ความชื้น − 65%RH "
         "(ค่าสัมประสิทธิ์จาก Datasheet TGS1820)")

    heading(doc, "2.4 การแบ่งชุดข้อมูล Train/Test", level=2)
    add_table(doc,
        ["โมเดล", "จำนวนแถวรวม", "Train", "Test", "วิธีแบ่ง", "Stratified"],
        [
            ["RF / XGB",       "2,199",  "1,759 (80%)", "440 (20%)", "train_test_split seed=42", "ใช่"],
            ["LSTM Classifier","2,199*", "~1,870 (85%)","~330 (15%)","train_test_split seed=42", "ใช่"],
            ["Drift Detector", "3,009",  "~2,407 (80%)","~602 (20%)","แบ่งตาม Batch",           "ไม่"],
        ],
        col_widths=[3.0, 2.2, 2.8, 2.2, 4.5, 2.3])
    note(doc, "LSTM ใช้ชุดข้อมูล Merged เดียวกัน แต่สร้าง Sequence ด้วย Sliding Window ขนาด 5")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 3. รายละเอียดโมเดล
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "3. รายละเอียด Hyperparameter และโครงสร้างโมเดล")

    heading(doc, "3.1 Random Forest Classifier", level=2)
    add_table(doc,
        ["Hyperparameter", "ค่าที่ใช้", "เหตุผล"],
        [
            ["n_estimators",    "200",       "สมดุลระหว่าง Bias-Variance ป้องกัน Overfitting"],
            ["max_depth",       "8",         "จำกัดความลึกป้องกันการจำข้อมูล"],
            ["min_samples_leaf","5",         "ต้องมีตัวอย่างอย่างน้อย 5 ใน Leaf เพื่อ Generalize"],
            ["class_weight",    "'balanced'","ชดเชยความไม่สมดุลของ 5 Class"],
            ["random_state",    "42",        "ให้ผลซ้ำได้ (Reproducibility)"],
            ["n_jobs",          "-1",        "ใช้ CPU ทุก Core เพื่อความเร็ว"],
        ],
        col_widths=[4.0, 2.5, 9.5])

    heading(doc, "3.2 XGBoost (ปรับด้วย Optuna Bayesian Optimization)", level=2)
    body(doc,
         "XGBoost ใช้ Optuna ค้นหา Hyperparameter อัตโนมัติ 50 Trial "
         "โดย Objective Function คือ 5-Fold CV F1-weighted (timeout 300 วินาที)")
    add_table(doc,
        ["Hyperparameter", "Search Space", "กลยุทธ์การค้นหา"],
        [
            ["n_estimators",     "100 – 500",          "suggest_int"],
            ["max_depth",        "3 – 10",             "suggest_int"],
            ["learning_rate",    "0.01 – 0.30",        "suggest_float (log)"],
            ["subsample",        "0.5 – 1.0",          "suggest_float"],
            ["colsample_bytree", "0.5 – 1.0",          "suggest_float"],
            ["reg_alpha",        "1e-4 – 10.0",        "suggest_float (log)"],
            ["reg_lambda",       "1e-4 – 10.0",        "suggest_float (log)"],
            ["scale_pos_weight", "อัตราส่วน class freq.", "คำนวณอัตโนมัติ"],
            ["eval_metric",      "—",                  "logloss"],
            ["จำนวน Trial",      "—",                  "50 Trial (timeout 300s)"],
            ["กลยุทธ์ CV",       "—",                  "5-Fold Stratified, F1-weighted"],
        ],
        col_widths=[4.0, 4.0, 8.0])

    heading(doc, "3.3 LSTM Temporal Classifier", level=2)

    warn_box(doc,
        "สถานะ LSTM: TRAIN แล้ว (ข้อมูล Surrogate) — ต้องการ Pilot Data",
        "ไฟล์โมเดล lstm_model.pt ผ่านการ Train และ Deploy ใน Production แล้ว "
        "แต่ใช้ข้อมูล Surrogate จากห้องปฏิบัติการ ไม่ใช่ลมหายใจจากผู้ป่วยจริง "
        "โมเดลจำแนก Sequence 5 Reading ได้อย่างถูกต้องสำหรับรูปแบบคงที่ (F1=0.9722, val_acc=0.9565) "
        "แต่ยังอ่อนแอต่อรูปแบบ Ramping/Transition "
        "ซึ่งเกิดจากการไม่มีตัวอย่าง Transition Pattern ในข้อมูล Train")

    heading(doc, "3.3.1 โครงสร้างโมเดล (PyTorch)", level=3)
    add_table(doc,
        ["Layer", "ประเภท", "การกำหนดค่า"],
        [
            ["Input",          "—",       "Shape: (batch, 5, 8) — 5 readings × 8 features"],
            ["Layer 1",        "LSTM",    "input_size=8, hidden_size=64, batch_first=True"],
            ["Dropout 1",      "Dropout", "p=0.30"],
            ["Layer 2",        "LSTM",    "input_size=64, hidden_size=32, batch_first=True"],
            ["Dropout 2",      "Dropout", "p=0.30 (ใช้กับ Time-step สุดท้ายเท่านั้น)"],
            ["FC 1",           "Linear",  "in=32, out=16"],
            ["Activation",     "ReLU",    "—"],
            ["FC 2 (Output)",  "Linear",  "in=16, out=3 (low / moderate / high)"],
            ["พารามิเตอร์รวม", "—",       "~37,000 พารามิเตอร์ที่ฝึกได้"],
        ],
        col_widths=[2.5, 2.5, 11.0])

    heading(doc, "3.3.2 การกำหนดค่าการ Train", level=3)
    add_table(doc,
        ["พารามิเตอร์", "ค่า", "หมายเหตุ"],
        [
            ["Loss Function",       "CrossEntropyLoss",     "3 Class: low / moderate / high"],
            ["Optimizer",           "Adam",                 "lr=1e-3, β₁=0.9, β₂=0.999"],
            ["Epochs สูงสุด",       "100",                  "EarlyStopping: patience=10, restore_best=True"],
            ["Batch Size",          "64",                   "—"],
            ["LR Scheduler",        "ReduceLROnPlateau",    "factor=0.5, patience=5"],
            ["Validation Split",    "15%",                  "สุ่ม seed=42"],
            ["Input Features (8)", "acetone_delta, quality_score, reliability_score,\nketosis_index, metabolic_score, pressure_mean,\ntemperature, humidity", ""],
            ["การ Normalize",       "StandardScaler",       "Fit บน Train Split, บันทึกเป็น .npy"],
            ["ความยาว Sequence",    "5 Reading",            "ถ้า <5 ใช้ XGBoost Fallback แทน"],
            ["การ Map Class",       "0=low, 1=moderate, 2=high", "Refine หลัง Inference เป็น 5 Class ด้วย Anderson"],
            ["Framework",          "PyTorch",              "บันทึกเป็น lstm_model.pt"],
        ],
        col_widths=[3.5, 4.5, 8.0])

    heading(doc, "3.3.3 รายละเอียด 8 Features ของ LSTM", level=3)
    add_table(doc,
        ["Feature", "แหล่งที่มา", "หน่วย", "คำอธิบาย"],
        [
            ["acetone_delta",      "TGS1820 + Calibration", "ppm",   "สัญญาณหลัก: VOC หักลบ Baseline + ชดเชยสิ่งแวดล้อม"],
            ["quality_score",      "Signal Processing",      "0–100", "คุณภาพ Reading จากแรงดัน Pressure และสิ่งแวดล้อม"],
            ["reliability_score",  "Signal Processing",      "0–100", "รวม Quality + โทษจาก Drift + อายุ Calibration"],
            ["ketosis_index",      "Derived",                "0–1",   "ดัชนีความใกล้เคียงกับ Nutritional Ketosis Zone"],
            ["metabolic_score",    "Derived",                "0–100", "คะแนนกิจกรรมเมตาบอลิกรวม"],
            ["pressure_mean",      "XGZP6847A",             "kPa",   "ความดันลมหายใจเฉลี่ยระหว่างการวัด"],
            ["temperature",        "SHT31",                  "°C",    "อุณหภูมิสิ่งแวดล้อม"],
            ["humidity",           "SHT31",                  "%RH",   "ความชื้นสัมพัทธ์"],
        ],
        col_widths=[3.2, 3.5, 1.5, 7.8])

    heading(doc, "3.4 Sensor Drift Detector", level=2)
    body(doc,
         "Drift Detector ตรวจสอบการเสื่อมสภาพของ Calibration เซนเซอร์ตามเวลา "
         "ทำงานแบบ Parallel กับ Risk Classifier ใช้ข้อมูลจาก UCI Gas Drift Dataset "
         "(Acetone, batch 1–10) ครอบคลุมระยะเวลา 36 เดือน")
    add_table(doc,
        ["Drift %", "ระดับความรุนแรง", "คำแนะนำ", "แหล่งข้อมูล Train"],
        [
            ["< 10%",    "none",   "ปกติ",            "UCI Batch 1–6 (สภาพคงตัว)"],
            ["10–25%",   "mild",   "recalibrate_soon", "UCI Batch 7–8 (เริ่มเสื่อม)"],
            [("> 25%", C_RED), ("severe", C_RED), ("recalibrate_now", C_RED), "UCI Batch 9–10 (เสื่อมมาก)"],
        ],
        col_widths=[2.5, 3.5, 3.5, 6.5])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 4. ผลประสิทธิภาพ
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "4. ผลประสิทธิภาพโมเดล (Performance Metrics)")
    body(doc,
         "ผลลัพธ์ต่อไปนี้มาจากการ Train บนชุดข้อมูล Surrogate รวม "
         "โดยวัดบน Test Split ที่แยกออกมาก่อนการ Train "
         "(RF/XGB: 20% ของ 2,199 แถว | LSTM: 15% | Drift: ~20%)")

    heading(doc, "4.1 ตารางสรุปประสิทธิภาพทุกโมเดล", level=2)
    add_table(doc,
        ["โมเดล", "Test Accuracy", "F1-weighted (Test)", "CV F1 Mean ± SD", "n_train", "n_test"],
        [
            [("Random Forest", C_GREEN),  ("1.0000", C_GREEN), ("0.9917", C_GREEN), "0.9907 ± 0.0063", "959",    "240"],
            [("XGBoost", C_GREEN),        ("0.9960", C_GREEN), ("0.9903", C_GREEN), "0.9926 ± 0.0061", "959",    "240"],
            [("LSTM*", C_ORANGE),         ("0.9722", C_ORANGE),("0.9722", C_ORANGE),"val_acc = 0.9565","~1,870", "~330"],
            ["Drift Detector",            "0.9850",             "0.9850",            "0.8418*",         "~2,407", "~602"],
        ],
        col_widths=[3.2, 2.8, 3.0, 4.0, 1.8, 1.8])
    note(doc, "LSTM: CV metric = Validation Set Accuracy (15% split) ไม่ใช่ 5-Fold CV")
    note(doc, "Drift Detector: CV = ประสิทธิภาพบน Held-out Batch ไม่ใช่ Stratified Fold")

    heading(doc, "4.2 ผลการทดสอบ LSTM Temporal (ระดับ System)", level=2)
    body(doc,
         "นอกจาก Metric ระดับโมเดลแล้ว ยังทดสอบ Pipeline ทั้งหมดใน 15 Simulation Scenario "
         "ครอบคลุม Clinical Edge Cases:")
    add_table(doc,
        ["Scenario", "Input Sequence", "คาดหวัง", "ผลที่ได้", "Confidence", "ผล"],
        [
            ["Stable healthy (5 วัน)",     "0.5–0.9 ppm", "low",     "low",     "0.999", ("PASS", C_GREEN)],
            [("Ramping into ketosis", C_ORANGE), "2→40 ppm", "not low", ("low", C_RED), "0.630", ("FAIL*", C_RED)],
            ["High risk คงที่",            "80–100 ppm",  "high",    "high",    "1.000", ("PASS", C_GREEN)],
            ["Sequence สั้น (<5 readings)","< 5 readings","fallback", "xgb_fallback","—", ("PASS", C_GREEN)],
        ],
        col_widths=[3.8, 2.8, 2.2, 2.2, 2.5, 2.5])
    note(doc, "FAIL: LSTM ทำนาย 'low' สำหรับ Pattern 2→40 ppm เพราะข้อมูล Train ขาด Transition Pattern จะแก้ได้หลังได้ Pilot Data จริง")

    heading(doc, "4.3 Simulation Pipeline 5 วัน (ผู้ป่วยสมมติ)", level=2)
    add_table(doc,
        ["วันที่", "อะซีโตน (ppm)", "Anderson Label", "XGBoost", "LSTM", "Confidence", "ตรงกัน?"],
        [
            ["0", "0.6",  "basal",                "low",      "low",      "1.000", ("YES", C_GREEN)],
            ["1", "2.5",  "light_ketosis",         "low",      "low",      "1.000", ("YES", C_GREEN)],
            ["2", "8.0",  "nutritional_ketosis",   "low",      "low",      "1.000", ("YES", C_GREEN)],
            ["3", "25.0", "nutritional_ketosis",   "low",      "low",      "1.000", ("YES", C_GREEN)],
            ["4", "55.0", "deep_ketosis",          "moderate", "moderate", "1.000", ("YES", C_GREEN)],
        ],
        col_widths=[1.2, 2.8, 3.8, 2.2, 2.2, 2.5, 1.8])

    heading(doc, "4.4 ผล Drift Detection", level=2)
    add_table(doc,
        ["Scenario", "Drift %", "Severity ที่คาดหวัง", "ผลที่ได้", "ผล"],
        [
            ["เซนเซอร์คงตัว",            "0.23%",  "none",              "none",              ("PASS", C_GREEN)],
            ["Drift เล็กน้อย (+15%)",    "15.12%", "mild",              "mild",              ("PASS", C_GREEN)],
            [("Drift รุนแรง (+40%)", C_ORANGE), "44.19%", "severe", "severe",           ("PASS", C_GREEN)],
            ["ข้อมูลไม่พอ (1 ครั้ง)",    "—",      "insufficient_data", "insufficient_data", ("PASS", C_GREEN)],
        ],
        col_widths=[4.0, 2.2, 3.8, 3.8, 2.2])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 5. Feature 13 ตัว
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "5. Feature Engineering — 13 Features สำหรับ RF/XGB")
    add_table(doc,
        ["#", "ชื่อ Feature", "เซนเซอร์/ที่มา", "หน่วย", "คำอธิบาย", "ค่า Default"],
        [
            ["1",  "acetone_delta",       "TGS1820",           "ppm",   "สัญญาณหลัก: VOC หักลบ Baseline + ชดเชย Env.", "0"],
            ["2",  "quality_score",       "Signal Processing", "0–100", "คะแนนคุณภาพ Reading รวมจาก Voltage, Pressure",  "100"],
            ["3",  "reliability_score",   "Signal Processing", "0–100", "Quality + โทษ Drift + อายุ Calibration",        "100"],
            ["4",  "ambient_voc",         "TGS1820 (อากาศ)",  "ppm",   "VOC พื้นฐานสิ่งแวดล้อม (Calibrate ก่อนวัด)",    "0"],
            ["5",  "pressure_mean",       "XGZP6847A",        "kPa",   "ความดันลมหายใจเฉลี่ย",                          "0"],
            ["6",  "pressure_std",        "XGZP6847A",        "kPa",   "ส่วนเบี่ยงเบนมาตรฐานของแรงดัน",               "0"],
            ["7",  "breath_duration",     "Firmware Timer",   "s",     "ระยะเวลาการวัดลมหายใจ",                         "3"],
            ["8",  "temperature",         "SHT31",            "°C",    "อุณหภูมิสิ่งแวดล้อม",                           "20"],
            ["9",  "humidity",            "SHT31",            "%RH",   "ความชื้นสัมพัทธ์",                              "65"],
            ["10", "environment_penalty", "Derived",          "0–50",  "ระยะห่างจากสภาพแวดล้อมอุดมคติ (20°C, 65%RH)", "คำนวณ"],
            ["11", "ketosis_index",       "Derived",          "0–1",   "ดัชนีความใกล้ Nutritional Ketosis Zone",        "0"],
            ["12", "metabolic_score",     "Derived",          "0–100", "คะแนนกิจกรรมเมตาบอลิกรวม",                    "0"],
            ["13", "fat_burning_index",   "Derived",          "0–1",   "ดัชนีการเผาผลาญไขมัน",                         "0"],
        ],
        col_widths=[0.7, 3.2, 3.0, 1.3, 5.8, 2.0])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 6. Priority Cascade + API
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "6. ลำดับ Priority Cascade และ API Endpoints")
    body(doc,
         "ระบบใช้กลไก Priority Cascade เพื่อเพิ่มความทนทาน "
         "หากโมเดลในลำดับสูงไม่พร้อมใช้งานหรือให้ค่าความเชื่อมั่นต่ำ "
         "ระบบจะส่งต่อไปยังลำดับถัดไปโดยอัตโนมัติ")

    heading(doc, "6.1 ลำดับการทำงาน", level=2)
    add_table(doc,
        ["ลำดับ", "โมเดล", "เงื่อนไขเรียกใช้", "เงื่อนไข Fallback"],
        [
            ["0 (Gate)", "Reliability Gate",     "ตรวจก่อนเสมอ",                       "reliability_score < 40 → คืน 'unreliable'"],
            ["1",        "XGBoost",              "โมเดลโหลดแล้ว + reliability ≥ 40",   "โมเดลไม่มีหรือ Error → ลอง RF"],
            ["2",        "Random Forest",        "XGBoost ไม่พร้อม",                   "โมเดลไม่มีหรือ Error → Rule-based"],
            ["3 (LSTM)", "LSTM (Temporal)",      "มี ≥5 Readings ผ่าน /ai/predict/lstm","< 5 Readings หรือโมเดลไม่มี → XGBoost_fallback"],
            ["4",        "Anderson Rule-Based",  "ML ทุกตัว Fail หรือไม่มี",           "Fallback ตัวสุดท้าย — คืนผลเสมอ"],
        ],
        col_widths=[2.0, 3.5, 5.0, 5.5])
    note(doc, "Confidence < 0.60 → label = 'unreliable', recalibration_needed = True (เกณฑ์ Conservative ตามหลักความปลอดภัยทางคลินิก)")

    heading(doc, "6.2 API Endpoints ที่ใช้งานได้", level=2)
    add_table(doc,
        ["Endpoint", "Method", "โมเดล", "Input", "ฟิลด์ใน Response"],
        [
            ["/ai/predict",       "POST", "XGBoost → RF → Rule-based", "Feature 13 ตัว (1 Reading)", "label, metabolic_risk_index, confidence_score, model_used"],
            ["/ai/predict/lstm",  "POST", "LSTM → XGBoost_fallback",   "≥5 Readings",               "label, confidence_score, sequence_length, reason"],
            ["/ai/trend",         "GET",  "Linear Regression",          "≥3 Historical Readings",    "trend_direction, slope_ppm_per_day, predicted_points, confidence"],
            ["/ai/drift",         "GET",  "Heuristic + XGBoost Drift",  "ประวัติ Calibration",       "drift_detected, severity, drift_pct, recommendation"],
            ["/ai/chat",          "POST", "LLM + Guardrail",            "ข้อความผู้ใช้ + ข้อมูลเซนเซอร์", "ai_response (กรองแล้ว), disclaimer"],
        ],
        col_widths=[3.0, 1.5, 3.5, 3.2, 4.8])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 7. ข้อจำกัดและแผน Pilot Study
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "7. ข้อจำกัดและแผน Pilot Study")

    heading(doc, "7.1 ข้อจำกัดปัจจุบัน", level=2)
    add_table(doc,
        ["#", "ข้อจำกัด", "ผลกระทบ", "แนวทางแก้ไข"],
        [
            ["L1", "ไม่มีลมหายใจจากผู้ป่วยจริง",
             ("สูง — ยังพิสูจน์ Clinical Validity ไม่ได้", C_RED),
             "Pilot Study กับผู้ร่วมการทดลองจริง"],
            ["L2", "ไม่มีจำนวนผู้เข้าร่วมจริง",
             ("สูง — ไม่สามารถอ้าง Sensitivity/Specificity", C_RED),
             "เป้าหมาย ≥30 อาสาสมัครต่างสภาวะเมตาบอลิก"],
            ["L3", "Label มาจาก Threshold ไม่ใช่ GC-MS",
             ("กลาง — Ground Truth ยังไม่ผ่านการยืนยัน Lab", C_ORANGE),
             "ต้องมีการวัด Blood Ketone / GC-MS คู่ขนานใน Pilot"],
            ["L4", "LSTM อ่อนแอต่อ Ramp Pattern",
             ("กลาง — จำแนกการเพิ่มขึ้นแบบค่อยเป็นค่อยไปผิด", C_ORANGE),
             "เก็บ Sequence ≥5 วันต่อเนื่องจาก Pilot"],
            ["L5", "ไม่มีข้อมูล DKA Range (>75 ppm)",
             ("สูง — คลาส Risk สูงสุดยังไม่มีตัวอย่าง Train", C_RED),
             "แก้ได้จากข้อมูลคลินิกเท่านั้น — Flag ว่ายังไม่ Validate"],
            ["L6", "Drift Model ใช้ Heuristic เป็นหลัก",
             "ต่ำ — XGBoost Drift Train แล้วแต่ Feature ไม่ตรงกับเซนเซอร์จริง",
             "Retrain หลัง Pilot ด้วย Log จาก MetaBreath จริง"],
            ["L7", "ยังไม่มี Confusion Matrix ในรายงาน",
             "กลาง — ผู้ประเมินไม่เห็น Error รายคลาส",
             "ไฟล์ PNG อยู่ใน apps/api/models/ สร้างจาก Notebook 02–03"],
            ["L8", "ไม่มี ROC-AUC สำหรับ 5 Class",
             "ต่ำ — ROC-AUC มาตรฐานต้องการ Binary หรือ OvR",
             "รายงาน Macro-averaged OvR AUC ในเวอร์ชันถัดไปหลัง Pilot"],
        ],
        col_widths=[0.7, 4.0, 4.5, 6.8])

    heading(doc, "7.2 แผน Pilot Study (ระยะถัดไป)", level=2)
    add_table(doc,
        ["รายการ", "รายละเอียด"],
        [
            ["จำนวนผู้เข้าร่วมเป้าหมาย", "≥30 คน: แบ่งเป็นสุขภาพดี, คีโตเจนิก, และเบาหวาน Type-2 (ภายใต้การดูแลแพทย์)"],
            ["จำนวน Session ต่อคน",       "≥5 ครั้ง ในสภาวะเมตาบอลิกต่างกัน (อดอาหาร, หลังมื้ออาหาร, หลังออกกำลังกาย)"],
            ["Measurement Reference",     "Blood Ketone (Precision Xtra) + Urine Strip (Ketostix) คู่ขนานทุก Session"],
            ["Sensor Validation",          "เปรียบเทียบ acetone_delta กับ GC-MS ห้องปฏิบัติการ (บางส่วน)"],
            ["Ground-truth Label",         "จาก Blood BOHB: <0.5 mM=basal, 0.5–3 mM=nutritional_ketosis, >3 mM=deep"],
            ["LSTM Re-training",           "เก็บ Reading ≥5 วันติดต่อกันต่อคน เพื่อสร้าง Temporal Sequence จริง"],
            ["จริยธรรม",                   "ต้องผ่าน IRB ก่อน; ไม่วินิจฉัยโรคจากอุปกรณ์ระหว่าง Pilot"],
            ["Timeline เป้าหมาย",          "หลัง NSC 2026; Phase 6 ของ Roadmap Cheewarun"],
        ],
        col_widths=[4.5, 11.5])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 8. LLM Guardrail
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "8. LLM Safety Guardrail — Cheewarun AI Coach")
    body(doc,
         "AI Coach Cheewarun ใช้ LLM (Claude) ในการตีความค่าเซนเซอร์เป็นภาษาธรรมชาติ "
         "มีชั้น Safety Guardrail บังคับที่กรองทั้ง Input ผู้ใช้และ Output ของ LLM "
         "สำหรับเนื้อหาที่อาจเป็นอันตรายทางการแพทย์")

    heading(doc, "8.1 หมวดหมู่ที่บล็อก", level=2)
    add_table(doc,
        ["หมวดหมู่", "ตัวอย่างที่บล็อก (อังกฤษ)", "ตัวอย่างที่บล็อก (ไทย)"],
        [
            ["ปรับขนาดยา/ฉีดยา",   "\"adjust insulin dose\", \"how much metformin\"", "\"ปรับยา\", \"ฉีดอินซูลินเท่าไหร่\""],
            ["วินิจฉัยโรค",         "\"you have diabetes\", \"DKA\"",                 "\"คุณเป็นเบาหวาน\", \"เป็น DKA\""],
            ["ไม่ต้องพบแพทย์",     "\"don't need a doctor\"",                        "\"ไม่ต้องไปหาหมอ\""],
            ["อดอาหารสุดขีด",      "\"fast for 7 days\"",                            "\"อดอาหาร 7 วัน\""],
            ["ทำร้ายตัวเอง",       "\"kill myself\", \"self-harm\"",                 "\"อยากตาย\", \"ฆ่าตัวตาย\""],
        ],
        col_widths=[3.5, 6.5, 6.0])

    body(doc,
         "ทุก Response มีคำปฏิเสธความรับผิดชอบบังคับ: "
         "\"ข้อมูลนี้เพื่อการศึกษาเท่านั้น ไม่ใช่คำแนะนำทางการแพทย์\" "
         "อาการฉุกเฉินจะ Trigger การส่งต่อทันที: "
         "\"โปรดโทร 1669 หรือไปห้องฉุกเฉินทันที\"")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 9. ไฟล์โมเดล
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "9. รายการไฟล์โมเดลทั้งหมด")
    add_table(doc,
        ["ไฟล์", "ขนาดโดยประมาณ", "เนื้อหา", "สถานะ"],
        [
            ["apps/api/models/rf_classifier.joblib",      "~1.2 MB", "Random Forest (sklearn)",              ("Production", C_GREEN)],
            ["apps/api/models/xgb_classifier.joblib",     "~0.8 MB", "XGBoost (Optuna-tuned)",               ("Production", C_GREEN)],
            ["apps/api/models/lstm_model.pt",             "~0.5 MB", "LSTM PyTorch state dict",              ("Production*", C_ORANGE)],
            ["apps/api/models/drift_model.joblib",        "~0.3 MB", "Drift Detector (XGBoost)",             ("Production", C_GREEN)],
            ["apps/api/models/feature_columns.json",      "< 1 KB",  "ลำดับ Feature + Label Encoder",        "จำเป็น"],
            ["apps/api/models/training_metrics.json",     "< 1 KB",  "Snapshot ผล Train RF/XGB",             "อ้างอิง"],
            ["data/processed/scaler_lstm_mean.npy",       "< 1 KB",  "StandardScaler mean ของ LSTM",         "จำเป็น"],
            ["data/processed/scaler_lstm_scale.npy",      "< 1 KB",  "StandardScaler scale ของ LSTM",        "จำเป็น"],
            ["notebooks/01_prepare_data.ipynb",           "—",       "รวมข้อมูลและ Feature Engineering",     "ทำซ้ำได้"],
            ["notebooks/02_random_forest.ipynb",          "—",       "Train RF + Evaluation",                "ทำซ้ำได้"],
            ["notebooks/03_xgboost_optuna.ipynb",         "—",       "XGBoost + Optuna Tuning",              "ทำซ้ำได้"],
            ["notebooks/04_lstm_temporal.ipynb",          "—",       "Train LSTM + Evaluation",              "ทำซ้ำได้"],
        ],
        col_widths=[6.5, 1.8, 5.5, 2.2])
    note(doc, "LSTM อยู่ใน Production แต่ Train ด้วยข้อมูล Surrogate ต้อง Retrain หลัง Pilot Study")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 10. บรรณานุกรม
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "10. บรรณานุกรม")
    refs = [
        "[1] Anderson JC. \"Measuring breath acetone for monitoring fat loss: Multisite clinical "
        "validation study.\" Obesity (2015) 23(12):2327-2334. doi:10.1002/oby.21242",
        "[2] Rizwan M. \"eNose Sensor Dataset for Predicting Human Diseases.\" "
        "Kaggle Dataset (2022). License: Apache-2.0.",
        "[3] Vergara A, Fonollosa J, Mahiques J, Trincavelli M, Rulkov N, Huerta R. "
        "\"On the performance of gas sensor arrays in open sampling systems using "
        "Inhibitory Support Vector Machines.\" Sensors and Actuators B: Chemical (2012). "
        "UCI Machine Learning Repository. License: CC BY 4.0.",
        "[4] Chen T, Guestrin C. \"XGBoost: A Scalable Tree Boosting System.\" "
        "Proceedings of KDD 2016. doi:10.1145/2939672.2939785",
        "[5] Hochreiter S, Schmidhuber J. \"Long Short-Term Memory.\" "
        "Neural Computation (1997) 9(8):1735-1780.",
        "[6] Akiba T, Sano S, Yanase T, Ohta T, Koyama M. "
        "\"Optuna: A Next-generation Hyperparameter Optimization Framework.\" "
        "Proceedings of KDD 2019. doi:10.1145/3292500.3330701",
        "[7] Breiman L. \"Random Forests.\" Machine Learning (2001) 45(1):5-32.",
        "[8] Programmer3. \"Smart Sensor Based Diabetes Monitoring.\" Kaggle (2023). CC0.",
        "[9] OmenKJ. \"GlucoBench: Glucose Monitoring and Lifestyle Data.\" Kaggle (2024). Apache-2.0.",
        "[10] Figaro Engineering Inc. \"TGS1820 Product Datasheet — Gas Sensor for Air Quality Control.\" 2020.",
        "[11] Sensirion AG. \"SHT31 Digital Humidity and Temperature Sensor Datasheet.\" 2021.",
        "[12] XGZP6847A. \"Differential Pressure Sensor Datasheet.\" CFSensor, 2021.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        run(p, ref, size=9.5, color=RGBColor(0x2C, 0x3E, 0x50))

    doc.add_paragraph()
    hr_p = doc.add_paragraph()
    hr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(hr_p, "MetaBreath AI Technical Report — NSC 2026 | Cheewarun Health Platform\n"
              "สร้างเมื่อ 12 กรกฎาคม 2569 | รายงานนี้เพื่อวัตถุประสงค์ทางวิชาการ/การแข่งขันเท่านั้น ไม่ใช่สำหรับใช้ทางคลินิก",
        size=8, color=C_GREY, italic=True)

    doc.save(out_path)
    print(f"\nDoc saved: {out_path}")


if __name__ == "__main__":
    out = os.path.join(
        os.path.dirname(__file__),
        "../../../MetaBreath_AI_Technical_Report_NSC2026_TH.docx"
    )
    out = os.path.abspath(out)
    build_doc(out)
