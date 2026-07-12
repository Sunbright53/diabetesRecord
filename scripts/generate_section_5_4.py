"""Generate Section 5.4 Software Specification as a Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT = "/Users/ciy_th/Desktop/diabetesRecord/5.4_Software_Specification.docx"


def set_thai_font(run, size=14, bold=False, color=None):
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "TH Sarabun New")
    rfonts.set(qn("w:hAnsi"), "TH Sarabun New")
    rfonts.set(qn("w:cs"), "TH Sarabun New")
    rfonts.set(qn("w:eastAsia"), "TH Sarabun New")


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_thai_font(run, size=18, bold=True)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_thai_font(run, size=16, bold=True)


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_thai_font(run, size=15, bold=True)


def add_para(doc, text, indent_first=True):
    p = doc.add_paragraph()
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_thai_font(run, size=14)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.7)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_thai_font(run, size=14)


def add_table(doc, header, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = False
    if col_widths:
        for row in tbl.rows:
            for cell, w in zip(row.cells, col_widths):
                cell.width = w
    hdr = tbl.rows[0].cells
    for i, h in enumerate(header):
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_thai_font(run, size=13, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_thai_font(run, size=13)


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    # Set default style
    style = doc.styles["Normal"]
    style.font.name = "TH Sarabun New"
    style.font.size = Pt(14)

    # ═══════════════════════ HEADER ═══════════════════════
    add_h1(doc, "5.4 รายละเอียดโปรแกรมที่ได้พัฒนาในเชิงเทคนิค (Software Specification)")

    add_para(
        doc,
        "หัวข้อนี้อธิบายรายละเอียดเชิงเทคนิคของระบบ MetaBreath ซึ่งเป็นระบบต้นแบบที่เชื่อมระหว่างฮาร์ดแวร์ตรวจวัดลมหายใจ "
        "(ESP32 + TGS1820 + Pressure Sensor + SHT31) กับซอฟต์แวร์ฝั่งเซิร์ฟเวอร์และเว็บแอปพลิเคชัน โดยครอบคลุมข้อกำหนดของ "
        "input/output, ข้อกำหนดเชิงหน้าที่ (functional specification), โครงสร้างของซอฟต์แวร์ (software design), "
        "รายละเอียดอื่น ๆ ที่เกี่ยวข้อง และแหล่งที่มาของ source code ที่นำมาประกอบ",
        indent_first=True,
    )

    # ═══════════════════════ 5.4.1 I/O ═══════════════════════
    add_h2(doc, "5.4.1 Input / Output Specifications")

    add_h3(doc, "ก. Input ของระบบ")
    add_para(
        doc,
        "ระบบรับข้อมูลเข้าจากหลายแหล่ง โดยแบ่งตามชั้นการทำงานดังนี้",
        indent_first=False,
    )
    add_bullet(doc, "Sensor Input จากอุปกรณ์ ESP32")
    add_bullet(doc, "TGS1820 VOC Sensor — ค่า analog voltage (mV) ผ่าน ADC ของ ESP32 (12-bit, 0–3.3 V)", level=1)
    add_bullet(doc, "Pressure Sensor — ค่า analog voltage แปลงเป็นความดัน (kPa) ในช่วง 0–10 kPa", level=1)
    add_bullet(doc, "SHT31 Temperature/Humidity Sensor — ค่าอุณหภูมิ (°C) และความชื้นสัมพัทธ์ (%RH) ผ่าน I²C", level=1)
    add_bullet(doc, "User Input จากผู้ใช้ผ่าน Web Application")
    add_bullet(doc, "ข้อมูลลงทะเบียน / login (username, password, email, role)", level=1)
    add_bullet(doc, "ข้อมูล onboarding profile (อายุ, เพศ, น้ำหนัก, ส่วนสูง, เป้าหมาย)", level=1)
    add_bullet(doc, "การ pair อุปกรณ์ ESP32 ผ่าน pairing token", level=1)
    add_bullet(doc, "การเริ่ม/หยุด session การวัด (Start / Stop Recording)", level=1)
    add_bullet(doc, "การบันทึกผล log ต่าง ๆ (weight, meal, activity, ketone strip)", level=1)
    add_bullet(doc, "การตั้งค่าการแจ้งเตือน (reminder) และ push notification", level=1)
    add_bullet(doc, "Network Input")
    add_bullet(doc, "MQTT payload (JSON) จาก ESP32 มายัง Mosquitto broker", level=1)
    add_bullet(doc, "REST API request จาก Web Application", level=1)
    add_bullet(doc, "WebSocket connection สำหรับ real-time streaming", level=1)
    add_bullet(doc, "Configuration / Calibration Input")
    add_bullet(doc, "Wi-Fi credentials ผ่าน WiFiManager captive portal", level=1)
    add_bullet(doc, "Firmware configuration parameters (sampling interval, threshold)", level=1)
    add_bullet(doc, "Calibration reference values สำหรับ baseline update", level=1)

    add_h3(doc, "ข. Output ของระบบ")
    add_bullet(doc, "Sensor Output ไปยัง MQTT Broker")
    add_bullet(
        doc,
        "JSON payload บน topic `metabreath/<MAC>/reading` ประกอบด้วย ambient_voc, breath_voc, acetone_delta, "
        "pressure_mean, pressure_std, breath_duration, temperature, humidity, timestamp",
        level=1,
    )
    add_bullet(doc, "API Response (JSON) จาก Backend FastAPI")
    add_bullet(doc, "SensorReadingOut — ข้อมูล reading ที่ผ่านการประมวลผลแล้ว รวม quality_score / reliability_score", level=1)
    add_bullet(doc, "SessionSummary — สรุปแต่ละ session (start, end, mean, peak, label)", level=1)
    add_bullet(doc, "TrendResponse — แนวโน้มย้อนหลัง 7/14/30 วัน", level=1)
    add_bullet(doc, "PredictResponse — ผลการทำนายจาก ML model (label + confidence)", level=1)
    add_bullet(doc, "CalibrationReportOut — รายงาน drift และ recalibration status", level=1)
    add_bullet(doc, "UI Output แสดงบน Web Application")
    add_bullet(doc, "กราฟ acetone แบบเรียลไทม์ (recharts line chart)", level=1)
    add_bullet(doc, "สถานะการเป่า BLOW / IDLE / DONE", level=1)
    add_bullet(doc, "Trend indicator (low / moderate / high) เป็นภาษาผู้ใช้", level=1)
    add_bullet(doc, "Sample quality indicator + reliability score", level=1)
    add_bullet(doc, "หน้าประวัติแนวโน้ม (weekly / monthly view)", level=1)
    add_bullet(doc, "Notification Output")
    add_bullet(doc, "Web Push Notification ผ่าน VAPID + Service Worker", level=1)
    add_bullet(doc, "In-app toast notification (Sonner)", level=1)
    add_bullet(doc, "Device Command Output กลับไปที่ ESP32")
    add_bullet(doc, "MQTT topic `metabreath/<MAC>/command` (เช่น reset_wifi, start_recording)", level=1)

    add_h3(doc, "ค. ตัวอย่าง MQTT Payload Schema")
    add_para(
        doc,
        "ข้อมูลจาก ESP32 ถูกส่งเป็น JSON payload มายัง MQTT broker ตามรูปแบบมาตรฐาน "
        "ประกอบด้วยค่าจากเซนเซอร์ (breath_voc, ambient_voc, acetone_delta, pressure_mean, pressure_std), "
        "ค่าสิ่งแวดล้อม (temperature, humidity), metadata (device_mac, timestamp, reading_number) และ "
        "สถานะปัจจุบัน (state: BLOW / IDLE) ซึ่ง MQTT subscriber ฝั่ง Python จะรับ payload นี้และ insert "
        "ลง TimescaleDB hypertable โดยอัตโนมัติ",
        indent_first=True,
    )

    # ═══════════════════════ 5.4.2 Functional Spec ═══════════════════════
    add_h2(doc, "5.4.2 Functional Specification")
    add_para(
        doc,
        "ระบบ MetaBreath ประกอบด้วยฟังก์ชันหลัก 8 กลุ่ม แต่ละกลุ่มมีขอบเขตความรับผิดชอบและ API endpoint ที่ชัดเจน ดังนี้",
        indent_first=True,
    )

    add_h3(doc, "F1. Authentication & Authorization")
    add_bullet(doc, "ลงทะเบียนผู้ใช้ใหม่ (POST /auth/register)")
    add_bullet(doc, "เข้าสู่ระบบด้วย JWT (POST /auth/login)")
    add_bullet(doc, "รีเฟรช token (POST /auth/refresh)")
    add_bullet(doc, "ดึงข้อมูลผู้ใช้ปัจจุบัน (GET /auth/me)")
    add_bullet(doc, "รองรับ role-based access control: user, admin, doctor, caregiver")

    add_h3(doc, "F2. Device Management")
    add_bullet(doc, "สร้าง pairing token (POST /provision/token)")
    add_bullet(doc, "จับคู่อุปกรณ์กับผู้ใช้ (POST /device/pair)")
    add_bullet(doc, "ดูรายการอุปกรณ์ (GET /devices)")
    add_bullet(doc, "อุปกรณ์แบบ shared pool — claim / release (POST /device/{id}/claim, /release)")
    add_bullet(doc, "รีเซ็ต Wi-Fi ระยะไกล (POST /device/{id}/reset-wifi)")
    add_bullet(doc, "ควบคุมการ recording (POST /device/{id}/recording/start | stop | status)")

    add_h3(doc, "F3. Sensor Reading & Session")
    add_bullet(doc, "บันทึก reading (POST /readings)")
    add_bullet(doc, "ดึงข้อมูล reading ย้อนหลัง (GET /readings)")
    add_bullet(doc, "สรุปแต่ละ session (GET /sessions)")
    add_bullet(doc, "สถิติรายวัน (GET /daily-stats)")
    add_bullet(doc, "รับข้อมูลแบบเรียลไทม์ผ่าน WebSocket (WS /ws)")

    add_h3(doc, "F4. Calibration & Signal Processing")
    add_bullet(doc, "บันทึกค่า calibration (POST /device/{id}/calibrate)")
    add_bullet(doc, "ดูประวัติ calibration (GET /device/{id}/calibration)")
    add_bullet(doc, "รายงาน drift / recalibration status (GET /device/{id}/calibration/report)")
    add_bullet(doc, "Signal processing pipeline: baseline correction, quality scoring, drift detection")

    add_h3(doc, "F5. AI / ML Inference")
    add_bullet(doc, "ทำนาย risk จาก snapshot (POST /ai/predict) — Random Forest / XGBoost")
    add_bullet(doc, "ทำนาย risk จาก sequence (POST /ai/predict/lstm) — PyTorch LSTM")
    add_bullet(doc, "จัดหมวดหมู่แนวโน้ม (POST /ai/predict/trend) — LSTM Trend Classifier")
    add_bullet(doc, "แนวโน้มย้อนหลัง (GET /ai/trend)")
    add_bullet(doc, "ตรวจสอบ drift ของเซนเซอร์ (GET /ai/drift)")
    add_bullet(doc, "AI chat สำหรับอธิบายผลลัพธ์ (POST /ai/chat) — Anthropic Claude + MCP")
    add_bullet(doc, "Metabolic flexibility scoring (POST /ai/flexibility)")

    add_h3(doc, "F6. Logs & History")
    add_bullet(doc, "บันทึก / ดูประวัติ ketone log (GET/POST /logs/ketone)")
    add_bullet(doc, "บันทึก / ดูประวัติ weight log (GET/POST /logs/weight)")
    add_bullet(doc, "บันทึก / ดูประวัติ meal log (GET/POST /logs/meal)")
    add_bullet(doc, "บันทึก / ดูประวัติ activity log (GET/POST /logs/activity)")

    add_h3(doc, "F7. Gamification & Content")
    add_bullet(doc, "ดูค่า XP และ level (GET /gamification/xp)")
    add_bullet(doc, "ดู streak (GET /gamification/streak)")
    add_bullet(doc, "ดู badges (GET /gamification/badges)")
    add_bullet(doc, "quest ประจำวัน (GET /gamification/quests/today)")
    add_bullet(doc, "อ่านบทความ / ทำภารกิจการเรียนรู้ (GET /content, POST /content/{slug}/complete)")

    add_h3(doc, "F8. Notification & Reminder")
    add_bullet(doc, "ดึง VAPID public key (GET /push/vapid-public)")
    add_bullet(doc, "subscribe push notification (POST /push/subscribe)")
    add_bullet(doc, "จัดการ reminder — CRUD (GET/POST/PATCH/DELETE /reminders)")

    # ═══════════════════════ 5.4.3 Software Design ═══════════════════════
    add_h2(doc, "5.4.3 โครงสร้างของ Software (Software Design)")

    add_h3(doc, "ก. สถาปัตยกรรมโดยรวม (Overall Architecture)")
    add_para(
        doc,
        "ระบบ MetaBreath ออกแบบด้วยหลัก Layered Microservices Architecture ประกอบด้วย 5 ชั้นหลัก "
        "แต่ละชั้นทำงานอิสระต่อกันและสื่อสารผ่าน well-defined protocol (MQTT, REST, WebSocket) "
        "รายละเอียดของแต่ละชั้นมีดังนี้",
        indent_first=True,
    )

    add_bullet(doc, "Layer 1 — Firmware Layer (ESP32 + Arduino C++)")
    add_bullet(doc, "ทำหน้าที่อ่านค่าเซนเซอร์, calibrate baseline, สร้าง JSON payload และส่งผ่าน MQTT", level=1)
    add_bullet(doc, "Layer 2 — Data Pipeline Layer (Mosquitto + Python MQTT Subscriber)")
    add_bullet(doc, "MQTT broker รับ payload จาก ESP32 หลายเครื่องพร้อมกัน (pub/sub pattern)", level=1)
    add_bullet(doc, "Python subscriber ทำหน้าที่ validate + insert ลง TimescaleDB", level=1)
    add_bullet(doc, "Layer 3 — Backend / Analytics Layer (FastAPI + SQLModel + Celery)")
    add_bullet(doc, "REST API + WebSocket endpoints", level=1)
    add_bullet(doc, "Background job สำหรับ session summary, notification, model inference", level=1)
    add_bullet(doc, "Signal processing pipeline สำหรับ quality scoring และ drift detection", level=1)
    add_bullet(doc, "Layer 4 — AI Layer (scikit-learn + XGBoost + PyTorch + Anthropic Claude)")
    add_bullet(doc, "Rule-based classifier (Anderson 5-pattern scale)", level=1)
    add_bullet(doc, "Random Forest / XGBoost — snapshot classifier (13 features)", level=1)
    add_bullet(doc, "LSTM Trend Classifier — sequence-based (PyTorch)", level=1)
    add_bullet(doc, "Drift detection model", level=1)
    add_bullet(doc, "Anthropic Claude ผ่าน MCP server สำหรับสร้างคำอธิบายผลลัพธ์", level=1)
    add_bullet(doc, "Layer 5 — Application Layer (Next.js + React + TypeScript + Tailwind)")
    add_bullet(doc, "Web application เป็น dashboard สำหรับผู้ใช้ทุก role", level=1)
    add_bullet(doc, "รับข้อมูลผ่าน REST API + WebSocket + MQTT WebSocket bridge", level=1)

    add_h3(doc, "ข. โครงสร้าง Directory ของโปรเจกต์ (Monorepo)")
    add_para(doc, "โปรเจกต์จัดเป็น monorepo ประกอบด้วยโฟลเดอร์หลักดังนี้", indent_first=False)

    add_table(
        doc,
        header=["Path", "หน้าที่"],
        rows=[
            ["apps/firmware/metabreath/", "Arduino C++ firmware สำหรับ ESP32"],
            ["apps/api/", "FastAPI backend, SQLModel, Alembic migration, Celery worker"],
            ["apps/api/app/routers/", "REST API endpoints (auth, sensor, ai, admin, ...)"],
            ["apps/api/app/services/", "Business logic — signal processing, ML inference, calibration"],
            ["apps/api/app/models/", "SQLModel database models"],
            ["apps/api/app/mqtt/", "MQTT publisher + subscriber logic"],
            ["apps/api/app/workers/", "Celery async jobs + MQTT worker"],
            ["apps/api/notebooks/", "Jupyter Notebook สำหรับ train ML/LSTM models"],
            ["apps/api/models/", "โมเดล ML/LSTM ที่ train แล้ว (.joblib, .pt)"],
            ["apps/web/src/app/", "Next.js App Router — pages ทั้งหมด"],
            ["apps/web/src/components/", "Reusable React components"],
            ["apps/web/src/hooks/", "Custom React hooks"],
            ["apps/web/src/lib/", "API client, utilities, MQTT client"],
            ["apps/mcp/src/", "MCP server สำหรับเชื่อม Claude AI"],
            ["packages/shared-types/", "TypeScript types ที่ frontend + backend ใช้ร่วม"],
            ["infra/mosquitto/", "MQTT broker configuration"],
            ["infra/nginx/", "Nginx reverse proxy configuration"],
            ["docker-compose.yml", "Orchestration ของทุก service"],
        ],
        col_widths=[Cm(6.0), Cm(10.0)],
    )

    add_h3(doc, "ค. Database Schema (Time-series + Relational)")
    add_para(
        doc,
        "ระบบใช้ PostgreSQL 16 พร้อม TimescaleDB extension เพื่อรองรับข้อมูล time-series ปริมาณมาก "
        "ประกอบด้วยตารางหลักดังนี้",
        indent_first=True,
    )
    add_table(
        doc,
        header=["ตาราง", "หน้าที่"],
        rows=[
            ["users", "ข้อมูลผู้ใช้ + role + profile"],
            ["devices", "อุปกรณ์ ESP32 ที่ pair กับ user"],
            ["sensor_readings (hypertable)", "ข้อมูลจากเซนเซอร์แบบ time-series"],
            ["device_calibration", "ประวัติ calibration + baseline"],
            ["device_sessions", "ช่วงเวลาที่ user claim อุปกรณ์ shared pool"],
            ["ketone_logs / weight_logs / meal_logs / activity_logs", "log ต่าง ๆ ของผู้ใช้"],
            ["xp / streak / badges / quests", "ระบบ gamification"],
            ["articles / article_progress", "content สำหรับการเรียนรู้"],
            ["reminders / push_subscriptions", "การแจ้งเตือน"],
            ["audit_log", "บันทึกการเปลี่ยนแปลงสำคัญ"],
        ],
        col_widths=[Cm(6.5), Cm(9.5)],
    )

    add_h3(doc, "ง. Deployment Architecture")
    add_para(
        doc,
        "ระบบทั้งหมด deploy ผ่าน Docker Compose ประกอบด้วย service ต่อไปนี้: "
        "(1) web — Next.js frontend, (2) api — FastAPI backend, (3) worker — Celery async worker, "
        "(4) beat — Celery scheduler, (5) mqtt-sub — MQTT subscriber worker, "
        "(6) db — TimescaleDB, (7) redis — cache + message broker, (8) mqtt — Mosquitto broker, "
        "(9) nginx — reverse proxy + TLS termination ที่ edge network",
        indent_first=True,
    )

    # ═══════════════════════ 5.4.4 อื่น ๆ ═══════════════════════
    add_h2(doc, "5.4.4 รายละเอียดอื่น ๆ (Additional Specifications)")

    add_h3(doc, "ก. Security & Authentication")
    add_bullet(doc, "JWT-based authentication (access + refresh token)")
    add_bullet(doc, "Password hashing ด้วย bcrypt (passlib)")
    add_bullet(doc, "Role-based access control (user, admin, doctor, caregiver)")
    add_bullet(doc, "MQTT authentication ด้วย username/password + ACL")
    add_bullet(doc, "CORS whitelist ที่ระดับ FastAPI middleware")
    add_bullet(doc, "TLS termination ที่ Nginx reverse proxy")
    add_bullet(doc, "Environment variable สำหรับ secret (.env, ไม่ commit)")
    add_bullet(doc, "LLM guardrail — ป้องกันการให้ข้อมูลที่ไม่เหมาะสมจาก AI chat")

    add_h3(doc, "ข. Performance Requirements")
    add_bullet(doc, "Sensor sampling — 1 Hz per reading")
    add_bullet(doc, "MQTT latency — < 500 ms (ESP32 → subscriber)")
    add_bullet(doc, "API response time — < 200 ms สำหรับ endpoint หลัก")
    add_bullet(doc, "TimescaleDB hypertable — รองรับ ≥ 1M readings ต่ออุปกรณ์")
    add_bullet(doc, "ML inference latency — < 100 ms ต่อ prediction")
    add_bullet(doc, "WebSocket real-time — < 100 ms end-to-end")

    add_h3(doc, "ค. Testing Strategy")
    add_bullet(doc, "Unit test — pytest สำหรับ signal processing, trend labeling, ML inference")
    add_bullet(doc, "Integration test — pytest + httpx client กับ real database")
    add_bullet(doc, "Frontend type checking — TypeScript strict mode")
    add_bullet(doc, "AI PDF Report — script สำหรับ export ผลการทดสอบ ML/LSTM เป็น PDF")
    add_bullet(doc, "Simulation scenarios — script `simulate_scenarios.py` สำหรับทดสอบ end-to-end")

    add_h3(doc, "ง. Migration & Schema Management")
    add_bullet(doc, "ใช้ Alembic สำหรับ schema versioning")
    add_bullet(doc, "TimescaleDB hypertable ถูกสร้างผ่าน migration script")
    add_bullet(doc, "รองรับ rollback และ zero-downtime migration")

    add_h3(doc, "จ. Observability & Logging")
    add_bullet(doc, "Structured logging ผ่าน Python logging module")
    add_bullet(doc, "Healthcheck endpoint — GET /healthz")
    add_bullet(doc, "Docker healthcheck สำหรับทุก service")
    add_bullet(doc, "Heartbeat file สำหรับ MQTT subscriber")

    # ═══════════════════════ 5.4.5 Self-developed + Sources ═══════════════════════
    add_h2(doc, "5.4.5 ส่วนที่ทีมพัฒนาขึ้นเอง และแหล่งที่มาของ Source Code")

    add_h3(doc, "ก. ส่วนที่ทีมพัฒนาขึ้นเอง 100%")
    add_para(
        doc,
        "ต่อไปนี้คือส่วนสำคัญของระบบที่ทีม MetaBreath ออกแบบและพัฒนาขึ้นเองทั้งหมด "
        "โดยเป็น intellectual property ของทีมและไม่ได้ลอกเลียนแบบจากแหล่งอื่น",
        indent_first=True,
    )

    add_table(
        doc,
        header=["ส่วนที่พัฒนา", "ไฟล์/ตำแหน่ง", "รายละเอียด"],
        rows=[
            [
                "MetaBreath Firmware v2",
                "apps/firmware/metabreath/metabreath.ino",
                "Arduino C++ firmware ทั้งหมด — sensor reading, baseline calibration, "
                "state machine BLOW/IDLE, MQTT publish, WiFiManager integration, remote reset",
            ],
            [
                "Signal Processing Pipeline",
                "apps/api/app/services/signal_processing.py",
                "TGS1820 temperature/humidity/pressure compensation, quality score, "
                "reliability score, breath duration estimation, drift indicator",
            ],
            [
                "ML Inference Layer",
                "apps/api/app/services/ml_inference.py",
                "โหลดโมเดล RF/XGB/LSTM/Drift มาให้บริการทำนาย พร้อม rule-based fallback "
                "ตาม Anderson 5-pattern scale",
            ],
            [
                "LSTM Trend Classifier",
                "apps/api/app/services/trend_label.py + notebooks/train_lstm_trend.py",
                "โมเดล LSTM สำหรับจัดหมวดหมู่แนวโน้ม (stable / increasing / decreasing / abnormal) "
                "จาก session sequence",
            ],
            [
                "Metabolic Flexibility Engine",
                "apps/api/app/services/flexibility_engine.py",
                "คำนวณ flexibility score จาก amplitude + return speed + appropriateness",
            ],
            [
                "Device Session (Shared Pool)",
                "apps/api/app/services/device_session.py",
                "ระบบ claim/release อุปกรณ์ shared pool + attribution ของ reading",
            ],
            [
                "MCP Server",
                "apps/mcp/src/server.py",
                "MCP server สำหรับให้ Claude AI เข้าถึงข้อมูลผู้ใช้ผ่าน tool interface",
            ],
            [
                "LLM Guardrail",
                "apps/api/app/services/llm_guardrail.py",
                "ตัวกรองคำตอบจาก LLM ไม่ให้ตอบคำถามนอกขอบเขต",
            ],
            [
                "Gamification Engine",
                "apps/api/app/services/gamification.py",
                "ระบบ XP, level, streak, badges, quest ประจำวัน",
            ],
            [
                "Web Dashboard UI",
                "apps/web/src/app/**, apps/web/src/components/**",
                "หน้า UI ทั้งหมด — home, breathing, trends, learn, chat, log, me, admin",
            ],
            [
                "MQTT Subscriber Worker",
                "apps/api/app/workers/mqtt_subscriber.py",
                "รับ MQTT payload จาก ESP32 มา validate + insert ลง TimescaleDB",
            ],
            [
                "Database Schema + Migrations",
                "apps/api/alembic/versions/**",
                "Schema design และ Alembic migration script ทั้งหมด",
            ],
            [
                "Simulation & Evaluation Scripts",
                "apps/api/notebooks/simulate_scenarios.py, generate_longitudinal_data.py",
                "สคริปต์สำหรับ simulate scenario และประเมินผลระบบ",
            ],
        ],
        col_widths=[Cm(4.5), Cm(5.0), Cm(6.5)],
    )

    add_h3(doc, "ข. แหล่งที่มาของ Open-source Library ที่นำมาประกอบ")
    add_para(
        doc,
        "ระบบ MetaBreath ใช้ open-source library หลายรายการเป็นส่วนประกอบ โดยทุกตัวมี license "
        "ที่อนุญาตให้ใช้งานเชิงพาณิชย์และวิจัยได้ ตารางต่อไปนี้แสดงแหล่งที่มา, version, "
        "และหน้าที่ในระบบ",
        indent_first=True,
    )

    add_para(doc, "1) Firmware — Arduino C++ Libraries", indent_first=False)
    add_table(
        doc,
        header=["Library", "Version", "License", "หน้าที่ในระบบ"],
        rows=[
            ["WiFiManager (tzapu)", "≥ 2.0.17", "MIT", "captive portal สำหรับตั้งค่า Wi-Fi"],
            ["PubSubClient (Nick O'Leary)", "latest", "MIT", "MQTT client"],
            ["ArduinoJson (Benoit Blanchon)", "latest", "MIT", "สร้าง/parse JSON payload"],
            ["Wire, WiFi, WiFiMulti, Preferences", "ESP32 Core", "LGPL", "core libraries ของ ESP32"],
            ["SHT31 driver", "Adafruit", "BSD", "อ่าน temp/humidity sensor"],
        ],
        col_widths=[Cm(4.5), Cm(2.5), Cm(2.0), Cm(7.0)],
    )

    add_para(doc, "2) Backend — Python Libraries", indent_first=False)
    add_table(
        doc,
        header=["Library", "Version", "License", "หน้าที่ในระบบ"],
        rows=[
            ["fastapi", "0.115.5", "MIT", "REST API framework"],
            ["uvicorn", "0.32.1", "BSD", "ASGI server"],
            ["sqlmodel", "0.0.22", "MIT", "ORM (SQLAlchemy + Pydantic)"],
            ["alembic", "1.14.0", "MIT", "database migration"],
            ["asyncpg / psycopg2", "0.30.0 / 2.9.10", "Apache 2.0 / LGPL", "Postgres driver"],
            ["celery + redis", "5.4.0 / 5.2.1", "BSD / BSD", "async worker + message broker"],
            ["aiomqtt / paho-mqtt", "2.3.0 / 2.x", "BSD / EPL", "MQTT client"],
            ["python-jose", "3.3.0", "MIT", "JWT encoding/decoding"],
            ["passlib + bcrypt", "1.7.4 / 4.0.1", "BSD", "password hashing"],
            ["pywebpush", "2.0.0", "MPL 2.0", "Web Push notification"],
            ["scikit-learn", "1.5.2", "BSD", "Random Forest classifier"],
            ["xgboost", "2.1.3", "Apache 2.0", "Gradient boosting classifier"],
            ["optuna", "4.1.0", "MIT", "Hyperparameter tuning"],
            ["torch (PyTorch)", "latest", "BSD", "LSTM model training + inference"],
            ["joblib", "1.4.2", "BSD", "โหลด/บันทึกโมเดล ML"],
            ["httpx", "0.28.1", "BSD", "async HTTP client"],
            ["anthropic", "0.49.0", "MIT", "Claude API SDK"],
            ["mcp[server]", "1.1.2", "MIT", "Model Context Protocol SDK"],
        ],
        col_widths=[Cm(4.0), Cm(2.5), Cm(2.5), Cm(7.0)],
    )

    add_para(doc, "3) Frontend — JavaScript / TypeScript Libraries", indent_first=False)
    add_table(
        doc,
        header=["Library", "Version", "License", "หน้าที่ในระบบ"],
        rows=[
            ["next", "16.2.10", "MIT", "React framework (App Router)"],
            ["react + react-dom", "19.2.4", "MIT", "UI library"],
            ["typescript", "5.x", "Apache 2.0", "static type checking"],
            ["tailwindcss + @tailwindcss/postcss", "4.x", "MIT", "utility-first CSS"],
            ["@radix-ui/react-*", "1.x / 2.x", "MIT", "accessible UI primitives"],
            ["@tanstack/react-query", "5.101.2", "MIT", "server state management"],
            ["react-hook-form", "7.80.0", "MIT", "form management"],
            ["zod + @hookform/resolvers", "4.4.3 / 5.4.0", "MIT", "schema validation"],
            ["recharts", "3.9.2", "MIT", "chart library"],
            ["framer-motion", "12.42.2", "MIT", "animation library"],
            ["sonner", "2.0.7", "MIT", "toast notification"],
            ["lucide-react", "1.23.0", "ISC", "icon library"],
            ["next-themes", "0.4.6", "MIT", "light/dark mode"],
            ["class-variance-authority + clsx + tailwind-merge", "0.7.1 / 2.1.1 / 3.6.0", "MIT", "className utilities"],
            ["vaul", "1.1.2", "MIT", "drawer/modal component"],
        ],
        col_widths=[Cm(5.0), Cm(2.5), Cm(2.0), Cm(6.5)],
    )

    add_para(doc, "4) Infrastructure", indent_first=False)
    add_table(
        doc,
        header=["Component", "Version", "License", "หน้าที่ในระบบ"],
        rows=[
            ["TimescaleDB (PostgreSQL 16)", "latest-pg16", "Apache 2.0 / TSL", "time-series database"],
            ["Redis", "7-alpine", "BSD", "cache + Celery broker"],
            ["Eclipse Mosquitto", "2.x", "EPL / EDL", "MQTT broker"],
            ["Nginx", "latest", "BSD", "reverse proxy + TLS"],
            ["Docker + Docker Compose", "latest", "Apache 2.0", "containerization"],
        ],
        col_widths=[Cm(5.0), Cm(2.5), Cm(2.5), Cm(6.0)],
    )

    add_h3(doc, "ค. อ้างอิงทางวิชาการ (Academic References)")
    add_bullet(
        doc,
        "Anderson JC. Measuring breath acetone for monitoring fat loss: Review. Obesity (2015) "
        "23(12):2327-2334. doi:10.1002/oby.21242 — ใช้เป็น reference สำหรับ 5-pattern classification "
        "ของ breath acetone",
    )
    add_bullet(doc, "TGS1820 sensor datasheet (Figaro Engineering) — ใช้เป็น reference สำหรับ temperature/humidity compensation")
    add_bullet(doc, "Model Context Protocol Specification (Anthropic, 2024) — สำหรับการ integrate Claude AI")

    add_h3(doc, "ง. สรุป")
    add_para(
        doc,
        "ระบบ MetaBreath ทั้งหมดถูกออกแบบ สถาปัตยกรรม, โค้ด business logic, ML pipeline, "
        "signal processing pipeline, firmware และ UI ทั้งหมดพัฒนาขึ้นโดยทีม โดยใช้ open-source library "
        "ที่มี license เหมาะสมมาประกอบเป็นเครื่องมือระดับ infrastructure และ framework "
        "(FastAPI, Next.js, PyTorch, ฯลฯ) เท่านั้น ส่วนตรรกะทางธุรกิจและตรรกะการวิเคราะห์ข้อมูลทั้งหมด "
        "เป็นผลงานต้นฉบับของทีม MetaBreath",
        indent_first=True,
    )

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
