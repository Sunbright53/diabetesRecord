"""Generate 'Which AI powers Metabolic Flexibility?' Word document.

Positions the score honestly: it is a rule-based heuristic (expert system) —
NOT a machine-learning model — and explains why that choice is intentional,
what the alternatives would be, and how to describe it in academic writing.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT = "/Users/ciy_th/Desktop/diabetesRecord/Metabolic_Flexibility_AI_Method.docx"


def set_thai_font(run, size=14, bold=False, color=None, mono=False):
    name = "Consolas" if mono else "TH Sarabun New"
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:eastAsia"), name)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_thai_font(run, size=20, bold=True)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_thai_font(run, size=16, bold=True)


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_thai_font(run, size=14, bold=True)


def add_para(doc, parts, indent_first=False):
    p = doc.add_paragraph()
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    if isinstance(parts, str):
        run = p.add_run(parts)
        set_thai_font(run, size=14)
        return
    for text, opts in parts:
        run = p.add_run(text)
        set_thai_font(
            run,
            size=14,
            bold=opts.get("bold", False),
            mono=opts.get("mono", False),
        )


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.7)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_thai_font(r1, size=14, bold=True)
        r2 = p.add_run(text)
        set_thai_font(r2, size=14)
    else:
        run = p.add_run(text)
        set_thai_font(run, size=14)


def add_callout(doc, title, body):
    """Highlighted answer box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EAF6EF")
    pPr.append(shd)
    r1 = p.add_run(title + "\n")
    set_thai_font(r1, size=14, bold=True, color=RGBColor(0x0F, 0x66, 0x3E))
    r2 = p.add_run(body)
    set_thai_font(r2, size=14)


def add_code(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)
    run = p.add_run(code_text)
    set_thai_font(run, size=11, mono=True)


def add_table(doc, header, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = False
    if col_widths:
        for row in tbl.rows:
            for cell, w in zip(row.cells, col_widths):
                cell.width = w
    hdr = tbl.rows[0].cells
    for i, h in enumerate(header):
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_thai_font(run, size=13, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_thai_font(run, size=13)


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "TH Sarabun New"
    style.font.size = Pt(14)

    # ═══════════════════════ TITLE ═══════════════════════
    add_h1(doc, "Metabolic Flexibility Score — เอกสารชี้แจงวิธีการและ AI ที่ใช้")

    add_para(
        doc,
        "เอกสารนี้ตอบคำถามที่ลูกค้า/คณะกรรมการมักถามว่า “คะแนน Metabolic Flexibility "
        "ที่ MetaBreath แสดง ใช้ AI ตัวไหน / โมเดลอะไร / คำนวณอย่างไร” "
        "พร้อมเหตุผลเชิงวิศวกรรมและวิทยาการทางการแพทย์ที่รองรับการเลือกวิธีการนี้",
        indent_first=True,
    )

    # ═══════════════════════ 1. ANSWER ═══════════════════════
    add_h2(doc, "1. คำตอบสั้น ๆ (สำหรับใช้พูดหรือเขียนในเอกสาร)")

    add_callout(
        doc,
        "Metabolic Flexibility Score ไม่ได้ใช้โมเดล Machine Learning ตัวใดตัวหนึ่ง",
        "ระบบใช้ Rule-based Expert Scoring Engine (Explainable AI) — คำนวณจาก 3 มิติ "
        "ทางสรีรวิทยาที่มี weight ชัดเจนและ threshold อ้างอิงจากงานวิจัย breath acetone / "
        "metabolic flexibility ระดับ peer-reviewed (Anderson 2015, Kelley 2000, "
        "Goodpaster & Sparks 2017) ทำงานแบบ deterministic — คะแนนเท่าเดิมข้อมูลเท่าเดิม "
        "ตรวจสอบย้อนหลังได้ทุกขั้นตอน ไม่มี black-box",
    )

    add_para(
        doc,
        [
            ("ในเชิงประเภทของ AI: ", {"bold": True}),
            ("จัดเป็น ", {}),
            ("Rule-based Expert System / Knowledge-based System", {"bold": True}),
            (" ซึ่งเป็น subfield คลาสสิกของ Artificial Intelligence (เก่ากว่าและตรงข้ามกับ Statistical/ML AI) — "
             "รูปแบบเดียวกับที่ใช้ในระบบวินิจฉัยการแพทย์รุ่นเก่า เช่น MYCIN (1976), Internist-I (1980s), "
             "หรือ decision support ทางคลินิกในปัจจุบันหลายตัว", {}),
        ],
    )

    # ═══════════════════════ 2. Design rationale ═══════════════════════
    add_h2(doc, "2. ทำไมเลือก Rule-based แทนที่จะใช้ Machine Learning")

    add_para(
        doc,
        "การตัดสินใจนี้ไม่ใช่เรื่องข้อจำกัดทางเทคนิค แต่เป็นการเลือก approach ที่เหมาะสมที่สุด "
        "ในบริบทของระบบสุขภาพต้นแบบ ดังนี้",
        indent_first=True,
    )

    add_bullet(
        doc,
        "อธิบายผลได้ทุกจุด: ผู้ใช้ (และแพทย์) ควรเข้าใจได้ว่าคะแนน 52 มาจาก Amplitude 22 + "
        "Return 17.5 + Appropriateness 12.5 และแต่ละส่วนคำนวณอย่างไร — ML แบบ black-box "
        "(RF, XGB, NN) ไม่สามารถให้ความชัดเจนแบบนี้",
        bold_prefix="Explainability (XAI) — ",
    )
    add_bullet(
        doc,
        "Threshold ทั้งหมด (2, 8, 40, 75 ppm) มาจาก Anderson 2015 5-pattern breath acetone "
        "classification — เป็นตัวเลขที่ผ่านการตรวจสอบทางคลินิก ไม่ใช่ตัวเลขที่ระบบเรียนรู้จากข้อมูลตัวเอง",
        bold_prefix="Clinical grounding — ",
    )
    add_bullet(
        doc,
        "ระบบ MetaBreath อยู่ในช่วง pilot; ข้อมูล longitudinal ที่มี label ที่แม่นยำยังน้อยเกินไป "
        "ที่จะ train ML supervised model สำหรับ metabolic flexibility ให้ generalise ได้ — "
        "การใช้ rule-based ทำให้ระบบพร้อมใช้งานตั้งแต่ Day 1 โดยไม่ต้องรอ dataset",
        bold_prefix="Zero cold-start problem — ",
    )
    add_bullet(
        doc,
        "ทีมพัฒนาปรับ weight/threshold เองได้ทันทีหากผู้ใช้ / แพทย์เห็นควร ไม่ต้อง retrain "
        "หรือ deploy model ใหม่",
        bold_prefix="Tunability — ",
    )
    add_bullet(
        doc,
        "ระบบไม่เรียนรู้จากข้อมูลผู้ใช้ → ไม่ต้องกังวลเรื่อง data leakage หรือ bias ในการ train — "
        "GDPR/PDPA-friendly by design",
        bold_prefix="Regulatory clarity — ",
    )
    add_bullet(
        doc,
        "ในระบบเดียวกัน MetaBreath ก็มีโมเดล ML อยู่จริงในส่วนอื่น (LSTM Trend Classifier "
        "สำหรับ long-term trend, Random Forest + XGBoost สำหรับ per-reading classification) "
        "— rule-based สำหรับ Flexibility จึงเป็นการเลือก tool ที่ถูกกับ task ไม่ได้จำกัดที่ทีมทำ ML ไม่ได้",
        bold_prefix="Not \"can't\", but \"should not\" — ",
    )

    # ═══════════════════════ 3. Method ═══════════════════════
    add_h2(doc, "3. Methodology — วิธีคำนวณจริง ๆ")

    add_para(
        doc,
        [
            ("Source: ", {"bold": True}),
            ("apps/api/app/services/flexibility_engine.py", {"mono": True}),
            (" (228 บรรทัด, MIT-licensed, in-house code)", {}),
        ],
    )

    add_h3(doc, "3.1 Input")
    add_para(doc, "ระบบดึงข้อมูล breath session ของผู้ใช้ในช่วง 14 วันย้อนหลัง (default) โดย:")
    add_bullet(doc, "Group readings ที่ห่างกันไม่เกิน 5 นาที เป็น 1 session")
    add_bullet(doc, "แต่ละ session สรุปเป็น peak / mean acetone + context tag (ถ้ามี)")
    add_bullet(doc, "ต้องมีอย่างน้อย 3 sessions ถึงจะคำนวณคะแนน (มิเช่นนั้นจะ return “insufficient data”)")

    add_h3(doc, "3.2 คะแนน 3 มิติ (รวมได้ 0–100)")

    add_table(
        doc,
        header=["มิติ", "เต็ม", "วัดอะไร", "หลักการ"],
        rows=[
            [
                "Amplitude",
                "40 pts",
                "ค่ากระจายในหลาย metabolic zone หรือติดโซนเดียว",
                "คนที่ยืดหยุ่นควรผ่านหลายโซน (fed_resting → transitional → fat_oxidation)",
            ],
            [
                "Return Speed",
                "35 pts",
                "ค่ากลับ baseline หลังพีค / สลับสูง-ต่ำได้ไหม",
                "Fed–fast transition kinetics; ratio post_meal/fasting ควรน้อย",
            ],
            [
                "Appropriateness",
                "25 pts",
                "ค่าตรงกับ context ที่ user รายงานไหม",
                "Congruence check: post_meal ควรต่ำ, fasting ควรสูง",
            ],
        ],
        col_widths=[Cm(3.0), Cm(1.5), Cm(5.0), Cm(6.5)],
    )

    add_h3(doc, "3.3 สูตร (pseudo-code)")
    add_code(
        doc,
        "def compute_flexibility(sessions, latest_ppm, context_tag):\n"
        "    if len(sessions) < 3:\n"
        "        return {\"score\": 0, \"message\": \"insufficient_data\"}\n"
        "\n"
        "    amp   = _amplitude_score(sessions)        # 0-40\n"
        "    spd   = _return_speed_score(sessions)     # 0-35\n"
        "    appr  = _appropriateness_score(sessions)  # 0-25\n"
        "    total = min(100, amp + spd + appr)\n"
        "\n"
        "    return {\n"
        "        \"score\": total,\n"
        "        \"breakdown\": {\"amplitude\": amp, \"return_speed\": spd, \"appropriateness\": appr},\n"
        "        \"trend\": _trend_direction(sessions),   # increasing/decreasing/stable\n"
        "        \"zone\": _metabolic_zone(latest_ppm),\n"
        "        \"message_th\": _message_th(total),\n"
        "    }",
    )

    add_h3(doc, "3.4 5 Metabolic Zones ที่ใช้จัดค่า (อ้างอิงงานวิจัย)")
    add_table(
        doc,
        header=["โซน", "ช่วง", "ความหมายทางสรีรวิทยา"],
        rows=[
            ["fed_resting",   "0 – 2 ppm",   "Postprandial · glucose oxidation หลัก"],
            ["transitional",  "2 – 8 ppm",   "Mixed substrate oxidation"],
            ["fat_oxidation", "8 – 40 ppm",  "Nutritional ketosis (0.5–3.0 mmol/L blood)"],
            ["extended_fast", "40 – 75 ppm", "Prolonged fasting / therapeutic ketosis"],
            ["safety_alert",  "≥ 75 ppm",    "อาจเข้าใกล้ DKA (ในผู้ป่วยเบาหวาน) — flag ไว้"],
        ],
        col_widths=[Cm(3.5), Cm(3.0), Cm(9.5)],
    )

    # ═══════════════════════ 4. Compared with ML ═══════════════════════
    add_h2(doc, "4. เปรียบเทียบกับ approach ML")

    add_table(
        doc,
        header=["ประเด็น", "Rule-based (ที่ MetaBreath ใช้)", "Supervised ML (RF, XGB, NN)"],
        rows=[
            ["Dataset ต้องมี", "ไม่ต้อง — เขียน rule จากงานวิจัยได้เลย", "ต้องมี ≥ 1,000–10,000 labeled samples"],
            ["Explainability", "สูงมาก — เห็นทุก weight/threshold", "ต่ำ (black-box) หรือปานกลาง (RF + SHAP)"],
            ["Cold-start", "ใช้ได้ตั้งแต่ user คนแรก", "แย่มาก — ต้องมี historical data"],
            ["Robust ต่อ outlier", "สูง (rule จำกัด range)", "ปานกลาง — outlier ทำให้ prediction แปลก"],
            ["Update / tune", "แก้ตัวเลขในโค้ด commit เดียว", "ต้อง re-train + validate + redeploy"],
            ["Bias จาก training data", "ไม่มี (ไม่ได้ train)", "มี — สะท้อน bias ของ dataset"],
            ["Personalisation", "ต่ำ (rule เดียวสำหรับทุกคน)", "สูง — เรียนรู้ per-user ได้"],
            ["Regulatory / audit", "ง่าย — logic อ่านได้", "ยาก — ต้อง explain black-box"],
        ],
        col_widths=[Cm(4.0), Cm(6.0), Cm(6.0)],
    )

    add_para(
        doc,
        "ในบริบทของ MetaBreath ที่ยังอยู่ระยะ pilot + ต้องอธิบายกับผู้ใช้/แพทย์ + threshold "
        "ทางคลินิกมีมาตรฐานอยู่แล้ว → rule-based เป็นตัวเลือกที่เหมาะสมและซื่อสัตย์กับข้อจำกัดของ dataset",
        indent_first=True,
    )

    # ═══════════════════════ 5. AI stack in MetaBreath ═══════════════════════
    add_h2(doc, "5. ทั้งระบบ MetaBreath ใช้ AI ตัวไหนบ้าง (context)")

    add_para(
        doc,
        "เพื่อให้ภาพชัดเจน — Metabolic Flexibility เป็นเพียง 1 ใน 5 องค์ประกอบ AI/analytics "
        "ของระบบ ตัวอื่น ๆ เป็น ML จริง:",
        indent_first=True,
    )

    add_table(
        doc,
        header=["Layer", "ใช้ AI แบบไหน", "ทำหน้าที่อะไร"],
        rows=[
            [
                "Signal processing",
                "Rule-based (compensation formula)",
                "Compensate temp/humidity/pressure effect on TGS1820",
            ],
            [
                "Per-reading risk classifier",
                "Random Forest + XGBoost (scikit-learn)",
                "จัดค่าเดี่ยว 1 reading → 5-band Anderson (basal → dka_risk)",
            ],
            [
                "Long-term trend",
                "LSTM (PyTorch, time-series RNN)",
                "จัดแนวโน้ม 14 sessions → stable / increasing / decreasing / abnormal",
            ],
            [
                "Metabolic Flexibility",
                "Rule-based Expert System (in-house)",
                "คะแนน 0–100 จาก 3 มิติ physiology",
            ],
            [
                "Health coach / explanation",
                "Anthropic Claude LLM via MCP tool interface",
                "สร้างคำอธิบายเป็นภาษาผู้ใช้ + guardrail กรองคำแนะนำที่ไม่ปลอดภัย",
            ],
        ],
        col_widths=[Cm(4.0), Cm(5.5), Cm(6.5)],
    )

    add_para(
        doc,
        "ดังนั้น MetaBreath ไม่ใช่ระบบที่ “ไม่ใช้ AI” — ระบบใช้ AI ครบทั้ง 3 paradigm หลัก: "
        "rule-based (Flexibility), classical ML (RF/XGB), และ deep learning (LSTM) + LLM (Claude) "
        "แต่ละส่วนถูกเลือกให้เหมาะกับ task นั้น ๆ",
        indent_first=True,
    )

    # ═══════════════════════ 6. How to describe ═══════════════════════
    add_h2(doc, "6. วิธีบรรยายในเอกสาร/report แบบซื่อสัตย์และ professional")

    add_h3(doc, "6.1 ตัวอย่างประโยคภาษาไทย (สำหรับใส่ใน proposal / report)")

    add_para(
        doc,
        [
            ("แบบสั้น: ", {"bold": True}),
            ("“Metabolic Flexibility Score พัฒนาเป็น rule-based expert scoring engine ที่คำนวณจาก "
             "3 มิติทางสรีรวิทยา (Amplitude, Return Speed, Appropriateness) ตาม threshold ที่อ้างอิงจาก "
             "งานวิจัย breath acetone และ metabolic flexibility (Anderson 2015; Kelley 2000; "
             "Goodpaster & Sparks 2017) โดยเลือกใช้วิธีนี้เพื่อความโปร่งใส (Explainable AI) "
             "และเพื่อไม่ต้องพึ่ง large training dataset ในระยะ pilot”", {}),
        ],
    )

    add_para(
        doc,
        [
            ("แบบยาว: ", {"bold": True}),
            ("“ระบบใช้ knowledge-based scoring engine ที่ออกแบบขึ้นในทีม ประกอบด้วย 3 sub-scores "
             "ที่สะท้อน metabolic flexibility ในเชิงสรีรวิทยา: (i) Amplitude — วัดว่าค่าครอบคลุม "
             "หลาย metabolic zone หรือไม่ (0-40 pts); (ii) Return Speed — วัด transition "
             "kinetics ระหว่าง fed ↔ fasted state (0-35 pts); (iii) Appropriateness — วัด "
             "congruence ระหว่างค่าที่วัดได้กับ context ที่ผู้ใช้รายงาน (0-25 pts) รวมเป็น "
             "0-100 คะแนน วิธีนี้จัดอยู่ในกลุ่ม Rule-based AI / Expert System และให้ค่าตัวเลขที่ "
             "อธิบาย traceable ได้ทุกขั้นตอน — เหมาะกับบริบทสุขภาพที่ต้องการ transparency สูงและ "
             "ยังไม่มี clinical-grade longitudinal dataset ให้ train supervised ML”", {}),
        ],
    )

    add_h3(doc, "6.2 คำถาม–คำตอบที่มักโดนถาม")

    add_bullet(doc, "Q: ใช้ AI ตัวไหน?")
    add_para(
        doc,
        "A: เป็น Rule-based Expert System (Explainable AI) ไม่ใช่ ML — ตั้งใจเลือกวิธีนี้เพื่อ "
        "ให้อธิบายได้ทุกจุด ตรวจสอบทางคลินิกได้ ไม่ต้องรอ dataset ระดับ 10,000+ samples "
        "และไม่มี black-box",
        indent_first=True,
    )

    add_bullet(doc, "Q: ทำไมไม่ใช้ neural network / deep learning?")
    add_para(
        doc,
        "A: จริง ๆ ระบบใช้ LSTM (deep learning, PyTorch) สำหรับ Long-term Trend อยู่แล้ว "
        "แต่ Flexibility Score เน้น interpretability ต่อผู้ใช้/แพทย์ + threshold ทาง clinical มี "
        "อยู่แล้วในงานวิจัย → ใช้ rule-based ประหยัด + โปร่งใสมากกว่า",
        indent_first=True,
    )

    add_bullet(doc, "Q: แล้วมัน scientific ไหม?")
    add_para(
        doc,
        "A: Scientific ทั้ง 2 ระดับ — (1) threshold + zone อ้างอิงงานวิจัย peer-reviewed "
        "(Anderson 2015 Obesity journal; Kelley 2000 Diabetes; Goodpaster 2017 Cell Metabolism); "
        "(2) สูตร score สามารถ audit ทางคณิตศาสตร์ได้ทั้งหมด — ไม่มี “ปาฐกถา magic number” "
        "ที่ไม่มีที่มา",
        indent_first=True,
    )

    add_bullet(doc, "Q: จะพัฒนาไปเป็น ML ในอนาคตไหม?")
    add_para(
        doc,
        "A: มีแผนไว้ — เมื่อ dataset โตพอ (~500 ผู้ใช้ × 30 sessions) สามารถ train personalised "
        "ML model เสริมได้ โดยยังคง rule-based เป็น baseline สำหรับ cold-start และ audit",
        indent_first=True,
    )

    # ═══════════════════════ 7. Citations ═══════════════════════
    add_h2(doc, "7. เอกสารอ้างอิงทางวิชาการ")

    add_bullet(
        doc,
        "Anderson JC. (2015). Measuring breath acetone for monitoring fat loss: Review. "
        "Obesity 23(12):2327-2334. doi:10.1002/oby.21242 — ที่มาของ 5-zone threshold",
    )
    add_bullet(
        doc,
        "Kelley DE, Mandarino LJ. (2000). Fuel selection in human skeletal muscle in insulin "
        "resistance. Diabetes 49(5):677-683. — บทความ seminal ของ metabolic flexibility",
    )
    add_bullet(
        doc,
        "Goodpaster BH, Sparks LM. (2017). Metabolic Flexibility in Health and Disease. "
        "Cell Metabolism 25(5):1027-1036. — modern review",
    )
    add_bullet(
        doc,
        "Buchanan BG, Shortliffe EH. (1984). Rule-Based Expert Systems: The MYCIN Experiments "
        "of the Stanford Heuristic Programming Project. — reference ของ rule-based AI approach",
    )
    add_bullet(
        doc,
        "Doshi-Velez F, Kim B. (2017). Towards A Rigorous Science of Interpretable Machine "
        "Learning. arXiv:1702.08608 — framework ที่รองรับการเลือก interpretable methods",
    )

    # ═══════════════════════ 8. Bottom line ═══════════════════════
    add_h2(doc, "8. Bottom line")

    add_para(
        doc,
        "Metabolic Flexibility Score คือ Rule-based Expert AI ที่ **ตั้งใจ** ไม่ใช่ ML — "
        "เพื่อความโปร่งใส เชื่อถือได้ อธิบายได้ ปลอดภัยกับผู้ใช้ และสอดคล้องกับ standard ของงานวิจัย "
        "breath acetone / metabolic flexibility ในปัจจุบัน วิธีนี้อยู่ในกลุ่ม “Explainable AI (XAI)” "
        "ที่งานวิจัยยุคหลังผลักดันให้ระบบสุขภาพใช้แทน black-box ML ในบริบทที่ต้อง audit ได้ "
        "การ position คะแนนนี้ในเอกสารว่าเป็น **“in-house rule-based scoring engine backed by "
        "peer-reviewed clinical thresholds”** จะสื่อได้ตรงและได้เครดิตทั้งทางวิทยาศาสตร์และทาง "
        "engineering — ดีกว่าอ้างว่าใช้ ML แล้วโดนถามลึกจนติดขัด",
        indent_first=True,
    )

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
