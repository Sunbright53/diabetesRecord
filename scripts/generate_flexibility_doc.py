"""Generate Metabolic Flexibility Score technical reference as a Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT = "/Users/ciy_th/Desktop/diabetesRecord/Metabolic_Flexibility_Reference.docx"


def set_thai_font(run, size=14, bold=False, color=None, mono=False):
    name = "Consolas" if mono else "TH Sarabun New"
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:eastAsia"), name)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_thai_font(run, size=20, bold=True)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_thai_font(run, size=16, bold=True)


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_thai_font(run, size=14, bold=True)


def add_para(doc, parts, indent_first=False):
    """parts: list of (text, {bold?, mono?}) tuples, or a plain string."""
    p = doc.add_paragraph()
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    if isinstance(parts, str):
        run = p.add_run(parts)
        set_thai_font(run, size=14)
        return
    for text, opts in parts:
        run = p.add_run(text)
        set_thai_font(
            run,
            size=14,
            bold=opts.get("bold", False),
            mono=opts.get("mono", False),
        )


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.7)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_thai_font(run, size=14)


def add_code(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    # light gray background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)
    run = p.add_run(code_text)
    set_thai_font(run, size=11, mono=True)


def add_table(doc, header, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = False
    if col_widths:
        for row in tbl.rows:
            for cell, w in zip(row.cells, col_widths):
                cell.width = w
    hdr = tbl.rows[0].cells
    for i, h in enumerate(header):
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_thai_font(run, size=13, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_thai_font(run, size=13)


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "TH Sarabun New"
    style.font.size = Pt(14)

    # ─── Header ───────────────────────────────────
    add_h1(doc, "วิธีคำนวณ Metabolic Flexibility Score")
    add_para(
        doc,
        [
            ("เอกสารอ้างอิงเชิงเทคนิคฉบับเต็ม — อธิบายทุกขั้นตอนของการคำนวณ Metabolic Flexibility Score ของระบบ MetaBreath\n", {}),
            ("Source code: ", {"bold": True}),
            ("apps/api/app/services/flexibility_engine.py", {"mono": True}),
            (" (228 บรรทัด) + endpoint ที่ ", {}),
            ("apps/api/app/routers/ai.py:424-482", {"mono": True}),
        ],
    )

    # ─── 1. Concept ───────────────────────────────
    add_h2(doc, "1. Concept — คืออะไร ทำไมมีอยู่")

    add_h3(doc, "ภาษาวิชาการ")
    add_para(
        doc,
        "Metabolic Flexibility คือ ความสามารถของร่างกายในการสลับ substrate หลักระหว่าง "
        "carbohydrate oxidation (เผาน้ำตาล) และ fatty acid oxidation (เผาไขมัน) "
        "ตาม nutrient availability และ energy demand — คนที่มี metabolic inflexibility "
        "(พบใน metabolic syndrome, type-2 diabetes, obesity) จะ “ติดค้าง” อยู่ในโหมด "
        "เผาน้ำตาลตลอด ไม่ยอมสลับไปเผาไขมันแม้จะอยู่ในสภาวะ fasting",
        indent_first=True,
    )

    add_h3(doc, "ภาษาง่าย")
    add_para(
        doc,
        "เหมือน “รถยนต์ที่วิ่งได้ทั้งเบนซินและ EV” — คนที่ metabolic ยืดหยุ่นดี "
        "จะกินคาร์บก็เผาคาร์บได้ อดข้าวก็ดึงไขมันมาใช้ได้ ไม่หิวจัด ไม่ล้า "
        "คนที่ไม่ยืดหยุ่นจะติดคาร์บอย่างเดียว ห้ามพลาดมื้อ เพราะจะ crash ทันที",
        indent_first=True,
    )
    add_para(
        doc,
        [
            ("ตัววัดที่ใช้: ", {"bold": True}),
            ("breath acetone (ppm) — เป็น proxy ของ ketone body ในเลือด → ยิ่งสูง = ร่างกายกำลังเผาไขมันมาก", {}),
        ],
    )

    # ─── 2. Input Data ────────────────────────────
    add_h2(doc, "2. Input Data — ข้อมูลที่ใช้")

    add_para(doc, [("Endpoint: ", {"bold": True}), ("POST /ai/flexibility", {"mono": True})])
    add_para(doc, [("Window: ", {"bold": True}), ("body.days (default 14 วัน)", {})])

    add_para(doc, [("Query ที่ยิงเข้า DB:", {"bold": True})])
    add_code(
        doc,
        "SELECT * FROM sensor_readings\n"
        "WHERE device_id = ? AND user_id = ? AND time >= NOW() - INTERVAL '14 days'\n"
        "ORDER BY time",
    )

    add_h3(doc, "Session grouping rule")
    add_bullet(doc, "ต่อ 1 session = readings ที่ห่างกันไม่เกิน 300 วินาที (5 นาที)")
    add_bullet(doc, "readings ห่างกัน > 5 นาที = แยกเป็น session ใหม่")
    add_bullet(doc, "ต่อ session แปลงเป็น dict ที่มี field: peak_ppm, mean_ppm, context_tag")
    add_code(
        doc,
        "{\n"
        "  \"peak_ppm\":    max(readings.acetone_delta),   # ค่าสูงสุดใน session\n"
        "  \"mean_ppm\":    mean(readings.acetone_delta),  # ค่าเฉลี่ยใน session\n"
        "  \"context_tag\": None  # ถ้ามี context จาก user เช่น \"fasting\"\n"
        "}",
    )
    add_para(
        doc,
        [
            ("หมายเหตุ: ", {"bold": True}),
            ("field ชื่อ ", {}),
            ("peak_ppm", {"mono": True}),
            (" แต่จริง ๆ backend ส่งค่า ", {}),
            ("acetone_delta", {"mono": True}),
            (" (หน่วย mV) เข้ามาโดยตรง — threshold ทั้งหมดถูก tune ตามตัวเลขนี้ ", {}),
            ("ต้องระวังถ้าจะแก้/tune", {"bold": True}),
        ],
    )

    # ─── 3. 5 Zones ───────────────────────────────
    add_h2(doc, "3. 5 Metabolic Zones (ที่ใช้ภายในเพื่อจัดหมวด)")

    add_para(
        doc,
        [
            ("ภาษาวิชาการ: ", {"bold": True}),
            ("discrete classification ตาม breath acetone level อ้างอิงจาก Anderson (2015) 5-pattern classification สำหรับ ketosis staging", {}),
        ],
    )
    add_para(doc, [("Function: ", {"bold": True}), ("_metabolic_zone(ppm) ที่ flexibility_engine.py:34-39", {"mono": True})])

    add_code(
        doc,
        "if ppm < 2:   return \"fed_resting\"      # โซน 1 — เพิ่งกิน / มีน้ำตาลในเลือดสูง\n"
        "if ppm < 8:   return \"transitional\"     # โซน 2 — เริ่มเปลี่ยนเป็น mixed oxidation\n"
        "if ppm < 40:  return \"fat_oxidation\"    # โซน 3 — เผาไขมันเต็มที่\n"
        "if ppm < 75:  return \"extended_fast\"    # โซน 4 — อดอาหารต่อเนื่อง / deep ketosis\n"
        "return \"safety_alert\"                    # โซน 5 — สูงผิดปกติ อาจเข้าใกล้ DKA",
    )
    add_para(doc, [("ภาษาง่าย: ", {"bold": True}), ("เหมือนแบ่งเกียร์รถ 5 เกียร์ตามความเร็ว engine เผาไขมันของร่างกาย", {})])

    # ─── 4. Amplitude ─────────────────────────────
    add_h2(doc, "4. Sub-score #1: Amplitude (40 คะแนน)")

    add_para(doc, [("Function: ", {"bold": True}), ("_amplitude_score(sessions) — flexibility_engine.py:42-78", {"mono": True})])

    add_h3(doc, "ภาษาวิชาการ")
    add_para(
        doc,
        "วัด breadth of metabolic dynamic range — คนที่ยืดหยุ่นควรผ่านหลายโซนใน 14 วัน "
        "(บางครั้งอิ่ม บางครั้งอด บางครั้งออกกำลัง) หากค่าอยู่โซนเดียวตลอด "
        "แสดงว่า metabolic state stuck",
        indent_first=True,
    )

    add_h3(doc, "สูตร")
    add_code(
        doc,
        "values     = [session.peak_ppm หรือ mean_ppm ของทุก session]\n"
        "zones_hit  = { _metabolic_zone(v) for v in values } - {\"safety_alert\"}\n"
        "n_zones    = |zones_hit|\n"
        "max_v      = max(values)\n"
        "mean_v     = mean(values)\n"
        "\n"
        "base = { 35    ถ้า n_zones >= 3\n"
        "       { 22    ถ้า n_zones == 2\n"
        "       { 8     ถ้า n_zones == 1\n"
        "\n"
        "fat_bonus     = 5  ถ้า max_v  >= 8    (แตะโซน fat_oxidation ได้อย่างน้อย 1 ครั้ง)\n"
        "stuck_penalty = 5  ถ้า mean_v >  30   (ค่าเฉลี่ยสูงเกินไป → ติดโซนสูง)\n"
        "\n"
        "score = min(40, base + fat_bonus - stuck_penalty)",
    )

    add_h3(doc, "ภาษาง่าย")
    add_para(
        doc,
        "วัดว่า “ตลอด 14 วัน คุณผ่านสภาวะเผาผลาญกี่แบบ” — เพิ่งกินก็ต่ำ อดก็สูง "
        "ออกกำลังก็สูง ผ่าน 3+ โซนคือดีเลิศ ผ่านโซนเดียวคือแย่ (ติดเบาะ) "
        "ถ้าแตะโซนเผาไขมัน (≥ 8 ppm) ได้อย่างน้อย 1 ครั้ง +5 bonus "
        "แต่ถ้าเฉลี่ยสูงเกิน 30 ppm ตลอด แสดงว่าติดที่โซนสูง −5 penalty",
        indent_first=True,
    )

    # ─── 5. Return Speed ───────────────────────────
    add_h2(doc, "5. Sub-score #2: Return Speed (35 คะแนน)")

    add_para(doc, [("Function: ", {"bold": True}), ("_return_speed_score(sessions) — flexibility_engine.py:81-128", {"mono": True})])

    add_h3(doc, "ภาษาวิชาการ")
    add_para(
        doc,
        "วัด transition kinetics ระหว่าง fed state ↔ fasted state — คนที่ยืดหยุ่น "
        "ระดับ acetone จะแกว่งขึ้นลงตาม nutrient state (ต่ำหลังกิน, สูงตอนอด/ออกกำลัง) "
        "หากค่าคงที่ตลอด แสดงว่า oxidation switching ไม่ทำงาน",
        indent_first=True,
    )

    add_h3(doc, "Path A — มี context_tag ครบทั้ง fasting + post_meal")
    add_code(
        doc,
        "fasting_avg  = mean( peaks ของ sessions ที่ tag = \"fasting\" หรือ \"post_exercise\" )\n"
        "postmeal_avg = mean( peaks ของ sessions ที่ tag = \"post_meal\" )\n"
        "\n"
        "ratio = postmeal_avg / fasting_avg\n"
        "        (คาดว่า postmeal < fasting → ratio ควรน้อย)\n"
        "\n"
        "score = clip( (1.5 - ratio) * 35, 0, 35 )\n"
        "        │  ratio = 0.5  → score = 35   (postmeal ต่ำมากกว่า fasting → excellent)\n"
        "        │  ratio = 1.0  → score = 17.5\n"
        "        │  ratio >= 1.5 → score = 0    (postmeal สูงเท่าหรือมากกว่า fasting → poor)",
    )

    add_h3(doc, "Path B — ไม่มี context tag (fallback proxy)")
    add_code(
        doc,
        "low_pct  = fraction ของ sessions ที่ peak < 8 ppm       (fed_resting + transitional)\n"
        "high_pct = fraction ของ sessions ที่ 8 <= peak < 75     (fat_oxidation + extended_fast)\n"
        "\n"
        "ถ้า low_pct == 0 หรือ high_pct == 0:\n"
        "    return 17.5   (neutral — ประเมินไม่ได้เพราะไม่มี variation)\n"
        "\n"
        "balance = min(low_pct, high_pct) / 0.5\n"
        "          (peak = 1.0 เมื่อ split 50/50)\n"
        "\n"
        "score = balance * 35\n"
        "        │  50/50 split → score = 35 (สมดุล)\n"
        "        │  80/20        → score = 14\n"
        "        │  95/5         → score = 3.5",
    )

    add_h3(doc, "ภาษาง่าย")
    add_para(
        doc,
        "วัดว่า “หลังกินอิ่มแล้ว ค่าลงตามที่ควรไหม” — คน metabolic ดีต้องมีทั้งค่าต่ำ "
        "(หลังกิน) และค่าสูง (อด) สลับกัน ถ้าค่าเดียวตลอด (สูงหรือต่ำ อย่างใดอย่างหนึ่ง) "
        "แปลว่ากลไก switching เพี้ยน",
        indent_first=True,
    )

    # ─── 6. Appropriateness ────────────────────────
    add_h2(doc, "6. Sub-score #3: Appropriateness (25 คะแนน)")

    add_para(doc, [("Function: ", {"bold": True}), ("_appropriateness_score(sessions) — flexibility_engine.py:131-151", {"mono": True})])

    add_h3(doc, "ภาษาวิชาการ")
    add_para(
        doc,
        "วัด context congruence — ค่า reading ตรงกับสถานะที่ user รายงานหรือไม่ "
        "ใช้ context tag เป็น ground-truth ของ physiological state",
        indent_first=True,
    )

    add_h3(doc, "Expected ranges per context (flexibility_engine.py:18-23)")
    add_table(
        doc,
        header=["context_tag", "คาดว่าอยู่ในช่วง (ppm)", "เหตุผล"],
        rows=[
            ["fasting", "2 – 40", "อดอาหาร → ควรอยู่ transitional → fat_oxidation"],
            ["post_meal", "0 – 8", "เพิ่งกิน → ควรต่ำ (fed_resting)"],
            ["post_exercise", "2 – 40", "ออกกำลัง → fat_oxidation"],
            ["evening", "0 – 20", "broad — เช้าเย็นได้ทั้งคู่"],
        ],
        col_widths=[Cm(3.5), Cm(4.5), Cm(8.0)],
    )

    add_h3(doc, "สูตร")
    add_code(
        doc,
        "tagged = [sessions ที่มี context_tag valid]\n"
        "\n"
        "ถ้าไม่มี tagged → return 12.5 (neutral)\n"
        "\n"
        "hits = 0\n"
        "สำหรับแต่ละ tagged session:\n"
        "    ppm = session.peak_ppm\n"
        "    lo, hi = expected range ของ tag นั้น\n"
        "    ถ้า lo <= ppm <= hi:\n"
        "        hits += 1\n"
        "    ถ้า tag == \"post_meal\" และ ppm < lo + 5 (คือ ppm < 5):\n"
        "        hits += 0.5   # partial credit ถ้าใกล้เคียง\n"
        "\n"
        "score = (hits / len(tagged)) * 25",
    )

    add_h3(doc, "ภาษาง่าย")
    add_para(
        doc,
        "เหมือนทดสอบ “ถ้าคุณบอกว่ากินข้าวเที่ยงมา แต่ค่าลม 30 ppm → ผิดปกติ” "
        "ตรวจว่าค่าที่วัดได้ *สมเหตุสมผล* กับ context ที่บอกไว้หรือเปล่า ยิ่งตรง ยิ่งได้คะแนน",
        indent_first=True,
    )
    add_para(
        doc,
        [
            ("ถ้า user ไม่เคยเลือก context เลย → คะแนน default = 12.5 (ครึ่ง) ", {"bold": True}),
            ("— ไม่ปรับขึ้น ไม่ปรับลง", {}),
        ],
    )

    # ─── 7. Total ─────────────────────────────────
    add_h2(doc, "7. คะแนนรวม + Trend Direction")

    add_para(doc, [("คะแนนรวม:", {"bold": True})])
    add_code(doc, "total = min(100, amplitude + return_speed + appropriateness)")

    add_para(doc, [("Trend direction — ", {"bold": True}), ("_trend_direction() ที่ flexibility_engine.py:154-165", {"mono": True})])
    add_code(
        doc,
        "แบ่ง sessions ครึ่งแรก vs ครึ่งหลัง (chronological)\n"
        "first_avg  = mean(peak_ppm ของครึ่งแรก)\n"
        "second_avg = mean(peak_ppm ของครึ่งหลัง)\n"
        "diff = second_avg - first_avg\n"
        "\n"
        "diff >  2.0   → \"increasing\"     # กำลังไต่ขึ้น\n"
        "diff < -2.0   → \"decreasing\"     # กำลังลง\n"
        "otherwise     → \"stable\"\n"
        "\n"
        "ต้องมี >= 4 sessions ถึงจะประเมินได้ (ไม่งั้น \"insufficient_data\")",
    )

    # ─── 8. Messages ──────────────────────────────
    add_h2(doc, "8. Message copy (ภาษาไทย)")

    add_para(doc, [("Function: ", {"bold": True}), ("_message_th(score, zone, context_tag) — flexibility_engine.py:168-176", {"mono": True})])

    add_table(
        doc,
        header=["Score", "Message"],
        rows=[
            ["≥ 80", "ระบบเผาผลาญยืดหยุ่นมาก — สลับระหว่างน้ำตาลและไขมันได้ดี"],
            ["≥ 60", "ระบบเผาผลาญยืดหยุ่นพอใช้ — ยังมีโอกาสพัฒนาต่อ"],
            ["≥ 40", "เริ่มมีความยืดหยุ่น — ลองวัดในหลายช่วงเวลาเพิ่มขึ้น"],
            ["< 40", "ข้อมูลยังไม่เพียงพอ — วัดเพิ่มในหลายบริบทเพื่อคำนวณความแม่นยำ"],
        ],
        col_widths=[Cm(2.5), Cm(13.5)],
    )

    add_para(
        doc,
        [
            ("Guard: ", {"bold": True}),
            ("ถ้ามี sessions < 3 → return score = 0 + message “วัดอย่างน้อย 3 ครั้งในหลายบริบทเพื่อคำนวณ Flexibility Score”", {}),
        ],
    )

    # ─── 9. Worked example ─────────────────────────
    add_h2(doc, "9. Worked example — เคส Wan (ข้อมูลจริงจาก DB)")

    add_para(doc, [("Sessions ใน 14 วัน: ", {"bold": True}), ("3 sessions ทั้งหมดเป็นวันเดียว (14 ก.ค. 2026)", {})])
    add_bullet(doc, "07:52 — peak 2.34 ppm, mean 1.83, no tag")
    add_bullet(doc, "10:14 — peak 0.81 ppm, mean 0.12, no tag")
    add_bullet(doc, "11:15 — peak 1.54 ppm, mean 0.87, no tag")

    add_h3(doc, "Amplitude")
    add_bullet(doc, "values = [2.34, 0.81, 1.54]")
    add_bullet(doc, "zones = {transitional (2.34), fed_resting (0.81, 1.54)} → 2 zones")
    add_bullet(doc, "base = 22")
    add_bullet(doc, "max = 2.34 < 8 → fat_bonus = 0")
    add_bullet(doc, "mean = 1.56 → no penalty")
    add_para(doc, [("= 22 pts", {"bold": True})])

    add_h3(doc, "Return Speed")
    add_bullet(doc, "ไม่มี tag ทั้งหมด → path B")
    add_bullet(doc, "all peaks < 8 → low_pct = 100%, high_pct = 0%")
    add_bullet(doc, "one side = 0 → return 17.5 (neutral)")
    add_para(doc, [("= 17.5 pts", {"bold": True})])

    add_h3(doc, "Appropriateness")
    add_bullet(doc, "ไม่มี tag → tagged = []")
    add_para(doc, [("= 12.5 pts (neutral)", {"bold": True})])

    add_h3(doc, "รวม")
    add_para(doc, [("Total = 22 + 17.5 + 12.5 = 52 pts", {"bold": True})])
    add_bullet(doc, "Trend: 3 sessions < 4 → \"insufficient_data\"")
    add_bullet(doc, "Zone: ใช้ peak_ppm ล่าสุด (1.54) → fed_resting")
    add_bullet(doc, "Message: \"เริ่มมีความยืดหยุ่น — ลองวัดในหลายช่วงเวลาเพิ่มขึ้น\"")

    # ─── 10. How to improve ────────────────────────
    add_h2(doc, "10. วิธีทำให้คะแนน “สูงขึ้น” (สำหรับ user)")

    add_table(
        doc,
        header=["ปัจจัย", "ทำอะไร"],
        rows=[
            ["Amplitude ต่ำ", "วัดในหลาย context — เพิ่งกิน / อด 4+ ชม / หลังออกกำลัง — ให้ค่าครอบคลุมหลายโซน"],
            ["Return Speed ต่ำ", "ถ้าค่าสูงตลอด → ลองวัดหลังกิน; ถ้าต่ำตลอด → ลองวัดตอนอดอาหาร 6+ ชม / หลังออกกำลัง"],
            ["Appropriateness ต่ำ", "ตอนกด START ให้เลือก context tag ทุกครั้ง (fasting/post_meal/post_exercise/evening) — ระบบจะเทียบว่าค่าตรงคาดหมายไหม"],
            ["ค่ารวม < 40", "เก็บข้อมูลเพิ่ม — น้อยกว่า 3 sessions ยังไม่คำนวณ, มี 3–5 sessions ก็ยังไม่พอ trend"],
        ],
        col_widths=[Cm(4.0), Cm(12.0)],
    )

    # ─── 11. Assumptions ───────────────────────────
    add_h2(doc, "11. Assumption + Limitation (ข้อสังเกตทางเทคนิค)")

    add_bullet(doc, "Session gap = 5 นาที — reading ห่างกันเกิน 300 วิ = คนละ session; ต่ำกว่านั้นยุบเป็นอันเดียว")
    add_bullet(doc, "Field naming lie — code ใช้ชื่อ peak_ppm แต่เก็บค่า acetone_delta (mV) โดยตรง; threshold (2/8/40/75) ถูก tune ให้ทำงานกับตัวเลข mV ตรง ๆ (บังเอิญ Anderson 2015 papers กำหนด ppm threshold ที่คล้ายกัน)")
    add_bullet(doc, "Zone safety_alert (≥ 75) ถูก excluded จาก zones_hit ตอนคำนวณ amplitude → คน DKA ที่ค่าสูงตลอดจะเข้าโซนเดียว → คะแนนต่ำ (ตั้งใจ)")
    add_bullet(doc, "Model-free — ไม่ใช้ ML ทั้งหมด เป็น rule-based heuristic ล้วน ๆ (โปร่งใส เข้าใจได้ อธิบายได้)")
    add_bullet(doc, "ไม่ทดสอบผ่าน RCT — เป็น monitoring signal ไม่ใช่ diagnostic tool (ต่างจาก HOMA-IR หรือ oral glucose tolerance test)")

    # ─── 12. Reference ────────────────────────────
    add_h2(doc, "12. Reference")

    add_bullet(doc, "Anderson JC. Measuring breath acetone for monitoring fat loss: Review. Obesity (2015) 23(12):2327-2334 — ใช้เป็น reference สำหรับ zone thresholds")
    add_bullet(doc, "Kelley DE, Mandarino LJ. Fuel selection in human skeletal muscle in insulin resistance. Diabetes (2000) 49(5):677-683 — concept origin ของ metabolic flexibility")
    add_bullet(doc, "Goodpaster BH, Sparks LM. Metabolic Flexibility in Health and Disease. Cell Metabolism (2017) 25(5):1027-1036 — modern review")

    add_para(
        doc,
        [
            ("คะแนน 0–100 ที่ใช้ในระบบนี้ไม่มี clinical validation ", {"bold": True}),
            ("— เป็น engineering heuristic ที่คำนวณจาก 3 มิติที่อธิบายด้านบน "
             "เอาไว้ track ตัวเอง หรือสังเกตแนวโน้ม ไม่ใช้ตัดสินสุขภาพในเชิงการแพทย์", {}),
        ],
    )

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
