"""Generate 'MetaBreath Chat + MCP' Word document.

A Thai-language architecture + operations reference covering:
  - What the chat + MCP subsystem does end-to-end
  - The 8 tools / 3 resources / 3 prompts we expose
  - Request/response flow across FastAPI + Anthropic + FastMCP
  - Persona rules and how to tune them
  - Auth paths for the in-process chat and external Claude Desktop clients
  - Config snippets + operational commands

Style matches the other generate_*.py scripts (TH Sarabun New, mono blocks).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT = "/Users/ciy_th/Desktop/diabetesRecord/MetaBreath_Chat_MCP_Documentation.docx"

MINT = RGBColor(0x01, 0xD1, 0x9B)
NAVY = RGBColor(0x14, 0x25, 0x52)
GREY = RGBColor(0x55, 0x5F, 0x6D)


def set_thai_font(run, size=14, bold=False, color=None, mono=False):
    name = "Consolas" if mono else "TH Sarabun New"
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:eastAsia"), name)


def add_title(doc, text, size=32):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_thai_font(run, size=size, bold=True, color=NAVY)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(text)
    set_thai_font(run, size=16, color=MINT)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_thai_font(run, size=22, bold=True, color=NAVY)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_thai_font(run, size=18, bold=True, color=NAVY)


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_thai_font(run, size=15, bold=True, color=MINT)


def add_p(doc, text, size=14, bold=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_thai_font(run, size=size, bold=bold, color=color)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8 + 0.4 * level)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_thai_font(run, size=14)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_thai_font(run, size=11, mono=True, color=GREY)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cell.paragraphs[0].add_run(h)
        set_thai_font(run, size=13, bold=True, color=NAVY)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            run = cell.paragraphs[0].add_run(str(val))
            set_thai_font(run, size=12)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # ─── Cover ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    add_title(doc, "MetaBreath Chat + MCP")
    add_subtitle(doc, "สถาปัตยกรรมและคู่มือระบบแชท AI ผ่าน Model Context Protocol")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("โครงการ MetaBreath — Cheewarun")
    set_thai_font(run, size=14, color=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("commit ที่ใช้อ้างอิง: f807dfe → 819ce08")
    set_thai_font(run, size=12, color=GREY, mono=True)

    doc.add_page_break()

    # ─── บทนำ ───────────────────────────────────────────────────────────
    add_h1(doc, "1. บทนำ — MetaBreath Chat คืออะไร")
    add_p(doc,
        "MetaBreath Chat คือระบบผู้ช่วย AI สนทนาที่รวมข้อมูลสุขภาพส่วนบุคคล "
        "(โปรไฟล์, ค่าลมหายใจจาก sensor, บันทึกกิจกรรม) เข้ากับความสามารถของ "
        "โมเดล Claude Haiku 4.5 ผ่าน Model Context Protocol (MCP) เพื่อให้ "
        "ผู้ใช้ถามได้เหมือนคุยกับผู้เชี่ยวชาญที่รู้ข้อมูลของคุณจริง ๆ")
    add_p(doc,
        "ระบบตอบเป็นภาษาไทย บุคลิก 'MetaBreath' (ผู้หญิงสุภาพ เป็นกันเอง), "
        "รองรับการวิเคราะห์ค่าลมหายใจ, ให้คำแนะนำพฤติกรรมสุขภาพเชิงบวก, "
        "และผูก AI กับข้อมูลจริงในฐานข้อมูลผ่านเครื่องมือ (tools) มาตรฐาน MCP")

    add_h2(doc, "จุดเด่นของระบบ")
    add_bullet(doc, "ใช้ MCP protocol จริง — ไม่ใช่ tool_use ธรรมดา")
    add_bullet(doc, "Streaming ตอบทีละคำ (SSE) — UX ลื่นเหมือน ChatGPT")
    add_bullet(doc, "Slash commands สำหรับคำถามที่ใช้บ่อย (/summary_today ฯลฯ)")
    add_bullet(doc, "Claude Desktop ต่อได้ผ่าน HTTPS + JWT")
    add_bullet(doc, "AI เข้าถึงข้อมูลผู้ใช้ตัวเองเท่านั้น (auth scope)")
    add_bullet(doc, "ไม่แสดง disclaimer footer + ไม่ใช้ emoji (มืออาชีพ)")

    doc.add_page_break()

    # ─── MCP คืออะไร ────────────────────────────────────────────────────
    add_h1(doc, "2. Model Context Protocol (MCP) คืออะไร")
    add_p(doc,
        "MCP เป็นมาตรฐานเปิดที่กำหนดวิธีให้ AI คุยกับ 'บริบทภายนอก' อย่างเป็นระเบียบ "
        "โดยแยก 3 primitives ที่ AI ใช้ได้:")

    add_h2(doc, "3 primitives ของ MCP")
    add_table(doc,
        ["Primitive", "หน้าที่", "ของเรามีกี่ตัว"],
        [
            ["Tools", "ฟังก์ชันที่ AI เรียกได้ (ทั้ง read + write)", "8"],
            ["Resources", "ข้อมูล read-only ที่ AI อ่านได้", "3"],
            ["Prompts", "Template คำถามให้ user เลือกใช้ (slash commands)", "3"],
        ])

    add_p(doc,
        "สรุปสั้น: Tools = ให้ AI ทำงาน, Resources = ให้ AI อ่านเอกสารอ้างอิง, "
        "Prompts = ให้ user คลิกใช้แทนพิมพ์เอง")

    add_h2(doc, "ทำไมต้อง MCP")
    add_bullet(doc, "Standardized: schema เดียวกันใช้ได้กับ Claude Desktop / IDE / เว็บ")
    add_bullet(doc, "Auth-safe: tools scope ตาม user ที่ล็อกอินอยู่")
    add_bullet(doc, "Reusable: MCP server เดียว ใช้ได้หลาย client")
    add_bullet(doc, "Future-proof: ถ้ามี MCP client ใหม่ ๆ ระบบเราต่อได้ทันที")

    doc.add_page_break()

    # ─── สถาปัตยกรรม ────────────────────────────────────────────────────
    add_h1(doc, "3. สถาปัตยกรรมของระบบ")

    add_h2(doc, "โครงสร้างไฟล์ (backend)")
    add_code(doc,
        "apps/api/app/\n"
        "├── mcp_context.py         ← contextvars + mcp_scope() helper\n"
        "├── mcp_server.py          ← FastMCP: 8 tools + 3 resources + 3 prompts\n"
        "├── main.py                ← mount /mcp + MCPAuthMiddleware\n"
        "├── routers/ai.py\n"
        "│   ├── POST /ai/chat           (non-stream)\n"
        "│   ├── POST /ai/chat/stream    (SSE — UI ใช้ตัวนี้)\n"
        "│   └── GET  /ai/prompts        (slash menu)\n"
        "├── services/\n"
        "│   ├── chat_tools.py       ← implementation ของ 8 tools\n"
        "│   └── llm_guardrail.py    ← persona + sanitiser + refusal")

    add_h2(doc, "โครงสร้างสองเส้นทาง (dual-path)")
    add_p(doc,
        "FastMCP instance เดียวกันถูกใช้จากสองที่:")
    add_bullet(doc, "เส้น A — Chat ในเว็บ: FastAPI /ai/chat/stream เรียก MCP client แบบ in-process (ไม่มี HTTP hop)")
    add_bullet(doc, "เส้น B — Claude Desktop / IDE: HTTPS → /api/mcp/ → Streamable HTTP transport")
    add_p(doc,
        "ทั้งสองเส้นเรียก tools ตัวเดียวกัน แชร์ business logic ใน chat_tools.py")

    doc.add_page_break()

    # ─── 8 Tools ───────────────────────────────────────────────────────
    add_h1(doc, "4. Tools (8 ตัว)")
    add_p(doc, "AI เลือกเรียกเองตามคำถาม — ไม่ต้องให้ผู้ใช้ระบุ")

    add_h2(doc, "Tools ประเภทอ่านข้อมูล (Read)")

    add_h3(doc, "get_user_profile")
    add_p(doc, "ดึงโปรไฟล์ผู้ใช้: display_name, sex, อายุ (คำนวณจาก dob), "
                "height_cm, weight_kg, BMI (คำนวณ), goal_type, สถานะ onboarding")

    add_h3(doc, "get_recent_readings (days: int = 30, limit: int = 10)")
    add_p(doc, "ดึงค่า sensor ลมหายใจล่าสุด acetone_delta, label, confidence, "
                "quality_score. ทำงานได้แม้อุปกรณ์ไม่ได้เชื่อมอยู่ (query by user_id)")
    add_p(doc, "Default window ถูกขยับเป็น 30 วัน + max 90 วัน เพื่อให้ AI เห็นประวัติ")

    add_h3(doc, "get_metabolic_trend (days: int = 14)")
    add_p(doc, "วิเคราะห์แนวโน้ม acetone (increasing/decreasing/stable) พร้อม slope + "
                "confidence โดยใช้ linear regression จาก ml_inference.predict_trend()")

    add_h3(doc, "get_recent_logs (days: int = 30)")
    add_p(doc, "ดึงบันทึก: meals (มื้ออาหาร), activities (กิจกรรม), weights (น้ำหนัก), "
                "ketones (blood/urine ketone) — ทั้งหมดใน N วันย้อนหลัง")

    add_h3(doc, "explain_reading (acetone_ppm: float, context?)")
    add_p(doc, "แปลค่า ppm หนึ่งค่าเป็นสถานะเมตาบอลิก + คำอธิบาย ไม่แตะ DB")

    add_h2(doc, "Tools ประเภทเขียนข้อมูล (Write)")

    add_h3(doc, "log_meal (name, kcal?, carbs_g?)")
    add_p(doc, "บันทึกมื้ออาหารของ user เข้า meal_logs — ใช้เวลาผู้ใช้บอกว่า "
                "'จดว่ากินไข่ต้ม 2 ฟอง'")

    add_h3(doc, "log_activity (kind, duration_min, kcal?)")
    add_p(doc, "บันทึกกิจกรรม (walk / run / cycle / gym / yoga) เข้า activity_logs")

    add_h3(doc, "calibrate_device (baseline_voc, temp_c?, humidity_pct?)")
    add_p(doc, "Zero-point calibrate อุปกรณ์ MetaBreath ด้วยค่า ambient VOC baseline. "
                "เขียน device_calibration + clear flag needs_recalibration")

    doc.add_page_break()

    # ─── Resources ─────────────────────────────────────────────────────
    add_h1(doc, "5. Resources (3 ตัว)")
    add_p(doc, "ข้อมูล read-only ที่ AI + external clients (Claude Desktop) อ่านได้")

    add_table(doc,
        ["URI", "เนื้อหา"],
        [
            ["metabreath://reference/acetone-ranges",
             "ค่าอ้างอิง acetone ppm แบ่งตามสถานะเมตาบอลิก (healthy / fat_burning / ketosis / elevated)"],
            ["metabreath://reference/tgs1820-datasheet",
             "สเปคเซ็นเซอร์ TGS1820: cross-sensitivity, operating range, drift"],
            ["metabreath://user/personal-snapshot",
             "Snapshot เฉพาะบุคคล — profile + reading ล่าสุด + สรุปจำนวน logs 7 วัน"],
        ])

    doc.add_page_break()

    # ─── Prompts ───────────────────────────────────────────────────────
    add_h1(doc, "6. Prompts / Slash Commands (3 ตัว)")
    add_p(doc, "เมนูคำถามที่ผู้ใช้เลือกใช้ได้จาก UI (พิมพ์ '/' → เห็น dropdown)")

    add_table(doc,
        ["Slash", "จุดประสงค์", "Template"],
        [
            ["/summary_today", "สรุปสุขภาพวันนี้",
             "โปรไฟล์ + ค่าล่าสุด + logs 7 วัน + คำแนะนำ 1-2 ข้อ"],
            ["/daily_coaching", "ข้อความโค้ชประจำวัน",
             "ตาม goal_type — พร้อมกิจกรรมที่ทำได้ 1 อย่าง"],
            ["/analyze_metabolic", "วิเคราะห์เมตาบอลิก",
             "จากข้อมูล 14 วัน — ค่าเฉลี่ย + แนวโน้ม + คำแนะนำ 2-3 ข้อ"],
        ])

    doc.add_page_break()

    # ─── Flow ──────────────────────────────────────────────────────────
    add_h1(doc, "7. Request/Response Flow เต็ม")
    add_p(doc, "ตัวอย่างเมื่อผู้ใช้พิมพ์ 'ค่าลมหายใจล่าสุดเป็นไง'")

    add_code(doc,
        "[User พิมพ์ในกล่องแชท]\n"
        "      |\n"
        "      v  POST /api/ai/chat/stream  {message, device_id}\n"
        "[FastAPI /ai/chat/stream]\n"
        "      |\n"
        "      +-- 1. Guardrail: มีคำต้องห้ามไหม (ยา/วินิจฉัย)?\n"
        "      |       -> มี -> emit refusal event -> done\n"
        "      |\n"
        "      +-- 2. Load Profile -> user_context\n"
        "      |       (display_name, goal_th, sex, ...)\n"
        "      |\n"
        "      +-- 3. Build system prompt (persona + few-shot + context)\n"
        "      |\n"
        "      +-- 4. mcp_scope(user.id, device_id)\n"
        "      |       [ContextVars set — tools อ่านได้]\n"
        "      |       |\n"
        "      |       v\n"
        "      |   [MCP client in-process]\n"
        "      |       |\n"
        "      |       +-- list_tools()  <- get 8 schemas from FastMCP\n"
        "      |       |\n"
        "      |       +-- Loop <=6 รอบ:\n"
        "      |       |     Anthropic messages.stream(model, tools, ...)\n"
        "      |       |         text delta -> emit SSE 'text'\n"
        "      |       |         stop_reason == 'tool_use'?\n"
        "      |       |             YES: emit 'tool_use', mcp.call_tool(),\n"
        "      |       |                  emit 'tool_result', continue\n"
        "      |       |             NO:  break\n"
        "      |       |\n"
        "      |       +-- exit scope (contextvars reset)\n"
        "      |\n"
        "      +-- 5. Sanitize: strip emoji + strip any self-written disclaimer\n"
        "      |       (ไม่ append disclaimer)\n"
        "      |\n"
        "      +-- emit SSE 'done' -> close stream\n"
        "\n"
        "[Web chat/page.tsx]\n"
        "      |\n"
        "      +-- SSE parser (fetch + ReadableStream)\n"
        "      +-- event 'text'         -> append delta to assistant message\n"
        "      +-- event 'tool_use'     -> show 'กำลังดึงข้อมูล: <name>'\n"
        "      +-- event 'tool_result'  -> ซ่อน status\n"
        "      +-- event 'done'         -> ปิด loading\n"
        "      +-- render ด้วย react-markdown + remark-gfm")

    doc.add_page_break()

    # ─── Auth ──────────────────────────────────────────────────────────
    add_h1(doc, "8. Authentication (สองเส้นทาง)")

    add_h2(doc, "เส้น A: Chat ในเว็บ")
    add_bullet(doc, "Frontend ส่ง Authorization: Bearer <access_token>")
    add_bullet(doc, "FastAPI Depends(get_current_user) → ได้ User object")
    add_bullet(doc, "mcp_scope(user.id, device_id) ตั้ง contextvar")
    add_bullet(doc, "Tools อ่าน user_id + db + device_id ผ่าน get_user_id() / get_db()")

    add_h2(doc, "เส้น B: External MCP Client (Claude Desktop / IDE)")
    add_bullet(doc, "Client ส่ง Authorization: Bearer <access_token>")
    add_bullet(doc, "Optional: X-MetaBreath-Device-Id: <uuid>")
    add_bullet(doc, "MCPAuthMiddleware validate JWT + สร้าง DB session → ตั้ง contextvar")
    add_bullet(doc, "FastMCP ASGI handle: list_tools / call_tool / read_resource ฯลฯ")

    add_h2(doc, "Config ตัวอย่างสำหรับ Claude Desktop")
    add_code(doc,
        '{\n'
        '  "mcpServers": {\n'
        '    "metabreath": {\n'
        '      "url": "https://metabreath.duckdns.org/api/mcp/",\n'
        '      "headers": {\n'
        '        "Authorization": "Bearer <YOUR_JWT>",\n'
        '        "X-MetaBreath-Device-Id": "<optional_device_uuid>"\n'
        '      }\n'
        '    }\n'
        '  }\n'
        '}')

    doc.add_page_break()

    # ─── Persona ───────────────────────────────────────────────────────
    add_h1(doc, "9. Persona ของ MetaBreath")

    add_h2(doc, "บุคลิกและโทน")
    add_bullet(doc, "ผู้หญิงสุภาพ อบอุ่น เป็นกันเอง — เหมือนเพื่อนที่มีความรู้")
    add_bullet(doc, "ตอบภาษาไทย (สลับอังกฤษได้ถ้าผู้ใช้พิมพ์อังกฤษ)")
    add_bullet(doc, "ใช้ 'ค่ะ/นะคะ' อย่างเป็นธรรมชาติ ไม่ทุกประโยค")
    add_bullet(doc, "ตอบสั้น 2-4 ประโยคเป็นหลัก ยาวได้ถ้าถูกขอ deep dive")
    add_bullet(doc, "ห้ามใช้ตาราง (|) ห้ามใช้หัวข้อ (###) ในการแชท")
    add_bullet(doc, "ห้ามใช้ emoji ทุกกรณี")
    add_bullet(doc, "ห้ามใส่ disclaimer footer")
    add_bullet(doc, "ห้ามคำซ้ำติดกัน (ลองลอง, ดีดี) หรือซ้อน emphasizer 3 ชั้น")

    add_h2(doc, "กฎ Tool-first")
    add_p(doc,
        "ก่อนจะพูดว่า 'ยังไม่มีข้อมูล' AI ต้องเรียก get_recent_readings(30) ก่อนเสมอ. "
        "ถ้า 30 วันไม่มี ต้องลอง 90 วัน. ถ้ายังไม่มีจริงถึงบอกได้ว่ายังไม่มีค่า. "
        "เป้าคือให้ AI ใช้ประโยชน์จาก historical data สูงสุด")

    add_h2(doc, "กฎความปลอดภัย (เก็บภายในระบบ)")
    add_bullet(doc, "ห้ามวินิจฉัยโรค — พูดว่า 'ค่ากำลังชี้ไปทาง...' แทน")
    add_bullet(doc, "ห้ามแนะนำยา / ปรับขนาดยา / ฉีดยา")
    add_bullet(doc, "ห้ามพูด 'ไม่ต้องไปหาหมอ'")
    add_bullet(doc, "อาการฉุกเฉิน (เจ็บหน้าอก, หมดสติ ฯลฯ) → ตอบ 'โปรดโทร 1669'")

    add_h2(doc, "การ Sanitize ก่อนตอบ user")
    add_bullet(doc, "Regex ตัดข้อความที่มีคำห้าม → แทนด้วย [ข้อมูลนี้ถูกซ่อน...]")
    add_bullet(doc, "ตัดข้อความจาก signature disclaimer เป็นต้นไปทิ้ง (เผื่อ AI ใส่เอง)")
    add_bullet(doc, "Strip emoji ทั้งหมด (Unicode blocks)")

    doc.add_page_break()

    # ─── Frontend ──────────────────────────────────────────────────────
    add_h1(doc, "10. Frontend Integration")

    add_h2(doc, "หน้า chat")
    add_p(doc, "อยู่ที่ apps/web/src/app/(app)/chat/page.tsx")
    add_bullet(doc, "ใช้ api.ai.chatStream() (SSE) แทน api.ai.chat() (single-shot)")
    add_bullet(doc, "Render ด้วย react-markdown + remark-gfm")
    add_bullet(doc, "แสดง typing indicator + tool status ('กำลังดึงข้อมูล: <tool>')")
    add_bullet(doc, "Slash menu โผล่เมื่อพิมพ์ '/' ต้นข้อความ")

    add_h2(doc, "SSE parser (apps/web/src/lib/api.ts → chatStream)")
    add_code(doc,
        "const res = await fetch('/api/ai/chat/stream', {method:'POST', headers, body});\n"
        "const reader = res.body.getReader();\n"
        "const decoder = new TextDecoder();\n"
        "let buf = '';\n"
        "while (true) {\n"
        "  const { value, done } = await reader.read();\n"
        "  if (done) break;\n"
        "  buf += decoder.decode(value, { stream: true });\n"
        "  // split by SSE '\\n\\n' delimiter\n"
        "  // parse each 'data: {...}' JSON payload\n"
        "  // callback onEvent(payload)\n"
        "}")

    add_h2(doc, "Event ประเภทที่ frontend รับ")
    add_table(doc,
        ["type", "payload", "การใช้งาน"],
        [
            ["text", "{delta: string}", "ต่อ delta เข้า assistant message"],
            ["tool_use", "{name: string}", "แสดง 'กำลังดึงข้อมูล: <name>'"],
            ["tool_result", "{name: string}", "ซ่อน status"],
            ["refusal", "{reply: string}", "ตั้ง content + refusal=true"],
            ["error", "{message: string}", "แสดงข้อความผิดพลาด"],
            ["done", "{}", "ปิด loading state"],
        ])

    doc.add_page_break()

    # ─── ปรับแต่ง ──────────────────────────────────────────────────────
    add_h1(doc, "11. คู่มือการปรับแต่งระบบ")

    add_h2(doc, "แก้ persona / โทน / กฎการตอบ")
    add_p(doc, "ทั้งหมดอยู่ในไฟล์เดียว: apps/api/app/services/llm_guardrail.py")
    add_bullet(doc, "SYSTEM_PROMPT_TEMPLATE — บุคลิก + กฎ + few-shot examples")
    add_bullet(doc, "BANNED_PATTERNS_TH / EN — regex คำต้องห้าม (guardrail)")
    add_bullet(doc, "build_refusal_response — ข้อความปฏิเสธ")
    add_bullet(doc, "_EMOJI_RE — regex strip emoji")

    add_h2(doc, "เพิ่ม tool ใหม่")
    add_p(doc, "ขั้นตอน:")
    add_bullet(doc, "1. เขียน impl ใน chat_tools.py (async fn รับ db, user, ...)")
    add_bullet(doc, "2. Wrap ด้วย @mcp.tool() ใน mcp_server.py")
    add_bullet(doc, "3. ทดสอบ + deploy — AI เห็น tool ใหม่ทันที (list_tools() re-discovers)")

    add_h2(doc, "เพิ่ม slash command (Prompt)")
    add_bullet(doc, "1. Add @mcp.prompt(description='...') function ใน mcp_server.py")
    add_bullet(doc, "2. Return text template")
    add_bullet(doc, "3. UI ดึงจาก GET /ai/prompts อัตโนมัติ — ไม่ต้องแก้ frontend")

    add_h2(doc, "เพิ่ม Resource")
    add_bullet(doc, "1. Add @mcp.resource('metabreath://...') function")
    add_bullet(doc, "2. Return string (JSON แนะนำ)")
    add_bullet(doc, "3. Claude Desktop เห็น resource ใหม่ทันที")

    add_h2(doc, "แก้พฤติกรรม 'ห้าม emoji' หรือ 'ไม่ใส่ disclaimer'")
    add_bullet(doc, "Emoji: ปิด _EMOJI_RE strip ใน sanitise_response()")
    add_bullet(doc, "Disclaimer: return llm_response.rstrip() + DISCLAIMER_TH (แทนบรรทัดเดิม)")

    doc.add_page_break()

    # ─── Ops ──────────────────────────────────────────────────────────
    add_h1(doc, "12. Operations & Deployment")

    add_h2(doc, "Environment variables ที่ต้องมี")
    add_bullet(doc, "CLAUDE_API_KEY (หรือ ANTHROPIC_API_KEY) — Anthropic API key")
    add_bullet(doc, "CLAUDE_MODEL — default claude-haiku-4-5-20251001")
    add_bullet(doc, "SECRET_KEY — JWT signing")
    add_bullet(doc, "DATABASE_URL — Postgres/TimescaleDB")

    add_h2(doc, "Deploy backend หลังแก้โค้ด")
    add_code(doc,
        "# สั่งจาก repo root:\n"
        "rsync -avz --delete \\\n"
        "  --exclude '__pycache__' --exclude '*.pyc' \\\n"
        "  apps/api/app/ \\\n"
        "  root@metabreath.duckdns.org:/root/cheewarun/apps/api/app/\n"
        "\n"
        "ssh root@metabreath.duckdns.org \\\n"
        "  \"cd /root/cheewarun && docker compose up -d --build api worker beat\"")

    add_h2(doc, "Deploy web หลังแก้ frontend")
    add_code(doc,
        "rsync -avz apps/web/src/ \\\n"
        "  root@metabreath.duckdns.org:/root/cheewarun/apps/web/src/\n"
        "\n"
        "ssh root@metabreath.duckdns.org \"cd /root/cheewarun && \\\n"
        "  docker build -t cheewarun-web:latest apps/web/ && \\\n"
        "  docker compose up --no-build -d web\"")

    add_h2(doc, "ดู logs")
    add_code(doc,
        "ssh root@metabreath.duckdns.org \\\n"
        "  \"docker compose -f /root/cheewarun/docker-compose.yml logs -f api\"")

    add_h2(doc, "ทดสอบ endpoint โดยไม่ต้อง login (ควรได้ 403/401)")
    add_code(doc,
        "curl -o /dev/null -w '%{http_code}\\n' \\\n"
        "  -X POST https://metabreath.duckdns.org/api/ai/chat \\\n"
        "  -H 'Content-Type: application/json' -d '{\"message\":\"t\"}'\n"
        "\n"
        "curl -o /dev/null -w '%{http_code}\\n' \\\n"
        "  https://metabreath.duckdns.org/api/mcp/")

    add_h2(doc, "ทดสอบ MCP endpoint พร้อม JWT")
    add_code(doc,
        "TOKEN='<paste JWT>'\n"
        "curl -X POST https://metabreath.duckdns.org/api/mcp/ \\\n"
        "  -H \"Authorization: Bearer $TOKEN\" \\\n"
        "  -H 'Content-Type: application/json' \\\n"
        "  -d '{\"jsonrpc\":\"2.0\",\"id\":1,\"method\":\"tools/list\"}'")

    doc.add_page_break()

    # ─── Trade-offs ────────────────────────────────────────────────────
    add_h1(doc, "13. Trade-offs และการตัดสินใจ")
    add_table(doc,
        ["ประเด็น", "เลือกอะไร", "ทำไม"],
        [
            ["Transport ใน chat", "in-process MCP client",
             "Latency ~0, ไม่ต้อง HTTP hop"],
            ["Transport ภายนอก", "Streamable HTTP (mounted /mcp)",
             "Claude Desktop รองรับได้"],
            ["MCP SDK version", "mcp==1.9.4",
             "รองรับ FastMCP + streamable_http_app"],
            ["Streaming", "SSE (fetch+ReadableStream)",
             "รองรับ Bearer header (EventSource ทำไม่ได้)"],
            ["Persona อยู่ไหน", "llm_guardrail.py",
             "รวมกับ guardrail rules ที่ที่เดียว"],
            ["Disclaimer", "ไม่ใส่ให้",
             "user request: ให้ระบบรู้กฎ แต่ไม่โชว์ footer"],
            ["Emoji", "strip ทั้งหมด",
             "user request: โทนมืออาชีพ ไม่ใช้ emoji"],
            ["History window default", "30 วัน",
             "AI เห็นประวัติพอสำหรับตอบทั่วไป"],
            ["History window max", "90 วัน",
             "พอสำหรับ deep analysis; ไม่โหลด DB หนัก"],
        ])

    doc.add_page_break()

    # ─── Future ────────────────────────────────────────────────────────
    add_h1(doc, "14. Future Enhancements")
    add_bullet(doc, "Keyboard nav ใน slash menu (ลูกศร + Enter)")
    add_bullet(doc, "Multi-turn conversation memory — เก็บ chat history ใน DB")
    add_bullet(doc, "Resource browsing UI — user เปิดดู personal-snapshot ตรง ๆ ได้")
    add_bullet(doc, "MCP prompts args — รับ input parameter ก่อนส่ง (ต้องปรับ UI ให้กรอก)")
    add_bullet(doc, "Parallel tool calls — เมื่อ AI ขอ 3 tools พร้อมกัน เรียกขนาน")
    add_bullet(doc, "Voice input/output — TTS/STT")
    add_bullet(doc, "Feedback loop — 'คำตอบนี้มีประโยชน์ไหม' → เก็บใน DB สำหรับ tune")

    # ─── Summary ───────────────────────────────────────────────────────
    add_h1(doc, "15. TL;DR")
    add_p(doc,
        "Chat ในเว็บ = FastAPI → in-process MCP client → FastMCP (8 tools + "
        "3 resources + 3 prompts) → Anthropic streaming → SSE ออกไปเว็บ → "
        "react-markdown render. FastMCP instance เดียวกันเปิดที่ /api/mcp/ "
        "ให้ Claude Desktop ต่อผ่าน HTTPS + JWT ได้. บุคลิก MetaBreath อยู่ใน "
        "llm_guardrail.py — แก้ที่นั่นถ้าจะปรับโทน.")

    doc.save(OUTPUT)
    print(f"Wrote: {OUTPUT}")


if __name__ == "__main__":
    build()
