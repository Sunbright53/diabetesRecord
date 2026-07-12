"""Generate hardware specification section as a Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT = "/Users/ciy_th/Desktop/diabetesRecord/5.5_Hardware_Specification.docx"


def set_thai_font(run, size=14, bold=False, color=None):
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "TH Sarabun New")
    rfonts.set(qn("w:hAnsi"), "TH Sarabun New")
    rfonts.set(qn("w:cs"), "TH Sarabun New")
    rfonts.set(qn("w:eastAsia"), "TH Sarabun New")


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_thai_font(run, size=18, bold=True)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_thai_font(run, size=16, bold=True)


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_thai_font(run, size=15, bold=True)


def add_para(doc, text, indent_first=True):
    p = doc.add_paragraph()
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_thai_font(run, size=14)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.7)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_thai_font(run, size=14)


def add_table(doc, header, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = False
    if col_widths:
        for row in tbl.rows:
            for cell, w in zip(row.cells, col_widths):
                cell.width = w
    hdr = tbl.rows[0].cells
    for i, h in enumerate(header):
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_thai_font(run, size=13, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_thai_font(run, size=13)


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "TH Sarabun New"
    style.font.size = Pt(14)

    # ═══════════════════════ HEADER ═══════════════════════
    add_h1(doc, "5.5 คุณลักษณะของอุปกรณ์ที่ใช้กับโปรแกรม (Hardware Specification)")

    add_para(
        doc,
        "หัวข้อนี้อธิบายคุณลักษณะของอุปกรณ์ฮาร์ดแวร์และสิ่งแวดล้อมทางเทคนิคที่ระบบ MetaBreath ใช้งาน "
        "ครอบคลุมทั้ง (1) อุปกรณ์ตรวจวัดที่พัฒนาขึ้น (MetaBreath Device), (2) เซนเซอร์และชิ้นส่วนอิเล็กทรอนิกส์ "
        "ประกอบภายใน, (3) อุปกรณ์ฝั่งผู้ใช้ (client) ที่ใช้เข้าถึง web application, (4) อุปกรณ์ฝั่งเซิร์ฟเวอร์ "
        "และ (5) โครงสร้างเครือข่ายและเงื่อนไขการใช้งาน",
        indent_first=True,
    )

    # ═══════════════════════ 5.5.1 MetaBreath Device ═══════════════════════
    add_h2(doc, "5.5.1 อุปกรณ์ตรวจวัดลมหายใจ (MetaBreath Device)")

    add_para(
        doc,
        "อุปกรณ์ MetaBreath เป็นเครื่องต้นแบบที่ทีมพัฒนาขึ้นเอง ใช้ ESP32 เป็นไมโครคอนโทรลเลอร์กลาง "
        "ทำหน้าที่อ่านค่าจากเซนเซอร์ตรวจจับสารระเหย (VOC), เซนเซอร์แรงดัน (Pressure) และเซนเซอร์อุณหภูมิ/ความชื้น "
        "แล้วส่งข้อมูลผ่าน Wi-Fi ไปยัง MQTT broker แบบเวลาจริง โดยมีคุณลักษณะโดยรวมดังนี้",
        indent_first=True,
    )

    add_table(
        doc,
        header=["คุณสมบัติ", "รายละเอียด"],
        rows=[
            ["ขนาดโดยประมาณ", "กล่องพลาสติกขนาดพกพา ประมาณ 10 × 7 × 4 ซม."],
            ["แหล่งจ่ายไฟ", "USB-C 5V หรือ Power Bank"],
            ["กระแสไฟฟ้าที่ใช้", "ประมาณ 200–400 mA ระหว่างการทำงาน"],
            ["การเชื่อมต่อ", "Wi-Fi 2.4 GHz (802.11 b/g/n)"],
            ["Protocol การส่งข้อมูล", "MQTT ผ่าน TCP port 1883"],
            ["ระยะทำงาน", "ในระยะสัญญาณ Wi-Fi ปกติ (~ 15 เมตร)"],
            ["การตั้งค่า Wi-Fi", "captive portal ผ่าน WiFiManager (ไม่ต้องใช้สาย USB)"],
            ["สถานะแสดงผล", "LED บนบอร์ด แสดงสถานะการเชื่อมต่อและการวัด"],
            ["Sampling Rate", "1 Hz (1 reading/วินาที)"],
            ["Baseline Calibration", "อัตโนมัติทุก 10 วินาที ขณะไม่มีการเป่า"],
        ],
        col_widths=[Cm(5.0), Cm(11.0)],
    )

    # ═══════════════════════ 5.5.2 Microcontroller ═══════════════════════
    add_h2(doc, "5.5.2 ไมโครคอนโทรลเลอร์ (Microcontroller)")

    add_h3(doc, "ESP32 DevKit v1 (ESP32-WROOM-32)")

    add_table(
        doc,
        header=["คุณสมบัติ", "รายละเอียด"],
        rows=[
            ["CPU", "Dual-core Xtensa LX6 32-bit @ 240 MHz"],
            ["RAM", "520 KB SRAM"],
            ["Flash Memory", "4 MB"],
            ["Wi-Fi", "802.11 b/g/n 2.4 GHz"],
            ["Bluetooth", "BLE 4.2 + Classic BT"],
            ["ADC", "12-bit SAR ADC จำนวน 18 ช่อง (ช่วง 0–3.3 V)"],
            ["GPIO", "34 pins"],
            ["I²C / SPI / UART", "รองรับครบทั้ง 3 protocol"],
            ["แรงดันไฟฟ้าใช้งาน", "3.3 V (regulated จาก USB 5V)"],
            ["Package", "DevKit V1 (30-pin หรือ 38-pin)"],
        ],
        col_widths=[Cm(5.0), Cm(11.0)],
    )

    add_para(
        doc,
        "เหตุผลที่เลือกใช้ ESP32 คือมี Wi-Fi และ ADC ในตัว ราคาไม่แพง มี ecosystem library ที่กว้างขวาง "
        "รองรับการเขียนโปรแกรมด้วย Arduino Framework และมีกำลังประมวลผลเพียงพอสำหรับการเก็บข้อมูลเซนเซอร์ "
        "แบบเรียลไทม์พร้อมส่งผ่าน MQTT ได้อย่างเสถียร",
        indent_first=True,
    )

    # ═══════════════════════ 5.5.3 Sensors ═══════════════════════
    add_h2(doc, "5.5.3 เซนเซอร์ที่ใช้ในอุปกรณ์ (Sensors)")

    add_h3(doc, "ก. Figaro TGS1820 — VOC / Breath Acetone Sensor")
    add_table(
        doc,
        header=["คุณสมบัติ", "รายละเอียด"],
        rows=[
            ["ประเภทเซนเซอร์", "MOX (Metal Oxide Semiconductor) — sensor for VOC / breath acetone"],
            ["ช่วงการวัด", "1–100 ppm (breath acetone equivalent)"],
            ["Output", "Analog voltage (0–3.3 V)"],
            ["Sensitivity", "ตอบสนองต่อการเปลี่ยนแปลง VOC ระดับ ppm"],
            ["Response time (T90)", "< 30 วินาที"],
            ["Recovery time", "30–60 วินาที (ใช้เวลาให้ baseline กลับสู่ปกติ)"],
            ["Heater voltage", "1.7 V (จ่ายจากวงจรภายในบอร์ด)"],
            ["Warm-up time", "แนะนำ ≥ 5 นาที ก่อนวัดจริงเพื่อความแม่นยำสูงสุด"],
            ["การต่อ", "Analog Input pin GPIO 34 ของ ESP32"],
            ["Sampling", "อ่านค่า 50 samples แล้ว average เพื่อลด noise"],
            ["ADC Attenuation", "11 dB (ทำให้อ่านช่วง 0–3.3 V ได้เต็มช่วง)"],
        ],
        col_widths=[Cm(5.0), Cm(11.0)],
    )
    add_para(
        doc,
        "การใช้งาน: TGS1820 ตอบสนองต่อ VOC ในลมหายใจ โดยเฉพาะ acetone ที่เป็น marker ของสภาวะ metabolic "
        "ทีมพัฒนาได้ทำ compensation ตามข้อมูลใน datasheet ของ Figaro โดยชดเชยผลของอุณหภูมิและความชื้น "
        "ผ่าน signal processing pipeline ในฝั่ง backend",
        indent_first=True,
    )

    add_h3(doc, "ข. Pressure Sensor (0–10 kPa)")
    add_table(
        doc,
        header=["คุณสมบัติ", "รายละเอียด"],
        rows=[
            ["ประเภทเซนเซอร์", "Analog pressure sensor (differential/gauge)"],
            ["ช่วงการวัด", "0–10 kPa"],
            ["Output", "Analog voltage (0–3.3 V)"],
            ["Accuracy", "±2 % full scale"],
            ["การต่อ", "Analog Input pin GPIO 32 ของ ESP32"],
            ["Sampling", "อ่านค่า 20 samples แล้ว average"],
            ["หน้าที่ในระบบ", "ตรวจจับการเป่าลมหายใจ (BLOW state) และตัดสินคุณภาพของ sample"],
        ],
        col_widths=[Cm(5.0), Cm(11.0)],
    )

    add_h3(doc, "ค. Sensirion SHT31 — Temperature & Humidity Sensor")
    add_table(
        doc,
        header=["คุณสมบัติ", "รายละเอียด"],
        rows=[
            ["ประเภทเซนเซอร์", "Digital temperature + humidity sensor"],
            ["ช่วง Temperature", "-40 °C ถึง +125 °C"],
            ["ช่วง Humidity", "0 – 100 %RH"],
            ["Accuracy Temperature", "±0.3 °C"],
            ["Accuracy Humidity", "±2 %RH"],
            ["Interface", "I²C (address 0x44)"],
            ["Voltage", "2.4 – 5.5 V"],
            ["Pin ที่ต่อกับ ESP32", "SDA = GPIO 21, SCL = GPIO 22"],
            ["หน้าที่ในระบบ", "ใช้ compensate ผลของสิ่งแวดล้อมต่อค่า VOC ของ TGS1820"],
        ],
        col_widths=[Cm(5.0), Cm(11.0)],
    )

    add_h3(doc, "ง. LED Indicator")
    add_bullet(doc, "LED บนบอร์ด (built-in) ต่อกับ GPIO 2")
    add_bullet(doc, "ใช้แสดงสถานะ: กะพริบระหว่าง Wi-Fi setup, ติดค้างเมื่อเชื่อมต่อสำเร็จ, กะพริบระหว่างการวัด")

    # ═══════════════════════ 5.5.4 Pin Mapping ═══════════════════════
    add_h2(doc, "5.5.4 การเชื่อมต่อของอุปกรณ์ (Pin Mapping)")

    add_para(
        doc,
        "ตารางแสดงการเชื่อมต่อระหว่าง ESP32 กับเซนเซอร์ต่าง ๆ ในอุปกรณ์ MetaBreath",
        indent_first=True,
    )
    add_table(
        doc,
        header=["ESP32 Pin", "ต่อกับ", "หน้าที่"],
        rows=[
            ["GPIO 34 (ADC1_CH6)", "TGS1820 Output", "Analog input สำหรับอ่านค่า VOC"],
            ["GPIO 32 (ADC1_CH4)", "Pressure Sensor Output", "Analog input สำหรับอ่านแรงดัน"],
            ["GPIO 21 (SDA)", "SHT31 SDA", "I²C data line"],
            ["GPIO 22 (SCL)", "SHT31 SCL", "I²C clock line"],
            ["GPIO 2", "LED (built-in)", "Digital output แสดงสถานะ"],
            ["3V3", "Power ให้ SHT31 + TGS1820 board", "แหล่งจ่ายไฟ 3.3 V"],
            ["GND", "Ground ทุกเซนเซอร์", "Common ground"],
            ["USB-C", "Power in", "จ่ายไฟ 5V จาก USB / Power Bank"],
        ],
        col_widths=[Cm(4.5), Cm(4.5), Cm(7.0)],
    )

    # ═══════════════════════ 5.5.5 Client Device ═══════════════════════
    add_h2(doc, "5.5.5 อุปกรณ์ฝั่งผู้ใช้ (Client Device)")

    add_para(
        doc,
        "ผู้ใช้เข้าถึง MetaBreath web application ผ่านเว็บเบราว์เซอร์ ระบบสามารถทำงานได้บนอุปกรณ์หลายประเภท "
        "โดยมีคุณสมบัติขั้นต่ำที่แนะนำดังนี้",
        indent_first=True,
    )

    add_h3(doc, "ก. สมาร์ตโฟน / แท็บเล็ต")
    add_table(
        doc,
        header=["ระบบปฏิบัติการ", "เบราว์เซอร์ที่รองรับ", "หมายเหตุ"],
        rows=[
            ["iOS 16.4 หรือใหม่กว่า", "Safari, Chrome, Edge", "รองรับ Web Push (iOS 16.4+)"],
            ["Android 10 หรือใหม่กว่า", "Chrome, Edge, Firefox", "รองรับ Web Push, Web Bluetooth"],
            ["iPadOS 16 หรือใหม่กว่า", "Safari, Chrome", "แนะนำสำหรับหน้าประวัติแนวโน้ม"],
        ],
        col_widths=[Cm(5.5), Cm(5.5), Cm(5.0)],
    )

    add_h3(doc, "ข. คอมพิวเตอร์ (PC / Notebook)")
    add_table(
        doc,
        header=["คุณสมบัติ", "ขั้นต่ำ", "แนะนำ"],
        rows=[
            ["CPU", "Dual-core 1.6 GHz", "Quad-core 2.4 GHz"],
            ["RAM", "4 GB", "8 GB"],
            ["OS", "Windows 10 / macOS 12 / Ubuntu 20.04", "Windows 11 / macOS 14 / Ubuntu 22.04"],
            ["เบราว์เซอร์", "Chrome 100+, Edge 100+, Firefox 100+, Safari 15+", "Chrome/Edge เวอร์ชันล่าสุด"],
            ["ความละเอียดหน้าจอ", "1280 × 720", "1920 × 1080 ขึ้นไป"],
            ["Internet", "≥ 5 Mbps", "≥ 10 Mbps"],
        ],
        col_widths=[Cm(4.0), Cm(6.0), Cm(6.0)],
    )

    # ═══════════════════════ 5.5.6 Server ═══════════════════════
    add_h2(doc, "5.5.6 อุปกรณ์ฝั่งเซิร์ฟเวอร์ (Server Hardware)")

    add_para(
        doc,
        "ระบบ backend, database, MQTT broker และ AI inference ทั้งหมด deploy บน server เดียวกันแบบ containerized "
        "ผ่าน Docker Compose โดยมีสเปคขั้นต่ำที่แนะนำดังนี้",
        indent_first=True,
    )

    add_table(
        doc,
        header=["ทรัพยากร", "ขั้นต่ำ (Dev/Demo)", "แนะนำ (Production Pilot)"],
        rows=[
            ["CPU", "2 vCPU", "4 vCPU"],
            ["RAM", "4 GB", "8 GB"],
            ["Storage", "40 GB SSD", "100 GB SSD"],
            ["OS", "Ubuntu 22.04 LTS", "Ubuntu 22.04 LTS / Debian 12"],
            ["Runtime", "Docker 24+, Docker Compose v2", "Docker 24+, Docker Compose v2"],
            ["Public IP", "1 IP (สำหรับ MQTT + HTTPS)", "1 IP + Domain name (SSL)"],
            ["Bandwidth", "≥ 10 Mbps", "≥ 100 Mbps"],
            ["Port ที่ต้องเปิด", "80, 443, 1883 (MQTT)", "80, 443, 1883, 9001 (MQTT WS)"],
        ],
        col_widths=[Cm(4.0), Cm(6.0), Cm(6.0)],
    )

    add_para(
        doc,
        "ในการพัฒนาและทดสอบระยะแรก ทีมใช้ VPS ระดับ entry-level (2 vCPU / 4 GB RAM / 60 GB SSD) "
        "ที่มี public IP สำหรับให้ ESP32 หลาย ๆ เครื่องส่งข้อมูลผ่าน MQTT ได้พร้อมกัน และมีการตั้งค่า "
        "Nginx + Let's Encrypt สำหรับ TLS termination",
        indent_first=True,
    )

    # ═══════════════════════ 5.5.7 Network ═══════════════════════
    add_h2(doc, "5.5.7 โครงสร้างเครือข่าย (Network Requirements)")

    add_bullet(doc, "Wi-Fi 2.4 GHz (802.11 b/g/n) — สำหรับอุปกรณ์ MetaBreath (ESP32 ไม่รองรับ 5 GHz)")
    add_bullet(doc, "SSID ต้องไม่มี captive portal ของโรงแรม (ต้องเป็น Wi-Fi ปกติ)")
    add_bullet(doc, "MQTT — TCP port 1883 (ไม่ block outbound)")
    add_bullet(doc, "HTTPS — TCP port 443 สำหรับ web application")
    add_bullet(doc, "WebSocket — TCP port 443 (wss://) หรือ 9001 สำหรับ MQTT WebSocket")
    add_bullet(doc, "DNS resolution สำหรับ MQTT broker domain (metabreath.duckdns.org)")
    add_bullet(doc, "อุปกรณ์รองรับ Wi-Fi หลายสล็อต (สูงสุด 3 SSID) ผ่าน WiFiMulti — เปลี่ยนสถานที่ได้อัตโนมัติ")

    # ═══════════════════════ 5.5.8 สภาพแวดล้อม ═══════════════════════
    add_h2(doc, "5.5.8 สภาพแวดล้อมการใช้งาน (Operating Environment)")

    add_bullet(doc, "อุณหภูมิใช้งาน: 15 – 35 °C")
    add_bullet(doc, "ความชื้นสัมพัทธ์: 30 – 80 %RH")
    add_bullet(doc, "ไม่ควรใช้ในบริเวณที่มี VOC สูงผิดปกติ (เช่น ห้องพ่นสี หรือใกล้ solvent) เพราะ TGS1820 จะอิ่มตัว")
    add_bullet(doc, "แนะนำให้ตั้งเครื่องในบริเวณที่อากาศไหลเวียนดีเป็นเวลา ≥ 5 นาที ก่อนใช้งาน")
    add_bullet(doc, "ไม่แนะนำให้ใช้กลางแดดจัด หรือในบริเวณที่มีความชื้นสูง (เช่น ห้องน้ำ)")

    # ═══════════════════════ 5.5.9 สรุป ═══════════════════════
    add_h2(doc, "5.5.9 สรุปคุณลักษณะโดยรวม")

    add_para(
        doc,
        "อุปกรณ์ MetaBreath ถูกออกแบบให้เป็นอุปกรณ์ต้นแบบขนาดพกพา ราคาไม่แพง และใช้ชิ้นส่วนอิเล็กทรอนิกส์ "
        "ที่หาซื้อได้ทั่วไป ทำให้สามารถทำซ้ำและขยายผลได้ง่าย ในขณะที่ฝั่ง server และ client "
        "สามารถใช้อุปกรณ์มาตรฐาน (VPS + smartphone/PC ทั่วไป) ทำให้ระบบทั้งหมดเข้าถึงได้ง่ายและ "
        "เหมาะกับการต่อยอดไปสู่การใช้งานจริงในอนาคต",
        indent_first=True,
    )

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
